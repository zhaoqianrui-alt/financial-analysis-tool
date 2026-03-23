import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import anthropic
from dotenv import load_dotenv
import os

load_dotenv()
API_KEY = os.getenv("ANTHROPIC_API_KEY")
if not API_KEY:
    st.error("❗ ANTHROPIC_API_KEY not found. Please add it to your .env file.")
    st.stop()
client = anthropic.Anthropic(api_key=API_KEY)

# ============================================================
# 数据获取函数
# ============================================================

def is_a_share(code):
    return code.strip().isdigit()

def process_us_data(code):
    import yfinance as yf, time
    for attempt in range(3):
        try:
            ticker = yf.Ticker(code.upper())
            info = ticker.info or {}
            company_name = info.get("longName") or info.get("shortName") or code.upper()
            income = ticker.financials
            balance = ticker.balance_sheet
            break
        except Exception as e:
            if "429" in str(e) or "Too Many" in str(e) or "rate" in str(e).lower():
                if attempt < 2:
                    time.sleep(3)
                    continue
            raise
    rows = []
    for date in income.columns[:5]:
        year = date.year
        try:
            revenue = income.loc["Total Revenue", date] / 1e8
            net_income = income.loc["Net Income", date] / 1e8
            total_assets = balance.loc["Total Assets", date] / 1e8
            equity = None
            for eq_key in ["Stockholders Equity", "Total Equity Gross Minority Interest", "Common Stock Equity"]:
                try:
                    if eq_key in balance.index and pd.notna(balance.loc[eq_key, date]):
                        equity = float(balance.loc[eq_key, date]) / 1e8
                        break
                except Exception:
                    continue
            if equity is None:
                equity = total_assets - total_liabilities
            roe = round(net_income / equity * 100, 2) if equity and equity != 0 else None
            net_margin = round(net_income / revenue * 100, 2) if revenue else None
            debt_ratio = round(total_liabilities / total_assets * 100, 2) if total_assets else None
            rows.append({
                "年份": year,
                "营业收入(亿美元)": round(revenue, 2),
                "净利润(亿美元)": round(net_income, 2),
                "总负债(亿美元)": round(total_liabilities, 2),
                "总资产(亿美元)": round(total_assets, 2),
                "净资产(亿美元)": round(equity, 2),
                "净利率%": net_margin,
                "资产负债率%": debt_ratio,
                "ROE%": roe,
            })
        except KeyError:
            continue
    df = pd.DataFrame(rows).dropna(subset=["营业收入(亿美元)","净利润(亿美元)"]).sort_values("年份").reset_index(drop=True)
    return df, company_name

def a_share_to_yahoo(code):
    code = code.strip()
    if code.startswith("6"):
        return code + ".SS"
    elif code.startswith("0") or code.startswith("3"):
        return code + ".SZ"
    elif code.startswith("8") or code.startswith("4"):
        return code + ".BJ"
    return code + ".SS"
def process_a_share_data(code):
    import yfinance as yf
    import time
    yahoo_code = a_share_to_yahoo(code)
    # 最多重试3次，处理速率限制
    for attempt in range(3):
        try:
            ticker = yf.Ticker(yahoo_code)
            try:
                info = ticker.info or {}
                company_name = info.get("longName") or info.get("shortName") or code
            except Exception:
                company_name = code
            income  = ticker.financials
            balance = ticker.balance_sheet
            break  # 成功则跳出重试循环
        except Exception as e:
            if "429" in str(e) or "Too Many" in str(e) or "rate" in str(e).lower():
                if attempt < 2:
                    time.sleep(3)  # 等3秒再试
                    continue
            raise ValueError(f"数据拉取失败（{yahoo_code}）：{e}")
    # 处理转置情况：yfinance有时返回行列互换的DataFrame
    if income is not None and not income.empty:
        # 正常情况：行是财务科目，列是日期
        # 如果列是字符串（科目名）而不是日期，说明需要转置
        if len(income.columns) > 0 and isinstance(income.columns[0], str):
            income = income.T
    if balance is not None and not balance.empty:
        if len(balance.columns) > 0 and isinstance(balance.columns[0], str):
            balance = balance.T
    if income is None or income.empty:
       raise ValueError(
    f"未找到 {yahoo_code} 的财务数据。\n\n"
    f"雅虎财经对部分A股财务数据覆盖不完整，建议试试：\n"
    f"• 600519（贵州茅台）\n"
    f"• 000858（五粮液）\n"
    f"• 601318（中国平安）\n"
    f"• 300750（宁德时代）\n"
    f"• 688017（绿的谐波）\n"
    f"• 600036（招商银行）"
)
    def safe_get(df, keys, date):
        for k in keys:
            try:
                if k in df.index and pd.notna(df.loc[k, date]):
                    return float(df.loc[k, date])
            except Exception:
                continue
        return None
    rows = []
    for date in income.columns[:5]:
        year = date.year
        try:
            revenue    = safe_get(income,  ["Total Revenue", "Operating Revenue"], date)
            net_income = safe_get(income,  ["Net Income", "Net Income Common Stockholders"], date)
            tot_assets = safe_get(balance, ["Total Assets"], date)
            equity     = safe_get(balance, ["Stockholders Equity", "Total Equity Gross Minority Interest", "Common Stock Equity"], date)
            if revenue is None or net_income is None:
                continue
            net_margin = round(net_income / revenue * 100, 2) if revenue else None
            debt_ratio = round((tot_assets - equity) / tot_assets * 100, 2) if (tot_assets and equity) else None
            roe        = round(net_income / equity * 100, 2) if (equity and equity != 0) else None
            rows.append({
                "年份":           year,
                "营业收入(亿元)": round(revenue    / 1e8, 2),
                "净利润(亿元)":   round(net_income / 1e8, 2),
                "净利率%":        net_margin,
                "资产负债率%":    debt_ratio,
                "ROE%":           roe,
                "总资产(亿元)":   round(tot_assets / 1e8, 2) if tot_assets else None,
                "净资产(亿元)":   round(equity     / 1e8, 2) if equity     else None,
            })
        except Exception:
            continue
    if not rows:
        raise ValueError(f"获取到数据但无法解析财务指标，股票代码：{yahoo_code}")
    df = pd.DataFrame(rows).sort_values("年份").reset_index(drop=True)
    df["收入增速%"] = df["营业收入(亿元)"].pct_change() * 100
    df["收入增速%"] = df["收入增速%"].apply(lambda x: round(x, 2) if pd.notna(x) else None)
    return df, company_name

# ============================================================
# 预测建模函数
# 背景知识：利润表结构
# 营业收入
# - 营业成本
# = 毛利润
# - 销售费用、管理费用、研发费用（统称"期间费用"）
# = 营业利润（EBIT）
# - 所得税（税率 × 税前利润）
# = 净利润
# ============================================================

def build_income_statement(
    base_revenue,        # 基准年营业收入
    driver_names,        # 收入驱动因子名称列表，如["销量(万辆)", "单车价值(元)"]
    driver_base,         # 基准年各因子数值列表
    driver_forecasts,    # 未来3年各因子预测值，格式：[[Y1因子1, Y1因子2], [Y2...], [Y3...]]
    gross_margin,        # 毛利率假设（%）
    expense_ratio,       # 期间费用率假设（%，占收入比）
    tax_rate,            # 所得税率假设（%）
    base_year,           # 基准年份
    hist_df=None,        # 历史数据（可选，用于展示历史+预测对比）
    rev_col=None,        # 历史数据中收入列名
    net_col=None,        # 历史数据中净利润列名
):
    """
    构建预测利润表
    
    收入预测逻辑：
    - 如果只有1个因子：收入 = 基准收入 × (因子预测值 / 因子基准值)
    - 如果有2个因子：收入 = 因子1预测值 × 因子2预测值
      （适合"销量×单价"这类乘积型驱动）
    - 如果有3个因子：收入 = 因子1 × 因子2 × 因子3
    
    这是真实财务建模中最常用的收入预测方法
    """
    forecast_rows = []
    
    for i, year_drivers in enumerate(driver_forecasts):
        year = base_year + i + 1
        
        # 根据因子数量计算预测收入
        if len(driver_names) == 1:
            # 单因子：按比例缩放
            forecast_rev = base_revenue * (year_drivers[0] / driver_base[0])
        elif len(driver_names) == 2:
            # 双因子：两个因子相乘（如销量×单价）
            # 但单位可能不匹配，所以用"相对基准年的变化倍数"来算
            ratio = (year_drivers[0] / driver_base[0]) * (year_drivers[1] / driver_base[1])
            forecast_rev = base_revenue * ratio
        else:
            # 三因子
            ratio = (year_drivers[0] / driver_base[0]) * (year_drivers[1] / driver_base[1]) * (year_drivers[2] / driver_base[2])
            forecast_rev = base_revenue * ratio
        
        # 利润表推导
        # 背景知识：毛利润 = 收入 × 毛利率
        gross_profit = forecast_rev * gross_margin / 100
        # 期间费用 = 收入 × 费用率
        operating_expense = forecast_rev * expense_ratio / 100
        # 营业利润（EBIT）= 毛利润 - 期间费用
        ebit = gross_profit - operating_expense
        # 净利润 = 营业利润 × (1 - 税率)
        net_income = ebit * (1 - tax_rate / 100)
        # 净利率 = 净利润 / 收入
        net_margin = net_income / forecast_rev * 100
        
        forecast_rows.append({
            "年份": str(year),
            "营业收入": round(forecast_rev, 2),
            "毛利润": round(gross_profit, 2),
            "营业利润(EBIT)": round(ebit, 2),
            "净利润": round(net_income, 2),
            "毛利率%": round(gross_margin, 1),
            "净利率%": round(net_margin, 1),
            "类型": "预测",
        })
    
    forecast_df = pd.DataFrame(forecast_rows)
    
    # 如果有历史数据，拼接在一起展示
    if hist_df is not None and rev_col and net_col:
        hist_rows = []
        for _, row in hist_df.iterrows():
            rev = row[rev_col]
            net = row[net_col]
            if pd.notna(rev) and pd.notna(net) and rev != 0:
                hist_rows.append({
                    "年份": str(int(row["年份"])),
                    "营业收入": rev,
                    "毛利润": None,
                    "营业利润(EBIT)": None,
                    "净利润": net,
                    "毛利率%": None,
                    "净利率%": round(net / rev * 100, 1),
                    "类型": "历史",
                })
        if hist_rows:
            hist_display = pd.DataFrame(hist_rows)
            combined = pd.concat([hist_display, forecast_df], ignore_index=True)
            return forecast_df, combined
    
    return forecast_df, forecast_df


# ============================================================
# AUTO-PILOT：AI自动推断所有参数并一键跑完全流程
# ============================================================

def run_autopilot(df, cname, is_us):
    """
    输入：历史财务数据 DataFrame + 公司名
    输出：所有模块的计算结果（写入 session_state）
    
    核心思路：
    1. 用 Claude 读取历史数据，推断合理的假设参数
    2. 用这些参数调用现有的建模函数
    3. 把结果存入 session_state，各 tab 直接读取展示
    """
    import numpy as np
    import json

    # ── 防护：验证输入数据 ──────────────────────────────────────
    if df is None:
        raise ValueError("No data loaded. Please fetch company data first.")
    if not hasattr(df, 'columns'):
        raise ValueError("Data format error — please re-fetch the data.")
    if df.empty:
        raise ValueError("Data is empty. Please try fetching the data again.")
    if "年份" not in df.columns:
        raise ValueError(f"Unexpected data format (columns: {list(df.columns)}). Please re-fetch.")

    # ── Step 1：整理历史数据摘要，发给 AI ──────────────────────
    if is_us:
        rev_col  = "营业收入(亿美元)"
        net_col  = "净利润(亿美元)"
        asset_col = "总资产(亿美元)"
        liab_col  = "总负债(亿美元)"
        currency  = "USD bn"
    else:
        rev_col  = "营业收入(亿元)" if "营业收入(亿元)" in df.columns else None
        net_col  = "净利润(亿元)"   if "净利润(亿元)"  in df.columns else None
        asset_col = None
        liab_col  = None
        currency  = "CNY bn"

    # 构建发给 AI 的数据摘要
    has_abs = rev_col and rev_col in df.columns and net_col and net_col in df.columns

    if has_abs:
        latest = df.dropna(subset=[rev_col]).iloc[-1]
        base_rev  = float(latest[rev_col])
        base_year = int(latest["年份"])
        hist_summary = "\n".join(
            f"  {int(r['年份'])}: Revenue {r[rev_col]} {currency}, Net Income {r[net_col]} {currency}, "
            f"Net Margin {round(r[net_col]/r[rev_col]*100,1) if r[rev_col] else 'N/A'}%"
            for _, r in df.dropna(subset=[rev_col]).iterrows()
        )
    else:
        # A股只有比率数据
        base_rev  = 100.0   # 占位，用户可在各tab修改
        base_year = int(df["年份"].max()) if "年份" in df.columns else 2024
        hist_summary = "\n".join(
            f"  {int(r['年份'])}: Net Margin {r.get('净利率%','N/A')}%, "
            f"Debt Ratio {r.get('资产负债率%','N/A')}%, ROE {r.get('ROE%','N/A')}%"
            for _, r in df.iterrows()
        )

    # 获取资产负债表数据（用于推断shares/net_debt等）
    if asset_col and asset_col in df.columns:
        latest_assets  = float(df.dropna(subset=[asset_col]).iloc[-1][asset_col])
        latest_liab    = float(df.dropna(subset=[liab_col]).iloc[-1][liab_col])  if liab_col and liab_col in df.columns else latest_assets * 0.5
    else:
        latest_assets  = base_rev * 2.5 if has_abs else 250.0
        latest_liab    = latest_assets * 0.5

    equity_est   = latest_assets - latest_liab
    net_debt_est = latest_liab - (latest_assets * 0.15)   # 粗估净债务
    shares_est   = round(equity_est / 20, 1)              # 粗估股数（假设BVPS~20）
    shares_est   = max(shares_est, 0.5)

    # ── Step 2：调用 Claude，让它推断所有参数 ─────────────────
    prompt = f"""You are a senior equity analyst. Based on the following historical financial data for {cname}, 
infer reasonable forward-looking assumptions for a 3-year financial model.

HISTORICAL DATA ({currency}):
{hist_summary}

TASK: Return ONLY a valid JSON object (no markdown, no explanation) with these exact keys:

{{
  "revenue_growth_y1": <% e.g. 8.0>,
  "revenue_growth_y2": <% e.g. 7.0>,
  "revenue_growth_y3": <% e.g. 6.0>,
  "gross_margin": <% e.g. 42.0>,
  "expense_ratio": <% e.g. 18.0>,
  "tax_rate": <% e.g. 25.0>,
  "depreciation_rate": <% of PP&E, e.g. 10.0>,
  "capex_rate": <% of revenue, e.g. 5.0>,
  "ar_days": <integer, e.g. 55>,
  "inv_days": <integer, e.g. 40>,
  "ap_days": <integer, e.g. 45>,
  "dividend_payout": <% e.g. 30.0>,
  "wacc": <% e.g. 9.0>,
  "terminal_growth": <% e.g. 3.0>,
  "pe_bear": <e.g. 18.0>,
  "pe_base": <e.g. 25.0>,
  "pe_bull": <e.g. 32.0>,
  "pb_bear": <e.g. 2.0>,
  "pb_base": <e.g. 3.5>,
  "pb_bull": <e.g. 5.0>,
  "reasoning": "<2-3 sentences explaining your key assumptions>"
}}

Base your estimates on the historical trends above. Be conservative but realistic.
If data is limited (ratio-only), use industry averages appropriate for the sector."""

    msg = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=1000,
        messages=[{"role": "user", "content": prompt}]
    )

    raw = msg.content[0].text.strip()
    # 去掉可能的 markdown 代码块
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    params = json.loads(raw.strip())

    reasoning = params.get("reasoning", "AI-generated assumptions based on historical data.")

    # ── Step 3：用 AI 参数跑 Forecast ─────────────────────────
    g1 = params["revenue_growth_y1"] / 100
    g2 = params["revenue_growth_y2"] / 100
    g3 = params["revenue_growth_y3"] / 100

    rev_y1 = base_rev * (1 + g1)
    rev_y2 = rev_y1  * (1 + g2)
    rev_y3 = rev_y2  * (1 + g3)

    driver_names     = ["Revenue Growth Rate (%)"]
    driver_base      = [100.0]
    driver_forecasts = [
        [100 * (1 + g1)],
        [100 * (1 + g1) * (1 + g2)],
        [100 * (1 + g1) * (1 + g2) * (1 + g3)],
    ]

    forecast_df, combined_df = build_income_statement(
        base_revenue   = base_rev,
        driver_names   = driver_names,
        driver_base    = driver_base,
        driver_forecasts = driver_forecasts,
        gross_margin   = params["gross_margin"],
        expense_ratio  = params["expense_ratio"],
        tax_rate       = params["tax_rate"],
        base_year      = base_year,
        hist_df        = df if has_abs else None,
        rev_col        = rev_col,
        net_col        = net_col,
    )

    st.session_state.forecast_df     = forecast_df
    st.session_state.combined_df     = combined_df
    st.session_state.forecast_params = {
        "gross_margin":  params["gross_margin"],
        "expense_ratio": params["expense_ratio"],
        "tax_rate":      params["tax_rate"],
        "cname":         cname,
    }

    # ── Step 4：用预测结果跑三表 ──────────────────────────────
    dep_rate  = params["depreciation_rate"]
    capex_rate = params["capex_rate"]
    ar_days   = int(params["ar_days"])
    inv_days  = int(params["inv_days"])
    ap_days   = int(params["ap_days"])
    div_pct   = params["dividend_payout"]

    # 估算基年资产负债表
    base_cash_val  = latest_assets * 0.15
    base_fa_val    = latest_assets * 0.40
    base_eq_val    = equity_est
    base_debt_val  = latest_liab * 0.6

    income_rows, cashflow_rows, balance_rows = [], [], []
    prev_cash, prev_fa, prev_eq, prev_debt = base_cash_val, base_fa_val, base_eq_val, base_debt_val

    for _, row in forecast_df.iterrows():
        revenue    = row["营业收入"]
        net_income = row["净利润"]
        ebit       = row["营业利润(EBIT)"]
        year       = row["年份"]

        capex        = revenue * capex_rate / 100
        depreciation = prev_fa * dep_rate / 100
        dividends    = net_income * div_pct / 100
        ar  = revenue * ar_days  / 365
        inv = revenue * inv_days / 365
        ap  = revenue * ap_days  / 365
        nwc = ar + inv - ap
        cfo = net_income + depreciation - (nwc * 0.1)
        cfi = -capex
        cff = -dividends
        fcf = cfo + cfi
        net_cash_change = cfo + cfi + cff

        end_cash   = prev_cash + net_cash_change
        end_fa     = prev_fa + capex - depreciation
        total_assets = end_cash + end_fa + ar + inv
        end_equity = prev_eq + (net_income - dividends)
        total_liab = ap + prev_debt
        check = total_assets - (total_liab + end_equity)

        income_rows.append({"Year": year, "Revenue": round(revenue,2), "D&A": round(depreciation,2), "EBIT": round(ebit,2), "Net Income": round(net_income,2), "Net Margin%": round(row["净利率%"],1)})
        cashflow_rows.append({"Year": year, "CFO": round(cfo,2), "Capex": round(-capex,2), "CFI": round(cfi,2), "CFF": round(cff,2), "FCF": round(fcf,2), "Net Change": round(net_cash_change,2)})
        balance_rows.append({"Year": year, "Cash": round(end_cash,2), "Receivables": round(ar,2), "Inventory": round(inv,2), "PP&E": round(end_fa,2), "Total Assets": round(total_assets,2), "Payables": round(ap,2), "Debt": round(prev_debt,2), "Total Liabilities": round(total_liab,2), "Equity": round(end_equity,2), "Check": round(check,2)})

        prev_cash, prev_fa, prev_eq = end_cash, end_fa, end_equity

    st.session_state.income_rows      = income_rows
    st.session_state.cashflow_rows    = cashflow_rows
    st.session_state.balance_rows     = balance_rows
    st.session_state.three_table_params = {
        "depreciation_rate": dep_rate, "capex_rate": capex_rate,
        "ar_days": ar_days, "inv_days": inv_days, "ap_days": ap_days,
        "dividend_payout": div_pct, "cname": cname,
    }

    # ── Step 5：跑 DCF ─────────────────────────────────────────
    fcf_list = [r["FCF"] for r in cashflow_rows]
    wacc     = params["wacc"]
    tg       = params["terminal_growth"]

    def dcf_val(fcf_list, wacc_pct, tg_pct, shares, nd):
        w, g = wacc_pct/100, tg_pct/100
        if w <= g:
            return None
        pv_fcf = sum(f/(1+w)**(i+1) for i,f in enumerate(fcf_list))
        tv     = fcf_list[-1]*(1+g)/(w-g)
        pv_tv  = tv/(1+w)**len(fcf_list)
        ev     = pv_fcf + pv_tv
        eq_val = ev - nd
        pps    = eq_val/shares if shares > 0 else 0
        return {"pv_fcf": round(pv_fcf,1), "pv_terminal": round(pv_tv,1),
                "ev": round(ev,1), "equity_value": round(eq_val,1),
                "price_per_share": round(pps,2),
                "terminal_pct": round(pv_tv/ev*100,1) if ev > 0 else 0}

    base_dcf = dcf_val(fcf_list, wacc, tg, shares_est, net_debt_est)
    if base_dcf is None:
        base_dcf = {"pv_fcf":0,"pv_terminal":0,"ev":0,"equity_value":0,"price_per_share":0,"terminal_pct":0}

    wacc_steps = np.arange(max(wacc-2,1), wacc+3, 1.0)
    tg_steps   = np.arange(max(tg-1,0.5), tg+2, 0.5)
    all_prices = []
    for w in wacc_steps:
        for g in tg_steps:
            r = dcf_val(fcf_list, w, g, shares_est, net_debt_est)
            if r:
                all_prices.append(r["price_per_share"])

    st.session_state.dcf_result = base_dcf
    st.session_state.dcf_params = {
        "wacc": wacc, "terminal_growth": tg,
        "shares": shares_est, "net_debt": net_debt_est,
        "price_low":  round(min(all_prices),2) if all_prices else 0,
        "price_high": round(max(all_prices),2) if all_prices else 0,
        "cname": cname,
    }

    # ── Step 6：跑 PE/PB ──────────────────────────────────────
    last_ni  = income_rows[-1]["Net Income"]
    last_eq  = balance_rows[-1]["Equity"]
    eps      = last_ni / shares_est  if shares_est > 0 else 0
    bvps     = last_eq / shares_est  if shares_est > 0 else 0

    st.session_state.pepb_result = {
        "eps": round(eps,2), "bvps": round(bvps,2),
        "pe_bear": params["pe_bear"], "pe_base": params["pe_base"], "pe_bull": params["pe_bull"],
        "pb_bear": params["pb_bear"], "pb_base": params["pb_base"], "pb_bull": params["pb_bull"],
        "pe_price_bear": round(eps * params["pe_bear"],2),
        "pe_price_base": round(eps * params["pe_base"],2),
        "pe_price_bull": round(eps * params["pe_bull"],2),
        "pb_price_bear": round(bvps * params["pb_bear"],2),
        "pb_price_base": round(bvps * params["pb_base"],2),
        "pb_price_bull": round(bvps * params["pb_bull"],2),
    }

    return params, reasoning, base_dcf, all_prices


# ============================================================
# 页面主体 - 标签页布局（每个Tab独立可用）
# ============================================================

st.title("📊 AI Financial Analysis Tool")
st.markdown("""
**Understand any company's financial health in minutes — no finance background needed.**

This tool walks you through four steps: fetch real data → build a forecast → model the full financials → estimate fair value.
Each step is independent: you can jump straight to Valuation if you already have your own numbers.
""")

# session_state 初始化
for key, default in {
    "fetched_df": None,
    "fetched_company": None,
    "is_us_stock": False,
    "forecast_df": None,
    "combined_df": None,
    "forecast_params": None,
    "cashflow_rows": None,
    "income_rows": None,
    "balance_rows": None,
    "three_table_params": None,
    "dcf_result": None,
    "dcf_params": None,
    "pepb_result": None,
    "ai_report_text": None,
    "autopilot_done": False,
    "autopilot_params": None,
    "autopilot_reasoning": None,
}.items():
    if key not in st.session_state:
        st.session_state[key] = default

tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "📊 Data",
    "📈 Forecast",
    "📑 3-Statement",
    "💰 Valuation",
    "📄 Export Report",
    "🔍 地雷扫描",
])

# ============================================================
# TAB 1: DATA
# ============================================================
with tab1:
    # ── 封面入口选择 ──────────────────────────────────────────
    if "entry_mode" not in st.session_state:
        st.session_state.entry_mode = None

    if st.session_state.entry_mode is None:
        st.markdown("""
<div style="text-align:center;padding:40px 0 20px;">
  <h2 style="color:#2E5C8A;">Where would you like to start?</h2>
  <p style="color:#555;font-size:1.05em;">Choose the path that fits your situation.</p>
</div>""", unsafe_allow_html=True)

        col_l, col_r = st.columns(2, gap="large")
        with col_l:
            st.markdown("""
<div style="background:#EBF2FA;border:2px solid #4472C4;border-radius:14px;padding:30px 24px;text-align:center;">
  <div style="font-size:2.5em;">🔍</div>
  <h3 style="color:#2E5C8A;margin:10px 0 6px;">I have a stock ticker</h3>
  <p style="color:#555;font-size:0.95em;">Enter a ticker symbol (e.g. AAPL, TSLA) and we'll automatically pull 5 years of real financial data from the market.</p>
  <p style="color:#777;font-size:0.85em;">Best for: investors researching public companies</p>
</div>""", unsafe_allow_html=True)
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("Start with a ticker →", use_container_width=True, key="entry_ticker"):
                st.session_state.entry_mode = "ticker"
                st.rerun()

        with col_r:
            st.markdown("""
<div style="background:#F0FAF0;border:2px solid #70AD47;border-radius:14px;padding:30px 24px;text-align:center;">
  <div style="font-size:2.5em;">📝</div>
  <h3 style="color:#2E5C8A;margin:10px 0 6px;">I'll enter my own numbers</h3>
  <p style="color:#555;font-size:0.95em;">You already have financial data — from an annual report, spreadsheet, or your own research. Skip straight to building forecasts.</p>
  <p style="color:#777;font-size:0.85em;">Best for: analysts with existing data</p>
</div>""", unsafe_allow_html=True)
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("Enter my own numbers →", use_container_width=True, key="entry_manual"):
                st.session_state.entry_mode = "manual"
                st.rerun()

    # ── 入口 A：股票代码 ──────────────────────────────────────
    elif st.session_state.entry_mode == "ticker":
        if st.button("← Back", key="back_ticker"):
            st.session_state.entry_mode = None
            st.rerun()

        st.header("📊 Fetch Company Data")
        st.write("Enter a ticker symbol — we'll look up the company name and pull all the numbers automatically.")
        stock_code = st.text_input(
            "Ticker Symbol",
            placeholder="US stock: AAPL, TSLA, NVDA  ·  A-share (China): 600031, 688017",
            help="A ticker is the short code used to identify a stock on an exchange. For example, Apple's ticker is AAPL."
        )

        if stock_code:
            if st.button("🔍 Fetch Data"):
                with st.spinner("Looking up company and fetching data..."):
                    try:
                        if is_a_share(stock_code):
                            df, auto_cname_raw = process_a_share_data(stock_code)
                            auto_cname = auto_cname_raw or stock_code # A股暂无自动名称
                            st.session_state.fetched_df = df
                            st.session_state.fetched_company = auto_cname
                            st.session_state.is_us_stock = False
                            st.success(f"✅ Data fetched for {auto_cname}!")
                        else:
                            df, auto_cname = process_us_data(stock_code)
                            st.session_state.fetched_df = df
                            st.session_state.fetched_company = auto_cname
                            st.session_state.is_us_stock = True
                            if df.empty:
                                st.warning("No data found. Please double-check the ticker symbol.")
                            else:
                                df["收入增速%"] = round(df["营业收入(亿美元)"].pct_change() * 100, 2)
                                st.session_state.fetched_df = df
                                st.success(f"✅ Found **{auto_cname}** — data loaded!")
                    except Exception as e:
                        st.error(f"Error: {e}")

        if st.session_state.fetched_df is not None:
            df = st.session_state.fetched_df
            cname = st.session_state.fetched_company
            is_us = st.session_state.is_us_stock
            latest = df.iloc[-1]

            st.subheader(f"{cname} — {'Historical Financials (USD bn)' if is_us else '历史财务指标（A股）'}")
            st.dataframe(df)
            def metric_card(label, value, good_range, explanation, is_good):
                color = "#d4edda" if is_good else "#f8d7da"
                icon = "✅" if is_good else "⚠️"
                st.markdown(f"""
<div style="background:{color};padding:14px 18px;border-radius:10px;margin-bottom:10px;">
<b>{icon} {label}: {value}</b><br>
<span style="font-size:0.9em;color:#333;">{explanation}</span><br>
<span style="font-size:0.82em;color:#555;">Healthy range: {good_range}</span>
</div>""", unsafe_allow_html=True)

            col_a, col_b = st.columns(2)
            with col_a:
                if pd.notna(latest.get("净利率%")):
                    nm = latest["净利率%"]
                    metric_card("Net Margin", f"{nm}%", "10–30% for most industries",
                        f"Out of every $100 in revenue, {cname} keeps ${round(nm,1)} as profit after all costs and taxes. {'Strong profitability.' if nm > 15 else 'Below average — watch cost trends.' if nm > 0 else 'The company is currently losing money.'}",
                        nm > 10)
                if pd.notna(latest.get("收入增速%")):
                    rg = latest["收入增速%"]
                    metric_card("Revenue Growth", f"{rg}%", ">10% = fast-growing, 0–10% = stable",
                        f"Revenue {'grew' if rg >= 0 else 'shrank'} by {abs(round(rg,1))}% last year. {'Strong momentum.' if rg > 15 else 'Steady growth.' if rg > 0 else 'Declining top line — needs investigation.'}",
                        rg > 5)
            with col_b:
                if pd.notna(latest.get("资产负债率%")):
                    dr = latest["资产负债率%"]
                    metric_card("Debt Ratio", f"{dr}%", "Below 60% is generally safe",
                        f"{round(dr,1)}% of {cname}'s assets are funded by debt. {'Conservative balance sheet.' if dr < 40 else 'Moderate leverage.' if dr < 60 else 'High leverage — financial risk is elevated.'}",
                        dr < 60)
                if pd.notna(latest.get("ROE%")):
                    roe = latest["ROE%"]
                    metric_card("Return on Equity (ROE)", f"{roe}%", "Above 15% is considered strong",
                        f"For every $100 shareholders invested, the company earned ${round(roe,1)}. {'Excellent capital efficiency.' if roe > 20 else 'Decent returns.' if roe > 10 else 'Low returns on equity.'}",
                        roe > 15)

            chart_cols = [c for c in ["净利率%", "资产负债率%", "ROE%", "收入增速%"] if c in df.columns]
            if chart_cols:
                st.subheader("📈 5-Year Trend")
                st.write("Trends often matter more than a single year's number. Look for consistent improvement or warning signs.")
                chart_df = df[["年份"] + chart_cols].copy()
                chart_df["年份"] = chart_df["年份"].astype(str)
                fig = px.line(chart_df.melt(id_vars="年份", var_name="Metric", value_name="Value"),
                    x="年份", y="Value", color="Metric", markers=True,
                    title=f"{cname} — Key Ratios (5-Year)")
                fig.update_layout(hovermode="x unified", plot_bgcolor="white",
                    legend=dict(orientation="h", yanchor="bottom", y=1.02))
                st.plotly_chart(fig, use_container_width=True)

            # ── 🚀 AUTO-PILOT 按钮（只在数据加载后显示）────────────
            st.divider()
            st.markdown("""
<div style="background:linear-gradient(135deg,#1a3a5c,#2E5C8A);padding:24px 28px;border-radius:14px;margin:10px 0;">
  <h3 style="color:white;margin:0 0 8px;">🚀 Auto-Pilot Mode</h3>
  <p style="color:#cce0ff;margin:0;font-size:0.98em;">
    Let AI read the historical data above and automatically set all assumptions —
    then run the full Forecast → 3-Statement → DCF → PE/PB pipeline in one click.
    You can review and adjust anything afterwards.
  </p>
</div>""", unsafe_allow_html=True)

            col_ap1, col_ap2 = st.columns([1, 2])
            with col_ap1:
                run_ap = st.button("🚀 Run Auto-Pilot", use_container_width=True, type="primary")
            with col_ap2:
                st.caption("⏱ Takes ~15 seconds · Uses AI to infer assumptions · All results editable afterwards")

            if run_ap:
                # 防护检查：确保数据已加载
                if st.session_state.fetched_df is None or st.session_state.fetched_df.empty:
                    st.error("⚠️ Please fetch company data first before running Auto-Pilot.")
                else:
                    with st.spinner("🤖 AI is reading the data and building your full financial model..."):
                        try:
                            ap_params, ap_reasoning, ap_dcf, ap_prices = run_autopilot(
                                df=st.session_state.fetched_df,
                                cname=st.session_state.fetched_company,
                                is_us=st.session_state.is_us_stock,
                            )
                            st.session_state.autopilot_params    = ap_params
                            st.session_state.autopilot_reasoning = ap_reasoning
                            st.session_state.autopilot_done      = True
                            st.success("✅ Auto-Pilot complete! All modules have been populated.")
                            st.rerun()
                        except Exception as e:
                            import traceback
                            st.error(f"Auto-Pilot failed: {e}")
                            with st.expander("🔍 Error details"):
                                st.code(traceback.format_exc())

        if st.session_state.get("autopilot_done"):
            ap_params = st.session_state.autopilot_params
            ap_reason = st.session_state.autopilot_reasoning

            with st.expander("📋 View AI-generated assumptions", expanded=True):
                st.info(f"💡 **AI Reasoning:** {ap_reason}")

                c1, c2, c3 = st.columns(3)
                with c1:
                    st.markdown("**📈 Revenue Forecast**")
                    st.metric("Year 1 growth", f"{ap_params['revenue_growth_y1']}%")
                    st.metric("Year 2 growth", f"{ap_params['revenue_growth_y2']}%")
                    st.metric("Year 3 growth", f"{ap_params['revenue_growth_y3']}%")
                with c2:
                    st.markdown("**💼 P&L Assumptions**")
                    st.metric("Gross Margin",  f"{ap_params['gross_margin']}%")
                    st.metric("Opex Ratio",    f"{ap_params['expense_ratio']}%")
                    st.metric("Tax Rate",      f"{ap_params['tax_rate']}%")
                with c3:
                    st.markdown("**💰 Valuation**")
                    st.metric("WACC",             f"{ap_params['wacc']}%")
                    st.metric("Terminal Growth",  f"{ap_params['terminal_growth']}%")
                    st.metric("Base PE",          f"{ap_params['pe_base']}x")

                st.markdown("---")
                dcf_r = st.session_state.dcf_result
                pepb  = st.session_state.pepb_result
                if dcf_r and pepb:
                    v1, v2, v3, v4 = st.columns(4)
                    v1.metric("DCF Intrinsic Value", f"${dcf_r['price_per_share']}/share")
                    v2.metric("PE Base Target",      f"${pepb['pe_price_base']}/share")
                    v3.metric("PB Base Target",      f"${pepb['pb_price_base']}/share")
                    dcf_lo = st.session_state.dcf_params.get("price_low", 0)
                    dcf_hi = st.session_state.dcf_params.get("price_high", 0)
                    consensus = round((dcf_r["price_per_share"] + pepb["pe_price_base"] + pepb["pb_price_base"]) / 3, 2)
                    v4.metric("🎯 Consensus", f"${consensus}/share")

                st.success("📄 All tabs are now populated. Head to **Export Report** to download the full analysis.")

    # ── 入口 B：手动上传 ──────────────────────────────────────
    elif st.session_state.entry_mode == "manual":
        if st.button("← Back", key="back_manual"):
            st.session_state.entry_mode = None
            st.rerun()

        st.header("📝 Upload Your Own Data")
        st.write("Upload an Excel file with your financial data. The file should have these columns:")
        st.code("年份 | 净利润 | 营业收入 | 总负债 | 总资产", language=None)
        st.info("💡 Once uploaded, head to the **Forecast** tab to start building your model. You can skip the Data tab entirely.")

        uploaded_file = st.file_uploader("Upload Excel file (.xlsx)", type=["xlsx"])
        manual_company = st.text_input("Company Name", placeholder="e.g. My Portfolio Company")

        if uploaded_file and manual_company:
            df_m = pd.read_excel(uploaded_file)
            df_m["净利率%"] = round(df_m["净利润"] / df_m["营业收入"] * 100, 2)
            df_m["资产负债率%"] = round(df_m["总负债"] / df_m["总资产"] * 100, 2)
            df_m["收入增速%"] = round(df_m["营业收入"].pct_change() * 100, 2)
            st.session_state.fetched_df = df_m
            st.session_state.fetched_company = manual_company
            st.session_state.is_us_stock = False
            st.success(f"✅ Data loaded for {manual_company}")
            st.dataframe(df_m[["年份", "净利率%", "资产负债率%", "收入增速%"]])
            if st.button("Generate AI Summary", key="manual_ai"):
                with st.spinner("Analyzing..."):
                    data_text = "".join(f"{int(r['年份'])}: Net margin {r['净利率%']}%, Debt ratio {r['资产负债率%']}%\n" for _, r in df_m.iterrows())
                    prompt = f"As a securities analyst, analyze {manual_company} in 200 words:\n{data_text}"
                    msg = client.messages.create(model="claude-opus-4-6", max_tokens=1024,
                        messages=[{"role": "user", "content": prompt}])
                    st.write(msg.content[0].text)

# ============================================================
# TAB 2: FORECAST
# ============================================================
with tab2:
    st.header("📈 Income Statement Forecast")
    st.write("Build a 3-year projection of revenue and profit based on your assumptions about the business.")

    # ---- 数据来源 ----
    if st.session_state.fetched_df is not None and st.session_state.is_us_stock:
        df = st.session_state.fetched_df
        cname_f = st.session_state.fetched_company
        latest_row = df.dropna(subset=["营业收入(亿美元)"]).iloc[-1]
        auto_base_year = int(latest_row["年份"])
        auto_base_rev = float(latest_row["营业收入(亿美元)"])
        st.success(f"✅ Using data from Data tab: **{cname_f}**, base year **{auto_base_year}**, revenue **{auto_base_rev} bn**")
        use_auto = True
    else:
        use_auto = False

    with st.expander("📥 Enter base year data manually", expanded=not use_auto):
        m_cname = st.text_input("Company name", value="My Company", key="f_cname")
        m_base_year = st.number_input("Base year", value=2024, min_value=2000, max_value=2030, step=1, key="f_year")
        m_base_rev = st.number_input("Base year revenue (USD bn)", value=100.0, min_value=0.1, step=1.0, key="f_rev",
            help="Total revenue the company earned in the most recent full year, in billions of USD.")

    if use_auto:
        f_cname = cname_f
        f_base_year = auto_base_year
        f_base_rev = auto_base_rev
        f_hist_df = df
    else:
        f_cname = m_cname
        f_base_year = int(m_base_year)
        f_base_rev = m_base_rev
        f_hist_df = None

    st.info(f"Base year: **{f_base_year}** · Revenue: **{f_base_rev} bn USD**")

    # ── 收入驱动因子 ──────────────────────────────────────────
    st.subheader("Step 1 — Revenue Drivers")

    # 解释弹窗
    with st.expander("❓ What are Revenue Drivers — and what should I put here?"):
        st.markdown("""
**Revenue drivers are the key business variables that determine how much money the company makes.**

The idea is simple: instead of just guessing "revenue will grow 10%", you break revenue down into its components and forecast each one separately. This makes your assumptions more transparent and easier to challenge.

**Examples by industry:**

| Industry | Driver 1 | Driver 2 |
|----------|----------|----------|
| Car manufacturer | Units sold (mn cars) | Average selling price (USD) |
| Streaming service | Subscribers (mn) | Monthly fee (USD) |
| Retailer | Number of stores | Revenue per store (mn USD) |
| SaaS company | Paying customers | Annual revenue per customer (USD) |

**How it works in this tool:**
- You set what each driver was in the base year
- Then you forecast where each driver will be in years 1, 2, and 3
- Revenue = base revenue × (driver changes combined)

**Not sure what to use?** Just pick 1 driver called "Revenue Growth Rate (%)" and set the base to 100. Then forecast 105, 110, 115 for 5% annual growth.
""")

    num_drivers = st.radio("How many drivers do you want to use?", [1, 2, 3], index=1, horizontal=True,
        help="More drivers = more detailed model. Start with 2 if you're new to this.")

    driver_names, driver_bases = [], []
    default_names = ["Units Sold (mn)", "ASP (USD)", "Market Share%"]
    default_help = [
        "How many units/customers/subscribers does the company have? Enter the number in the base year.",
        "What is the average price per unit/customer/subscription? Enter in USD.",
        "What percentage of the total market does this company serve?"
    ]
    for i in range(num_drivers):
        st.markdown(f"**Driver {i+1}**")
        c1, c2 = st.columns([2, 1])
        with c1:
            name = st.text_input(f"Driver {i+1} name — what does it measure?",
                value=default_names[i], key=f"driver_name_{i}",
                help=default_help[i])
        with c2:
            base_val = st.number_input(f"Base value in {f_base_year}",
                value=100.0, key=f"driver_base_{i}",
                help=f"What was the value of this driver in {f_base_year}? Use the same unit as your driver name.")
        driver_names.append(name)
        driver_bases.append(base_val)

    # ── 3年预测 ───────────────────────────────────────────────
    st.subheader("Step 2 — 3-Year Forecasts")
    st.write("Now enter where you think each driver will be in each of the next 3 years.")
    forecast_years = [f_base_year + 1, f_base_year + 2, f_base_year + 3]
    year_cols = st.columns(3)
    year_driver_values = []
    for yi, year in enumerate(forecast_years):
        with year_cols[yi]:
            st.markdown(f"**{year}**")
            year_vals = []
            for di in range(num_drivers):
                val = st.number_input(driver_names[di], value=float(driver_bases[di]), key=f"forecast_{yi}_{di}")
                year_vals.append(val)
            year_driver_values.append(year_vals)

    # ── P&L 假设 ──────────────────────────────────────────────
    st.subheader("Step 3 — Profitability Assumptions")
    st.write("These ratios determine how much of the revenue turns into profit.")

    col_a, col_b, col_c = st.columns(3)
    with col_a:
        gross_margin = st.number_input("Gross Margin %", value=40.0, min_value=0.0, max_value=100.0, step=0.5,
            help="Revenue minus the direct cost of making the product/service, as a % of revenue.\n\nExample: If Apple sells a phone for $1,000 and it costs $600 to make, gross margin = 40%.\n\nTypical ranges: Software 60–80% | Manufacturing 20–40% | Retail 25–45%")
    with col_b:
        expense_ratio = st.number_input("Opex Ratio %", value=20.0, min_value=0.0, max_value=100.0, step=0.5,
            help="Operating expenses (sales, marketing, admin, R&D) as a % of revenue.\n\nThese are costs NOT directly tied to making the product — think office rent, salaries, advertising.\n\nLower is better. Best-in-class companies run at 10–20%.")
    with col_c:
        tax_rate = st.number_input("Tax Rate %", value=25.0, min_value=0.0, max_value=50.0, step=0.5,
            help="The effective corporate income tax rate.\n\nUS corporate tax: ~21%. Including state taxes: ~25–28%.\nChina: 25% standard rate. Tech companies may qualify for 15%.\n\nIf unsure, use 25%.")

    with st.expander("📊 How does this model calculate profit?"):
        st.markdown(f"""
The income statement is built step by step:

```
Revenue                    (your driver forecast)
− Cost of Goods Sold       (Revenue × (1 − Gross Margin%))
= Gross Profit             (Revenue × Gross Margin%)
− Operating Expenses       (Revenue × Opex Ratio%)
= Operating Profit (EBIT)
− Income Tax               (EBIT × Tax Rate%)
= Net Income               ← the bottom line
```

**Net Margin** = Net Income / Revenue — this is the % that ends up as profit.
""")

    if st.button("📈 Generate Forecast"):
        forecast_df, combined_df = build_income_statement(
            base_revenue=f_base_rev, driver_names=driver_names,
            driver_base=driver_bases, driver_forecasts=year_driver_values,
            gross_margin=gross_margin, expense_ratio=expense_ratio,
            tax_rate=tax_rate, base_year=f_base_year,
            hist_df=f_hist_df,
            rev_col="营业收入(亿美元)" if f_hist_df is not None else None,
            net_col="净利润(亿美元)" if f_hist_df is not None else None,
        )
        st.session_state.forecast_df = forecast_df
        st.session_state.combined_df = combined_df
        st.session_state.forecast_params = {
            "gross_margin": gross_margin, "expense_ratio": expense_ratio,
            "tax_rate": tax_rate, "cname": f_cname,
        }

    if st.session_state.forecast_df is not None:
        forecast_df = st.session_state.forecast_df
        combined_df = st.session_state.combined_df
        params = st.session_state.forecast_params

        st.subheader("Projected Income Statement (USD bn)")
        st.dataframe(forecast_df[["年份","营业收入","毛利润","营业利润(EBIT)","净利润","毛利率%","净利率%"]].set_index("年份"))

        hist_data = combined_df[combined_df["类型"] == "历史"]
        forecast_data = combined_df[combined_df["类型"] == "预测"]
        fig = go.Figure()
        fig.add_trace(go.Bar(x=hist_data["年份"], y=hist_data["营业收入"], name="Historical Revenue", marker_color="#4472C4", opacity=0.8))
        fig.add_trace(go.Bar(x=forecast_data["年份"], y=forecast_data["营业收入"], name="Forecast Revenue", marker_color="#4472C4", opacity=0.4, marker_pattern_shape="/"))
        fig.add_trace(go.Scatter(x=hist_data["年份"], y=hist_data["净利润"], name="Historical Net Income", mode="lines+markers", line=dict(color="#ED7D31", width=2)))
        fig.add_trace(go.Scatter(x=forecast_data["年份"], y=forecast_data["净利润"], name="Forecast Net Income", mode="lines+markers", line=dict(color="#ED7D31", width=2, dash="dash")))
        fig.update_layout(title=f"{params['cname']} — Revenue & Net Income",
            barmode="group", hovermode="x unified", plot_bgcolor="white", yaxis_title="USD bn",
            legend=dict(orientation="h", yanchor="bottom", y=1.02))
        st.plotly_chart(fig, use_container_width=True)

        if st.button("🤖 AI Commentary"):
            with st.spinner("Analyzing..."):
                ft = "".join(f"{r['年份']}: Rev {r['营业收入']} bn, NI {r['净利润']} bn, Margin {r['净利率%']}%\n" for _, r in forecast_df.iterrows())
                prompt = f"""Senior equity analyst. Forecast for {params['cname']}:
{ft}
Assumptions: GM {params['gross_margin']}%, Opex {params['expense_ratio']}%, Tax {params['tax_rate']}%
Commentary (~300 words): 1) Bull/base/bear? 2) Margin assumptions reasonable? 3) Key risks?"""
                msg = client.messages.create(model="claude-opus-4-6", max_tokens=1024,
                    messages=[{"role": "user", "content": prompt}])
                st.subheader("🤖 AI Commentary")
                st.write(msg.content[0].text)

# ============================================================
# TAB 3: 3-STATEMENT
# ============================================================
with tab3:
    st.header("3-Statement Model")

    # ---- 数据来源：自动填入 or 手动输入 ----
    if st.session_state.forecast_df is not None:
        auto_forecast = st.session_state.forecast_df
        auto_params = st.session_state.forecast_params
        st.success(f"✅ Using forecast from Forecast tab: **{auto_params['cname']}**")
        use_auto_3s = True
    else:
        use_auto_3s = False

    with st.expander("📥 Manual Input — enter projected P&L directly", expanded=not use_auto_3s):
        st.write("Enter your 3-year net income and revenue forecasts:")
        m3_cname = st.text_input("Company name", value="My Company", key="3s_cname")
        cols_y = st.columns(3)
        manual_years = []
        manual_revenues = []
        manual_net_incomes = []
        manual_ebits = []
        manual_margins = []
        for i, col in enumerate(cols_y):
            with col:
                yr = st.number_input(f"Year {i+1}", value=2025+i, min_value=2020, max_value=2035, step=1, key=f"3s_year_{i}")
                rev = st.number_input("Revenue (bn)", value=100.0*(1+0.05*i), step=1.0, key=f"3s_rev_{i}")
                ni = st.number_input("Net Income (bn)", value=15.0*(1+0.05*i), step=0.5, key=f"3s_ni_{i}")
                manual_years.append(str(int(yr)))
                manual_revenues.append(rev)
                manual_net_incomes.append(ni)
                ebit = ni / 0.75  # 反推EBIT（假设税率25%）
                manual_ebits.append(round(ebit, 2))
                manual_margins.append(round(ni/rev*100, 1) if rev > 0 else 0)

    # 构建 forecast_df 供三表使用
    if use_auto_3s:
        s3_forecast_df = auto_forecast
        s3_cname = auto_params["cname"]
    else:
        s3_forecast_df = pd.DataFrame({
            "年份": manual_years,
            "营业收入": manual_revenues,
            "净利润": manual_net_incomes,
            "营业利润(EBIT)": manual_ebits,
            "净利率%": manual_margins,
            "毛利率%": [40.0]*3,
            "毛利润": [r*0.4 for r in manual_revenues],
            "类型": ["预测"]*3,
        })
        s3_cname = m3_cname

    st.subheader("Balance Sheet & Cash Flow Assumptions")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("**Depreciation & Capex**")
        depreciation_rate = st.number_input("D&A rate % (of PP&E)", value=10.0, min_value=0.0, max_value=50.0, step=0.5)
        capex_rate = st.number_input("Capex % (of revenue)", value=5.0, min_value=0.0, max_value=50.0, step=0.5)
    with col2:
        st.markdown("**Working Capital (Days)**")
        ar_days = st.number_input("DSO — Receivables days", value=60, min_value=0, max_value=365, step=5)
        inv_days = st.number_input("DIO — Inventory days", value=45, min_value=0, max_value=365, step=5)
        ap_days = st.number_input("DPO — Payables days", value=45, min_value=0, max_value=365, step=5)
    with col3:
        st.markdown("**Base Year Balance Sheet**")
        dividend_payout = st.number_input("Dividend payout %", value=30.0, min_value=0.0, max_value=100.0, step=5.0)
        base_cash = st.number_input("Cash (bn)", value=50.0, step=1.0)
        base_fixed_assets = st.number_input("PP&E net (bn)", value=100.0, step=1.0)
        base_equity = st.number_input("Equity (bn)", value=80.0, step=1.0)
        base_debt = st.number_input("Debt (bn)", value=50.0, step=1.0)

    if st.button("📑 Build 3-Statement Model"):
        income_rows, cashflow_rows, balance_rows = [], [], []
        prev_cash, prev_fa, prev_eq, prev_debt = base_cash, base_fixed_assets, base_equity, base_debt

        for _, row in s3_forecast_df.iterrows():
            year = row["年份"]
            revenue = row["营业收入"]
            net_income = row["净利润"]
            ebit = row["营业利润(EBIT)"]

            capex = revenue * capex_rate / 100
            depreciation = prev_fa * depreciation_rate / 100
            dividends = net_income * dividend_payout / 100
            ar = revenue * ar_days / 365
            inventory = revenue * inv_days / 365
            ap = revenue * ap_days / 365
            nwc = ar + inventory - ap
            cfo = net_income + depreciation - (nwc * 0.1)
            cfi = -capex
            cff = -dividends
            fcf = cfo + cfi
            net_cash_change = cfo + cfi + cff

            end_cash = prev_cash + net_cash_change
            end_fa = prev_fa + capex - depreciation
            total_assets = end_cash + end_fa + ar + inventory
            end_equity = prev_eq + (net_income - dividends)
            total_liabilities = ap + prev_debt
            check = total_assets - (total_liabilities + end_equity)

            income_rows.append({"Year": year, "Revenue": round(revenue,2), "D&A": round(depreciation,2), "EBIT": round(ebit,2), "Net Income": round(net_income,2), "Net Margin%": round(row["净利率%"],1)})
            cashflow_rows.append({"Year": year, "CFO": round(cfo,2), "Capex": round(-capex,2), "CFI": round(cfi,2), "CFF": round(cff,2), "FCF": round(fcf,2), "Net Change": round(net_cash_change,2)})
            balance_rows.append({"Year": year, "Cash": round(end_cash,2), "Receivables": round(ar,2), "Inventory": round(inventory,2), "PP&E": round(end_fa,2), "Total Assets": round(total_assets,2), "Payables": round(ap,2), "Debt": round(prev_debt,2), "Total Liabilities": round(total_liabilities,2), "Equity": round(end_equity,2), "Check": round(check,2)})

            prev_cash, prev_fa, prev_eq = end_cash, end_fa, end_equity

        st.session_state.income_rows = income_rows
        st.session_state.cashflow_rows = cashflow_rows
        st.session_state.balance_rows = balance_rows
        st.session_state.three_table_params = {"depreciation_rate": depreciation_rate, "capex_rate": capex_rate, "ar_days": ar_days, "inv_days": inv_days, "ap_days": ap_days, "dividend_payout": dividend_payout, "cname": s3_cname}

    if st.session_state.cashflow_rows is not None:
        income_rows = st.session_state.income_rows
        cashflow_rows = st.session_state.cashflow_rows
        balance_rows = st.session_state.balance_rows

        st.subheader("Income Statement")
        st.dataframe(pd.DataFrame(income_rows).set_index("Year"))
        st.subheader("Cash Flow Statement")
        st.dataframe(pd.DataFrame(cashflow_rows).set_index("Year"))
        st.subheader("Balance Sheet")
        bd = pd.DataFrame(balance_rows).set_index("Year")
        st.dataframe(bd)

        max_diff = bd["Check"].abs().max()
        if max_diff < 1:
            st.success(f"✅ Balance sheet checks out. Max discrepancy: {round(max_diff,3)} bn")
        else:
            st.warning(f"⚠️ Discrepancy: {round(max_diff,2)} bn")

        fig_cf = go.Figure()
        fig_cf.add_trace(go.Bar(x=[r["Year"] for r in cashflow_rows], y=[r["FCF"] for r in cashflow_rows], name="FCF", marker_color="#70AD47"))
        fig_cf.add_trace(go.Scatter(x=[r["Year"] for r in income_rows], y=[r["Net Income"] for r in income_rows], name="Net Income", mode="lines+markers", line=dict(color="#ED7D31", width=2)))
        fig_cf.update_layout(title="FCF vs Net Income", plot_bgcolor="white", hovermode="x unified", yaxis_title="bn")
        st.plotly_chart(fig_cf, use_container_width=True)

        if st.button("🤖 AI 3-Statement Commentary"):
            with st.spinner("Analyzing..."):
                tp = st.session_state.three_table_params
                prompt = f"""Senior financial analyst. 3-statement model for {tp['cname']}:
Net Income: {[r['Net Income'] for r in income_rows]} bn
FCF: {[r['FCF'] for r in cashflow_rows]} bn
Debt ratio: {[round(r['Total Liabilities']/r['Total Assets']*100,1) for r in balance_rows]}%
Assumptions: D&A {tp['depreciation_rate']}%, Capex {tp['capex_rate']}%, DSO {tp['ar_days']}d, DPO {tp['ap_days']}d
Analysis (~300 words): 1) Earnings quality? 2) Capex intensity? 3) Working capital norms? 4) Financial health?"""
                msg = client.messages.create(model="claude-opus-4-6", max_tokens=1500,
                    messages=[{"role": "user", "content": prompt}])
                st.subheader("🤖 AI Commentary")
                st.write(msg.content[0].text)

# ============================================================
# TAB 4: VALUATION
# ============================================================
with tab4:
    st.header("DCF Valuation")

    # ---- 数据来源：自动填入 or 手动输入 ----
    if st.session_state.cashflow_rows is not None:
        auto_fcf = [r["FCF"] for r in st.session_state.cashflow_rows]
        auto_years = [r["Year"] for r in st.session_state.cashflow_rows]
        auto_vcname = st.session_state.three_table_params.get("cname", "Company")
        st.success(f"✅ Using FCF from 3-Statement tab: **{auto_vcname}** · FCF: {auto_fcf} bn")
        use_auto_v = True
    else:
        use_auto_v = False

    with st.expander("📥 Manual Input — enter FCF directly", expanded=not use_auto_v):
        v_cname = st.text_input("Company name", value="My Company", key="v_cname")
        vcols = st.columns(3)
        manual_fcf = []
        manual_vyears = []
        for i, col in enumerate(vcols):
            with col:
                yr = st.number_input(f"Year {i+1}", value=2025+i, min_value=2020, max_value=2035, step=1, key=f"v_year_{i}")
                fcf_val = st.number_input(f"FCF (bn)", value=50.0 + i*5, step=1.0, key=f"v_fcf_{i}")
                manual_fcf.append(fcf_val)
                manual_vyears.append(str(int(yr)))

    # 决定用哪个数据源
    if use_auto_v:
        v_fcf_list = auto_fcf
        v_years = auto_years
        v_company = auto_vcname
    else:
        v_fcf_list = manual_fcf
        v_years = manual_vyears
        v_company = v_cname

    st.info(f"FCF inputs: **{v_years[0]}: {v_fcf_list[0]} bn** · **{v_years[1]}: {v_fcf_list[1]} bn** · **{v_years[2]}: {v_fcf_list[2]} bn**")

    st.subheader("DCF Assumptions")
    col1, col2, col3 = st.columns(3)
    with col1:
        wacc = st.number_input("WACC %", value=9.0, min_value=1.0, max_value=30.0, step=0.5,
            help="Weighted avg cost of capital. Tech: typically 8–12%")
    with col2:
        terminal_growth = st.number_input("Terminal growth rate %", value=3.0, min_value=0.0, max_value=10.0, step=0.5,
            help="Perpetual growth after Year 3. Typically ~GDP (2–4%)")
    with col3:
        shares_outstanding = st.number_input("Shares outstanding (bn)", value=152.0, min_value=0.1, step=1.0)
        net_debt = st.number_input("Net debt (bn)", value=310.0, step=10.0,
            help="Debt minus cash. Negative = net cash")

    st.subheader("Sensitivity Analysis Range")
    col_a, col_b = st.columns(2)
    with col_a:
        wacc_low = st.number_input("WACC low %", value=7.0, min_value=1.0, max_value=29.0, step=0.5)
        wacc_high = st.number_input("WACC high %", value=12.0, min_value=2.0, max_value=30.0, step=0.5)
    with col_b:
        tg_low = st.number_input("Terminal growth low %", value=1.0, min_value=0.0, max_value=9.0, step=0.5)
        tg_high = st.number_input("Terminal growth high %", value=5.0, min_value=0.5, max_value=10.0, step=0.5)

    if st.button("💰 Run DCF Valuation"):
        import numpy as np

        def dcf_valuation(fcf_list, wacc_pct, tg_pct, shares, nd):
            w, g = wacc_pct/100, tg_pct/100
            pv_fcf = sum(f/(1+w)**(i+1) for i,f in enumerate(fcf_list))
            tv = fcf_list[-1]*(1+g)/(w-g)
            pv_tv = tv/(1+w)**len(fcf_list)
            ev = pv_fcf + pv_tv
            eq = ev - nd
            pps = eq/shares if shares > 0 else 0
            return {"pv_fcf": round(pv_fcf,1), "pv_terminal": round(pv_tv,1),
                "ev": round(ev,1), "equity_value": round(eq,1),
                "price_per_share": round(pps,2),
                "terminal_pct": round(pv_tv/ev*100,1) if ev > 0 else 0}

        base = dcf_valuation(v_fcf_list, wacc, terminal_growth, shares_outstanding, net_debt)

        c1,c2,c3,c4 = st.columns(4)
        c1.metric("PV of FCFs", f"{base['pv_fcf']} bn")
        c2.metric("PV of Terminal Value", f"{base['pv_terminal']} bn", f"{base['terminal_pct']}% of EV")
        c3.metric("Enterprise Value", f"{base['ev']} bn")
        c4.metric("Intrinsic Value / Share", f"${base['price_per_share']}")

        if base["terminal_pct"] > 80:
            st.warning(f"⚠️ Terminal value is {base['terminal_pct']}% of EV — highly sensitive to long-term assumptions")
        else:
            st.success(f"✅ Terminal value is {base['terminal_pct']}% of EV")

        wacc_steps = np.arange(wacc_low, wacc_high+0.1, 1.0)
        tg_steps = np.arange(tg_low, tg_high+0.1, 1.0)
        matrix_data, all_prices = {}, []
        for w in wacc_steps:
            row_data = {}
            for g in tg_steps:
                if w <= g:
                    row_data[f"g={g:.0f}%"] = "N/A"
                else:
                    r = dcf_valuation(v_fcf_list, w, g, shares_outstanding, net_debt)
                    row_data[f"g={g:.0f}%"] = f"${r['price_per_share']}"
                    all_prices.append(r["price_per_share"])
            matrix_data[f"WACC={w:.0f}%"] = row_data

        st.subheader("Sensitivity Matrix — Intrinsic Value per Share ($)")
        st.dataframe(pd.DataFrame(matrix_data).T, use_container_width=True)

        if all_prices:
            fig_val = go.Figure()
            fig_val.add_trace(go.Box(y=all_prices, name="Valuation range",
                marker_color="#4472C4", boxpoints="all", jitter=0.3))
            fig_val.add_hline(y=base["price_per_share"], line_dash="dash", line_color="red",
                annotation_text=f"Base ${base['price_per_share']}")
            fig_val.update_layout(title="Intrinsic Value Distribution",
                plot_bgcolor="white", yaxis_title="USD per share", showlegend=False)
            st.plotly_chart(fig_val, use_container_width=True)
            low, high, mid = round(min(all_prices),2), round(max(all_prices),2), round(sum(all_prices)/len(all_prices),2)
            st.info(f"📌 Range: **${low} — ${high}** · Median: **${mid}**")

        st.session_state.dcf_result = base
        st.session_state.dcf_params = {"wacc": wacc, "terminal_growth": terminal_growth,
            "shares": shares_outstanding, "net_debt": net_debt,
            "price_low": min(all_prices) if all_prices else 0,
            "price_high": max(all_prices) if all_prices else 0,
            "cname": v_company}

    if st.session_state.dcf_result is not None:
        if st.button("🤖 AI Valuation Report"):
            with st.spinner("Generating report..."):
                base = st.session_state.dcf_result
                dp = st.session_state.dcf_params
                prompt = f"""Senior equity research analyst. Valuation report for {dp['cname']}.
DCF: base case ${base['price_per_share']}/share, EV {base['ev']} bn, terminal {base['terminal_pct']}% of EV
Range: ${dp['price_low']:.1f}–${dp['price_high']:.1f}
WACC {dp['wacc']}%, terminal growth {dp['terminal_growth']}%
FCF forecast: {v_fcf_list} bn
Report (~400 words): 1) Under/overvalued? 2) Assumption sensitivity 3) Risks 4) Buy/Hold/Sell + target price"""
                msg = client.messages.create(model="claude-opus-4-6", max_tokens=2000,
                    messages=[{"role": "user", "content": prompt}])
                st.subheader("🤖 AI Valuation Report")
                st.write(msg.content[0].text)

    # ============================================================
    # PE/PB 相对估值
    # 背景知识：
    # PE（市盈率）= 股价 / 每股收益(EPS)，反映市场愿意为每1元利润付多少钱
    # PB（市净率）= 股价 / 每股净资产(BVPS)，反映市场愿意为每1元净资产付多少钱
    # 相对估值的逻辑：给定一个"合理倍数"，反推目标价
    # ============================================================
    st.divider()
    st.subheader("📐 Relative Valuation — PE & PB")
    st.write("Estimate target price based on earnings and book value multiples.")

    # ---- 数据来源 ----
    # 优先从三表/预测模块读取净利润和股东权益
    auto_ni, auto_eq, auto_shares_rel = None, None, None
    if st.session_state.get("income_rows"):
        # 取最后一年预测净利润
        auto_ni = st.session_state.income_rows[-1]["Net Income"]
    if st.session_state.get("balance_rows"):
        auto_eq = st.session_state.balance_rows[-1]["Equity"]
    if st.session_state.get("dcf_params"):
        auto_shares_rel = st.session_state.dcf_params.get("shares")

    with st.expander("📥 Manual Input — enter financials directly",
                     expanded=(auto_ni is None)):
        rel_cname  = st.text_input("Company name", value=st.session_state.dcf_params.get("cname", "Company") if st.session_state.dcf_params else "Company", key="rel_cname")
        rel_ni     = st.number_input("Net Income — forward year (bn)", value=float(auto_ni) if auto_ni else 50.0, step=1.0, key="rel_ni")
        rel_eq     = st.number_input("Shareholders' Equity (bn)",      value=float(auto_eq) if auto_eq else 200.0, step=1.0, key="rel_eq")
        rel_shares = st.number_input("Shares outstanding (bn)",        value=float(auto_shares_rel) if auto_shares_rel else 152.0, step=1.0, key="rel_shares")

    # 如果上游有数据就用，否则用手动输入
    f_ni     = auto_ni     if auto_ni     else st.session_state.get("rel_ni",     50.0)
    f_eq     = auto_eq     if auto_eq     else st.session_state.get("rel_eq",     200.0)
    f_shares = auto_shares_rel if auto_shares_rel else st.session_state.get("rel_shares", 152.0)
    f_ni     = st.session_state.get("rel_ni",     f_ni)
    f_eq     = st.session_state.get("rel_eq",     f_eq)
    f_shares = st.session_state.get("rel_shares", f_shares)

    # ---- PE 假设 ----
    st.markdown("**PE Valuation**")
    col_pe1, col_pe2, col_pe3 = st.columns(3)
    with col_pe1:
        pe_bear = st.number_input("Bear case PE", value=20.0, min_value=1.0, step=1.0,
            help="Conservative multiple — e.g. trough historical PE")
    with col_pe2:
        pe_base = st.number_input("Base case PE", value=28.0, min_value=1.0, step=1.0,
            help="Mid-cycle or industry average PE")
    with col_pe3:
        pe_bull = st.number_input("Bull case PE", value=35.0, min_value=1.0, step=1.0,
            help="Premium multiple — e.g. peak historical or high-growth peers")

    # ---- PB 假设 ----
    st.markdown("**PB Valuation**")
    col_pb1, col_pb2, col_pb3 = st.columns(3)
    with col_pb1:
        pb_bear = st.number_input("Bear case PB", value=5.0, min_value=0.1, step=0.5,
            help="Low end of historical PB range")
    with col_pb2:
        pb_base = st.number_input("Base case PB", value=8.0, min_value=0.1, step=0.5,
            help="Average historical PB")
    with col_pb3:
        pb_bull = st.number_input("Bull case PB", value=12.0, min_value=0.1, step=0.5,
            help="Peak PB or premium peers")

    if st.button("📐 Run PE/PB Valuation"):
        # EPS = 净利润 / 总股数
        eps = f_ni / f_shares if f_shares > 0 else 0
        # BVPS = 所有者权益 / 总股数
        bvps = f_eq / f_shares if f_shares > 0 else 0

        # PE目标价
        pe_price_bear = round(eps * pe_bear, 2)
        pe_price_base = round(eps * pe_base, 2)
        pe_price_bull = round(eps * pe_bull, 2)

        # PB目标价
        pb_price_bear = round(bvps * pb_bear, 2)
        pb_price_base = round(bvps * pb_base, 2)
        pb_price_bull = round(bvps * pb_bull, 2)

        st.markdown(f"**EPS (forward):** ${round(eps,2)}  ·  **BVPS:** ${round(bvps,2)}")

        # ---- 展示结果表 ----
        st.subheader("PE & PB Target Price Summary")
        summary_df = pd.DataFrame({
            "Method": ["PE Valuation", "PB Valuation"],
            "Bear Case ($)": [pe_price_bear, pb_price_bear],
            "Base Case ($)": [pe_price_base, pb_price_base],
            "Bull Case ($)": [pe_price_bull, pb_price_bull],
        }).set_index("Method")
        st.dataframe(summary_df)

        # ---- 可视化：三法对比瀑布图 ----
        dcf_low  = st.session_state.dcf_params.get("price_low",  0) if st.session_state.dcf_params else 0
        dcf_high = st.session_state.dcf_params.get("price_high", 0) if st.session_state.dcf_params else 0
        dcf_base_price = st.session_state.dcf_result.get("price_per_share", 0) if st.session_state.dcf_result else 0

        fig_rel = go.Figure()

        # DCF 区间
        if dcf_low and dcf_high:
            fig_rel.add_trace(go.Bar(
                name="DCF", x=["DCF"],
                y=[dcf_high - dcf_low],
                base=[dcf_low],
                marker_color="#4472C4",
                width=0.4,
            ))
            fig_rel.add_trace(go.Scatter(
                x=["DCF"], y=[dcf_base_price],
                mode="markers", name="DCF Base",
                marker=dict(color="#4472C4", size=12, symbol="diamond"),
                showlegend=True,
            ))

        # PE 区间
        fig_rel.add_trace(go.Bar(
            name="PE", x=["PE"],
            y=[pe_price_bull - pe_price_bear],
            base=[pe_price_bear],
            marker_color="#ED7D31",
            width=0.4,
        ))
        fig_rel.add_trace(go.Scatter(
            x=["PE"], y=[pe_price_base],
            mode="markers", name="PE Base",
            marker=dict(color="#ED7D31", size=12, symbol="diamond"),
            showlegend=True,
        ))

        # PB 区间
        fig_rel.add_trace(go.Bar(
            name="PB", x=["PB"],
            y=[pb_price_bull - pb_price_bear],
            base=[pb_price_bear],
            marker_color="#70AD47",
            width=0.4,
        ))
        fig_rel.add_trace(go.Scatter(
            x=["PB"], y=[pb_price_base],
            mode="markers", name="PB Base",
            marker=dict(color="#70AD47", size=12, symbol="diamond"),
            showlegend=True,
        ))

        fig_rel.update_layout(
            title="Valuation Range Comparison — DCF vs PE vs PB",
            plot_bgcolor="white",
            yaxis_title="Target Price (USD)",
            barmode="overlay",
            showlegend=True,
            hovermode="x unified",
        )
        st.plotly_chart(fig_rel, use_container_width=True)

        # ---- 三法交叉：取重叠区间 ----
        st.subheader("📌 Cross-Validation Summary")
        all_lows  = [dcf_low, pe_price_bear, pb_price_bear] if dcf_low else [pe_price_bear, pb_price_bear]
        all_highs = [dcf_high, pe_price_bull, pb_price_bull] if dcf_high else [pe_price_bull, pb_price_bull]
        all_bases = [dcf_base_price, pe_price_base, pb_price_base] if dcf_base_price else [pe_price_base, pb_price_base]

        consensus_low  = round(max(all_lows), 2)   # 三法下限的最大值（最保守的共识）
        consensus_high = round(min(all_highs), 2)  # 三法上限的最小值（最谨慎的上限）
        consensus_mid  = round(sum(all_bases) / len(all_bases), 2)

        if consensus_low < consensus_high:
            st.success(f"✅ **Consensus range: ${consensus_low} — ${consensus_high}** · Average base case: **${consensus_mid}**")
            st.write("Three methods agree on this overlap zone — higher confidence in this range.")
        else:
            st.warning(f"⚠️ No overlapping range across methods. Average base case: **${consensus_mid}**")
            st.write("Methods diverge significantly — review assumptions for consistency.")

        # 存储PE/PB结果
        st.session_state.pepb_result = {
            "pe_bear": pe_price_bear, "pe_base": pe_price_base, "pe_bull": pe_price_bull,
            "pb_bear": pb_price_bear, "pb_base": pb_price_base, "pb_bull": pb_price_bull,
            "eps": round(eps, 2), "bvps": round(bvps, 2),
            "consensus_low": consensus_low, "consensus_high": consensus_high,
            "consensus_mid": consensus_mid,
            "cname": rel_cname,
        }

    # ---- AI综合三法报告 ----
    if st.session_state.get("pepb_result") is not None:
        if st.button("🤖 AI Cross-Valuation Report"):
            with st.spinner("Generating comprehensive valuation report..."):
                pr = st.session_state.pepb_result
                dcf_r = st.session_state.dcf_result
                dp = st.session_state.dcf_params

                dcf_summary = f"DCF base: ${dcf_r['price_per_share']}, range ${dp['price_low']:.1f}–${dp['price_high']:.1f}" if dcf_r else "DCF: not run"

                prompt = f"""You are a senior equity research analyst. Write a comprehensive valuation report for {pr['cname']}.

Valuation results:
- {dcf_summary}
- PE valuation: bear ${pr['pe_bear']} / base ${pr['pe_base']} / bull ${pr['pe_bull']} (EPS: ${pr['eps']})
- PB valuation: bear ${pr['pb_bear']} / base ${pr['pb_base']} / bull ${pr['pb_bull']} (BVPS: ${pr['bvps']})
- Consensus range: ${pr.get('consensus_low','N/A')} — ${pr.get('consensus_high','N/A')}, average base ${pr.get('consensus_mid','N/A')}

Write a professional valuation conclusion (~400 words):
1. Which method is most appropriate for this company and why?
2. What does the convergence (or divergence) between methods tell us?
3. Final target price recommendation with Bull/Base/Bear scenarios
4. Investment rating: Strong Buy / Buy / Hold / Sell / Strong Sell"""

                msg = client.messages.create(model="claude-opus-4-6", max_tokens=2000,
                    messages=[{"role": "user", "content": prompt}])
                st.session_state.ai_report_text = msg.content[0].text
                st.subheader("🤖 AI Cross-Valuation Report")
                st.write(msg.content[0].text)

# ============================================================
# TAB 5: EXPORT REPORT
# ============================================================
with tab5:
    st.header("📄 Export Research Report")
    st.write("Generate a professional Word document (.docx) summarising all your analysis.")

    # 检查有哪些模块已完成
    has_data     = st.session_state.fetched_df is not None
    has_forecast = st.session_state.forecast_df is not None
    has_3stmt    = st.session_state.cashflow_rows is not None
    has_dcf      = st.session_state.dcf_result is not None
    has_pepb     = st.session_state.pepb_result is not None
    has_autopilot = st.session_state.get("autopilot_done", False)

    # Auto-Pilot 完成提示
    if has_autopilot:
        ap = st.session_state.autopilot_params
        dcf_r = st.session_state.dcf_result
        pepb  = st.session_state.pepb_result
        if dcf_r and pepb:
            consensus = round((dcf_r["price_per_share"] + pepb.get("pe_price_base",0) + pepb.get("pb_price_base",0)) / 3, 2)
            st.success(f"🚀 **Auto-Pilot results ready** — Consensus target price: **${consensus}/share** · Click Generate Report below to export.")
        else:
            st.success("🚀 Auto-Pilot complete — all modules populated.")

    st.subheader("Completed Sections")
    col_s1, col_s2, col_s3, col_s4, col_s5 = st.columns(5)
    col_s1.metric("📊 Data",        "✅ Done" if has_data     else "⬜ Pending")
    col_s2.metric("📈 Forecast",    "✅ Done" if has_forecast else "⬜ Pending")
    col_s3.metric("📑 3-Statement", "✅ Done" if has_3stmt    else "⬜ Pending")
    col_s4.metric("💰 DCF",         "✅ Done" if has_dcf      else "⬜ Pending")
    col_s5.metric("📐 PE/PB",       "✅ Done" if has_pepb     else "⬜ Pending")

    if not (has_data or has_forecast or has_dcf or has_pepb):
        st.info("👈 Complete at least one analysis tab to generate a report.")
    else:
        report_title = st.text_input("Report title", value=f"{st.session_state.fetched_company or 'Company'} — Equity Research Report")
        analyst_name = st.text_input("Analyst name", value="")
        include_ai_summary = st.checkbox("Include AI executive summary", value=True)

        if st.button("📄 Generate Word Report"):
            with st.spinner("Building report..."):
                try:
                    import subprocess, json, tempfile, base64
                    from datetime import date

                    # ---- 收集所有数据 ----
                    cname_r = st.session_state.fetched_company or "Company"
                    today   = date.today().strftime("%B %d, %Y")

                    # AI执行摘要
                    exec_summary = ""
                    if include_ai_summary:
                        parts = []
                        if has_data:
                            latest = st.session_state.fetched_df.iloc[-1]
                            nm  = latest.get("净利率%", "N/A")
                            dr  = latest.get("资产负债率%", "N/A")
                            roe = latest.get("ROE%", "N/A")
                            parts.append(f"Historical: Net margin {nm}%, Debt ratio {dr}%, ROE {roe}%")
                        if has_forecast:
                            fp = st.session_state.forecast_params
                            fd = st.session_state.forecast_df
                            parts.append(f"Forecast: GM {fp['gross_margin']}%, Opex {fp['expense_ratio']}%, Tax {fp['tax_rate']}%. Projected net income: {list(fd['净利润'])}")
                        if has_dcf:
                            d = st.session_state.dcf_result
                            dp2 = st.session_state.dcf_params
                            parts.append(f"DCF: base ${d['price_per_share']}/share, EV ${d['ev']}bn, range ${dp2['price_low']:.1f}–${dp2['price_high']:.1f}")
                        if has_pepb:
                            pr = st.session_state.pepb_result
                            _c_low  = pr.get('consensus_low',  round((pr.get('pe_price_base', pr.get('pe_base',0)) + pr.get('pb_price_base', pr.get('pb_base',0)))/2*0.9, 2))
                            _c_high = pr.get('consensus_high', round((pr.get('pe_price_base', pr.get('pe_base',0)) + pr.get('pb_price_base', pr.get('pb_base',0)))/2*1.1, 2))
                            _c_mid  = pr.get('consensus_mid',  round((pr.get('pe_price_base', pr.get('pe_base',0)) + pr.get('pb_price_base', pr.get('pb_base',0)))/2, 2))
                            parts.append(f"PE/PB consensus: ${_c_low}–${_c_high}, mid ${_c_mid}")

                        ai_prompt = f"""Write a concise executive summary (200 words) for an equity research report on {cname_r}.
Data: {'; '.join(parts)}
Cover: investment thesis, key financial strengths/risks, valuation conclusion, and rating (Buy/Hold/Sell)."""
                        ai_msg = client.messages.create(model="claude-opus-4-6", max_tokens=600,
                            messages=[{"role": "user", "content": ai_prompt}])
                        exec_summary = ai_msg.content[0].text

                    # ---- 用 python-docx + 内嵌模板生成报告 ----
                    try:
                        from docx import Document as _D
                    except ImportError:
                        import subprocess, sys
                        subprocess.run([sys.executable,"-m","pip","install","python-docx",
                                        "--quiet","--disable-pip-version-check"],
                                       capture_output=True)
                        from docx import Document as _D

                    from docx import Document
                    from docx.shared import Pt, RGBColor, Inches
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement
                    import io, base64 as _b64

                    # 内嵌空白 docx 模板（保证 Word 能正常打开）
                    _BLANK_B64 = "UEsDBBQAAAAIAHgka1ytUqWRlQEAAMoGAAATAAAAW0NvbnRlbnRfVHlwZXNdLnhtbLWVTU/bQBCG7/0Vli8+IHtDDxWq4nAocCyRGkSvm/U4Wdgv7UwC+ffMOolV0VCHBi6RnJn3fR7bsj2+fLYmW0NE7V1dnFejIgOnfKPdoi7uZjflRZEhSddI4x3UxQawuJx8Gc82ATDjsMM6XxKF70KgWoKVWPkAjietj1YSH8aFCFI9ygWIr6PRN6G8I3BUUurIJ+MraOXKUHb9zH93IvlDgEWe/dguJlada5sKuoE4mIlg8FVGhmC0ksRzsXbNK7NyZ1VxstvBpQ54xgtvENLkbcAud8tXM+oGsqmM9FNa3hJqheTtb2uEJrDT6AOeV/9uO6Dr21YraLxaWY5UfWnqg0gaevdDDpzrwIIpJ7MhXZQGmjK8j618hPfD9/cppY8kPvnYiF731NNNbcxVgMgPhjVVP7FSu0GPlskzOTf/cepDIn31oIRb2TlETn28RF89KIFAxHv48Q775mEF2hj4DIGu90j8vabldduComNMLJYpW/2VHaQRv5Fh+3v6C6erGUQ+wfzXp93lP8r3IqL7FE1eAFBLAwQUAAAACAB4JGtceSZLQPgAAADeAgAACwAAAF9yZWxzLy5yZWxzrZLNSgMxEIDvPkXIJadutlVEpNleROhNpD7AmMzupm5+SKbavr1RRF1YFsEe5+/jY2bWm6Mb2CumbINXYlnVgqHXwVjfKfG0u1/cCJYJvIEheFTihFlsmov1Iw5AZSb3NmZWID4r3hPFWymz7tFBrkJEXyptSA6ohKmTEfQLdChXdX0t028Gb0ZMtjWKp6255Gx3ivg/tnRIYIBA6pBwEVOZTmQxFzikDklxE/RDSefPjqqQuZwWuvq7UGhbq/Eu6INDT1NeeCT0Bs28EsQ4Z7Q8p9G440fmLSQjzVd6zmZ13oNRf3DPHuwwsZfvWrWP2H0IydFbNu9QSwMEFAAAAAgAeCRrXIiGC1NpAQAA0QIAABEAAABkb2NQcm9wcy9jb3JlLnhtbJ2Sy07DMBBF93xF1E1WifMQCEVJKgHqikpIFIHYufY0NU1sy542zd/jpG1aoCt2Ht87x/NwPt03tbcDY4WShR+Hke+BZIoLWRX+22IW3PueRSo5rZWEwu/A+tPyJmc6Y8rAi1EaDAqwngNJmzFdTNaIOiPEsjU01IbOIZ24Uqah6EJTEU3ZhlZAkii6Iw0g5RQp6YGBHomTI5KzEam3ph4AnBGooQGJlsRhTM5eBNPYqwmDcuFsBHYarlpP4ujeWzEa27YN23Swuvpj8jF/fh1aDYTsR8VgUuacZSiwBjIc7Xb5BQwPATNAUZlSd7hWMuCK7XNycd/PdgNdqwy3hwwOlhmh0e2orECCoQjcW3beb8SlscfU1OLcLXMlgD90ZLgzsBP9tss4J5dhfpzdoQ7Hdz1nhwmdlPf08Wkxm5RJFKdBnARJukjSLL7Nouizf/9H/hnYHCv4N/EEGOpnDl4p03dD/vzC8htQSwMEFAAAAAgAeCRrXPTb2xfrAQAAbAQAABAAAABkb2NQcm9wcy9hcHAueG1snVTLbtswELz7KwRddIppB0FRGJKC1kHRQ90asJKct9TKIkqRBLkx4n59+YgVOYYv9Yk7szv7tMr710FmB7ROaFUVy/miyFBx3Qq1r4rH5tvN5yJzBKoFqRVWxRFdcV/Pyq3VBi0JdJlXUK7KeyKzYszxHgdwc08rz3TaDkDetHumu05wfND8ZUBF7Hax+MTwlVC12N6YUTBPiqsD/a9oq3mozz01R+P16lmWlQ0ORgJh/TMEy3mraSjZiEYXTSAbMWC98MxoBGoLe3T1smTpEaBnbVsXPNMjQOseLHDy0wz4xArkF2Ok4EB+0PVGcKud7ijbABeKtOuzIFOyqVeI8o3tkL9YQcegOTUD/UMojMnSI5VqYW/B9BGfWIHccZC49rOpO5AOS/YOBPo7Qtj8FkQq2kMHWh2Qk7aZE3+xym/z7Dc4DJOt8gNYAYry5PvmnbATlEBpHNm6ESR9ztE+RbHLsKtK4i6sIT2uxicklh37Yh8bK2Mp7lfn50PXWl1OW40VnzUaEXYl4YV+uQHlbycFlGs9GFBHdlriH/doGv0QLvFtMefg+XU9C+p3Bjh+uLMJHpftCWz9yYzLHoG4bN+XlT7NV98kO4ecF1V7bE+Rl8TbST+lT0e9vJsv/C8e8Amb+fMb/9X17B9QSwMEFAAAAAgAeCRrXA5skMb3AQAADAYAABEAAAB3b3JkL2RvY3VtZW50LnhtbKWUzY7aMBDH732KKJecIA7QlEYb9sCK1R4qodI+gHGcxNrYY9mGlD59J5+kqoTocorn6+e/x848Pf+SlXfmxgpQaRDNSeBxxSATqkiDnz92s3XgWUdVRitQPA0u3AbPm09PdZIBO0munIcEZZNas9QvndNJGFpWckntXApmwELu5gxkCHkuGA9rMFm4IBFpV9oA49bidluqztT6PU7CfTRJ2bBcELJGW6iR8a8i0FxhMAcjqUPTFFhh3k96hkxNnTiKSrhLw4pHzDn1T0YlPWM26mhqEhSQnGU1JMOt3E5o/xkqzD0iu5KXvuWtvNDwCgWDsqXQ1759lIbBcoDcPPDksLWOVo9d+ouhNX6uwHvkZ12RrDrlt4kRueNGGsRYcY+Ev/cclEwfX/2x1kybWzzW21cDJ32licdob+p9ZOEg+B9Wf0fTo9nHxBxKqvEHkix5KxQYeqxQEXbca16kv8HpdITs0nwtZ25vvDrBCZd9T31Cdtv463LnD669aZwkJvFyOzgPWNR6l6s4ilueLg6/MYqPI1osVqTJLHH9eY3rsEv4Rpt9HOAbjlZdihFFiaRoTVrzCM6BvIYrnk+iJacZRzVfFq2ZA7iJWZxca/bbMagseq2mjHc5rRuH8qsRWcMWiu+FY6hyGbfRcOhGu+waFF7n+OYPUEsDBBQAAAAIAHgka1xugBsSMgEAAMsEAAAcAAAAd29yZC9fcmVscy9kb2N1bWVudC54bWwucmVsc62UQU+DMBiG7/4KwoWTFKZuixnsoia7KkavpXyFRtqS9kPl31vdZCxD4oHj9zZ9nydt0832U9beOxgrtEqCOIwCDxTThVBlEjxnD5frwLNIVUFrrSAJOrDBNr3YPEJN0e2xlWis50qUTfwKsbklxLIKJLWhbkC5Fa6NpOhGU5KGsjdaAllE0ZKYYYefnnR6uyLxza648r2sa+A/3ZpzweBOs1aCwhEEsdjVYF0jNSVg4u/n0PX4ZBx//QdeCma01RxDpuWB/E1cjRJfBFb3nAPDM/hgacrjZtZjAER3v0OXQzKlsJxT4QPypzOLQTglsppThGuFGc1rOGr00ZTEek4JdHsHAj/jPoynHOI5HVhrUctXR+s9wvCYEoEgJ20Wc9qoVuZg3Es42vTRrwQ5+YPSL1BLAwQUAAAACAB4JGtcB9SvmXMvAAASVQUADwAAAHdvcmQvc3R5bGVzLnhtbO1dXZPiRrJ9v7+io1/85G2QhADHzm4AknYcYXu9nrHvM00z0+zQ0Bdoj+1ffyUhQB9VUlVWSqqSsjvCnhZQKeVXnZNUZf39n3+8bO9+Xx+Om/3u3TfDvw2+uVvvVvunze7zu29+/Rh8O/nm7nha7p6W2/1u/e6bP9fHb/75j//5+9fvjqc/t+vjXfj53fG7l9W7++fT6fW7h4fj6nn9sjz+bf+63oUvftofXpan8M/D54eX5eHL2+u3q/3L6/K0edxsN6c/H6zBwL1PhjmIjLL/9GmzWnv71dvLeneKP/9wWG/DEfe74/Pm9XgZ7avIaF/3h6fXw361Ph7DZ37Znsd7WW5212GGTmGgl83qsD/uP53+Fj5MckfxUOHHh4P4Xy/b+7uX1Xfff97tD8vH7frdfTjQ/T9CzT3tV9760/JtezpGfx5+PiR/Jn/F/wv2u9Px7ut3y+Nqs/kYSg0HeNmEY72f7Y6b+/CV9fJ4mh03y/SLfnItev05eiPzk6vjKXV5vnna3D9EQo9/hS/+vty+u7esy5XFMX9tu9x9vlxb77799UP6ZlKXHsNx390vD99+mEUffEie7SH/xK/5v2LBr8vVJpaz/HRah34RmiUadLsJvfDeGruXP355i1S7fDvtEyGviZD0sA8FpYfuEjrPh7MPh6+uP/2wX31ZP304hS+8u49lhRd//f7nw2Z/CP303f10mlz8sH7ZvN88Pa137+6HlzfunjdP6/99Xu9+Pa6fbtf/E8S+loy42r/tTufbj2/i+OT/sVq/Rp4bvrpbRjb5KfrANnr3MSUn/vjb5nY35ws5qfHF/7uIHCb2Ykl5Xi+jGL8bVgqa4giymONKDWGrD+GoDzFSH8JVH2KsPsREfYgpfIjTfnV2vvTH7WnFJwpeVPmJgtNUfqLgI5WfKLhE5ScKHlD5iYLBKz9RsG/lJwrmLP3Eahn/XfjMSNgHPm5O23VlAhoqprok7d/9vDwsPx+Wr8930dxakFIywoe3x5PYrQ7VbvXD6bDffa4UY1lqYvyX1+flcXOsFqSo+o8R8Ln712HzVClqxJln+IP/vF2u1s/77dP6cPdx/cdJ9vM/7e8+nFFGtV3V1PDD5vPz6e7Dc5w0K4W5HKVXjf/D5niqHpzzKFWDC9nQ5fglf/Af10+bt5eLagTQiGsrirCqRThAEZEBRB5hpDK+wP27wPEjG4vc/1hlfIH7n6iMb1ePL51pvJC3ioXXWDp2F/vt/vDpbSucHsbSEXwVIfYI0kF8HV8oSYylIziTPu9mq1XI3ET8VCGPSkhRSKgSUpQzq4Qs5RQrIUst10oIkk66v6x/3xwv+FbKvMcU1qy8MZujAVFs8Z+3/akamFqKLP773Wm9O67vxKTZirAxM99J2Fht4pMQpDYDSghSmwolBMHnRHEh6pOjhCy1WVJCkNp0KSEIZ94UwF8I86aAFIR5U0AK2rwpIAtt3qydo0gIUiMrEoJwkreAIJzkXTuPkRCknryrheAlbwFZOMlbQBBO8hYQhJO8BcgtQvIWkIKQvAWkoCVvAVloyVtAFk7yFhCEk7wFBOEkbwFBOMlbQBBO8q61GiUuBC95C8jCSd4CgnCSt4AgnOTtNJK8BaQgJG8BKWjJW0AWWvIWkIWTvAUE4SRvAUE4yVtAEE7yFhCEk7wFBKkn72oheMlbQBZO8hYQhJO8BQThJO9RI8lbQApC8haQgpa8BWShJW8BWTjJW0AQTvIWEISTvAUE4SRvAUE4yVtAkHryrhaCl7wFZOEkbwFBOMlbQBBO8nYbSd4CUhCSt4AUtOQtIAsteQvIwkneAoJwkreAIJzkLSAIJ3kLCMJJ3gKC1JN3tRC85C0gCyd5CwjCSd4CgqRzQ7TOdru+E16eOkRa1SC+HlZ1fe/5AX9Zf1of1ruVwEoKRYGXJ5SQqLi2eL7ff7kTW9htcxxEWNTmcbvZx8ts/iyMPS5blvzvxd379XW5XW7Fe0H8w9fMdqFo2HjzW/jG05+v4Xiv6dU+T+fl5smi4fiN3z9dt/VEH45u4i7ZQJVcju81kRr/+3AMQy15z2AQLNypHST3Eg9ZcRNXsdFjrg8Fsc/ny7Gox2Wo93/vWHe03ey+XK6fR1o8L5OP3bR2ecc02S2QtSjjcXx3OJkH5zcn+71Oy8dj8v/L+6I0E95j+Ofr/vju3nEnSe5IvecQ4aPrW6a2O0iUdBmvsI8sdq9kF5lz/YO7i4yj7FWohuUqub3V2/G0f4mdI2/1lNLyJji/dHdTaM4OybaF60qyeNMCxypVFuGpX9abgv3+xPCmT+fLMt50Hom8ScqbUkrLm+D8kqo3BSlD1u9NSQoeMrPTeTtAlUvt1n+cRBJXJKbU2cQz8NXJvqzXrz+F8h8uf/wQmv74kPWTx/Wn/SHUgDOJvePqNvHb9m+nyF1++H17FZR2mIrNwMv/lmwGjl7kbgbOfPK2GTi6fNsM/Hj+7+L8RKsIA17u0nZHwTR2zfijMT4M/T0GhrfLEQSOZulEa6nNxZPLldTm4kny5IfyUCn1JIvrSRamJ1kCnsTIWvU5V7I3usq5hkY4lxNMhnOP51x5V3IZruQiuJLNdSUb05VsQ13J6oYrKTqJw3USB9NJHAEnuREtbX3G1tVnNuf/tuFBI64HjTA9aNQND3L08aCMl1iOHZy/QRDAQ+MAwW9crt+4mH7jdsNvRvr4TUmuad6LxlwvGmN60bgbXuQa4UXOIPrNe9Ep1MXNhz5uoi5EcwwXmnBdaILpQpNuuNBYHxdS4FwDBucaIPjSlOtLU0xfmnbDlyb6+BJiOsJytExJlfOVDLMmmndBTvcgjvsMxdyHf9+nqGNOyT3HHXVKv0u6i99SVcOtdvDT4zYppj9uv99F/v01qXef7/Tpj+X95Y2L9Xb74/L87v0r/63b9afT+dXhYMJ4/XF/Ou1f+J+PC/T8AR6yN/NwfQi+vndvL4/rQ/JFIPeru7hxRlHd54YaipqWTZY/7S9dixg3dHmp3D2lcpcG36Bdq/f5J35/+aIA42u0+KuI8mmBryx9qhm6VOolDWyVGthCMrDVNQM3Vi2XNKddak4byZx278wJhdjnFTl5e5yvYmDreKQyYD0cAOae1/nTIYML4rdGjZqT5UV/RTj47jxJRd+yxmo/K01ElZfxC3OcPRCZ5SJZuwjLvi23ycyrDSbPuNVwHE4EBV1Ed25xJ4GrSm4ltIi7HK4ucpscrm9idI0eWbj55eZoTGdWTSypiOD7sJ5pxTSLsxPVtddq3rzXFzDS1WWw0owFQcshnzj/Y7MtfvGevKhHglD51qvgLMNRAWs4DKzh4OaCjBV5/qKaEbJ+x3cTPZOCxlZmx39EqW/d8/JGzTXXq0oFRWvZDiCoN3H5IypeRMvnB9VTv+xDz/dPf8YtjPPPG71wbm5c9ahpl70Mh7K8cjYbehOvvCQwtDIL19QjO/MEXKWohvZV7RU64ikEaubiOrXbI1WvVGM9QfmSNGxTX5FxsqqxzvpP9glL9IblDPwSQU3eUFxqdnuq6sVmrEcoX1VWY+Bf57rbDDFk1ByGyDWH7HOXaBPLR/h1hwofQVYQfwplzpyA+VLFW9LTpn1e2fC83H2Ojpa6T9bW406j0TMWc2vSNr3GZ7ctN5gOSiFDI89ezCTxs1cnkfqefTiYNPTw87ftds32+7vktWbVcKWC4T++v741xwXr0gMnDM4vNh4NbFVYzaiCExWJKpoODrYq7LpV8VP8PSdbE8lrOuhh1IweONFxfrHe6LCmrj31BFThNqMKTnQkqqg1OoRVMa5bFYtwvM3urVh0jHVxfbVZXfDAdhFY1TKfXp6aEyuXlxuPFjG11FKmSauFEzdXtTQdOWJqieEYul5+XK4Oe2b96iV6pcijrh9AIaoMbTA2AEcKiO463ts7Giesi/eG4fDy1Qb3HePL1yG8d1j2wKl4x4SxCznzDtsZVdypE86uSX48P3XF1wtRP4+3w+ZMquOC8u1KQkSvAA1rAV4Jd8+6Qt5/4ldRan03H5Wi7mnfalOd7MA7H8eSV9r5alX6EfmaLB6pLEYtqY3TiQJLvpMYxD/s5aKobnd7Mqb2VL0tZQK+0oxQVGYPYl5Xl/U8DtJ6HocbnEnkZNdS6vmV2+P5v41sLpS04qjUiiMkK466YMX6t2ZJ2s4ttZ2LZDu3C7ZrepOdpCXHpZYcI1ly3HFL4m90kzTjpNSMEyQzTrpgxnY2m0nac1pqzymSPaddsKeGG77YBGmRnFGft+rl7HooSWIsLBox7ae6c/BW1hHccXN1jCwMhUfgkLEHZAjZA3Jbtnc+5T5vk+SyXIQx2JUFoKRpZUEf69pFNP9g1xeUH01qDT2DRELDKGkjyi43ZM+GxSg7pMWVVR/suvYUOKh7Ctgbe60Joz47teO+uvE+x/Nf1aGtB8Ms2KzUTVQn04xDVniHVPA3qs7MQubtmptA8n2RVfPIELlqNxlMkmUeVXM+iFHlfYyrp0I/Z+WEK7UFQDN3uvV85vjT7Q2qerIhejqGuX8bAjSGZhaD0cDhaOayPjOXudXdiq+vYhdtZYWpghRV9bE3+yAqNeoDzt50mOoQrqxGG1mNLLWAt1z+e3FpMp5XQboBOUsH2e3oEiSknvYlxeYj0zQuEehmcVNKdCU6R6Cok+iV+IgBpkrSjS84Tz+q/F4Fo6WBXGeM+f7wtD6cv4uOO2NUoM1BCm3etpkmfTNAnxXFuexPXzpugD682YWWWL9X+/hvsI8/FNRvcpuSYiDFJwMlp4wwlqKkTkGChpNbiZ9xwulwWdMl+PXm5WJm++rDdZx0dMZM5Zf91/ly9/Rh89dVP8NrfMbvCIfnvwMjwiccZ634Fld847vEoCYGxs1UPx+uH/q0ORxPoXHvma54Id3ZXloAv2SVhpIbO7vAKrmyqtUT0lPAbrOtzT1yKf8qKpfLc9d/y11/yOjj4aKlh7QhOWbdLsmq3bNqHKzhXd1X2EDUQ5CGegzT/vC39eG8crHC/Exj4es1tO/zdcJdbdfLQx7ehH9+2mxjohf9Xq0exBezs2R07Vx7uR4gJG61WD3v94e/eq8eKDT7dpaUc0oh2uVQNfaBJ5pjNUCPMTPRmsDXZpDULVIGJMSm3dwu4A1os7uALEJtZFlCbsZAE8/2At/PQZP8nNln7IaqIEX0xtoCx0Bv7J1wmqO3qWO7tsP7rqhD6E3gSzFIAq8cltCbjnO8gDegzfECsgi9kWUJvRkDTvwghCe32TENTrJX+4reUBWkiN5YO/UZ6I29YV9z9DZ2p5a9YCcgu0vobTqfz0dT3oOCE3jlsITedJzjBbwBbY4XkEXojSxL6M0ccOL6vjdighM7c7W36A1TQYrorXjGNhO9sQ/c1hy9jQJnOp6xE9CtJNcB9DYZuM7M4j0oOIFXDkvoTcc5XsAb0OZ4AVmE3siyhN6MASde4E38CROcOJmrfUVvqApSRG8jMfQ2MhG92cOJM52zE9ANPHcAvTnz2WLh8h4UnMArhyX0puMcL+ANeKujqmUReiPLEnozB5xY/izILuAqzpm9Rm+YClJEb64YenNNRG++7S4GnNrbLS91AL0F46nrcDKtC0/glcMSetNxjhfwBrQ5XkAWoTeyLKE3Y8BJ4PmOl99QmZ8z+4zeUBUkjd44Bz9G+uAe/ygC0ypPuMbvq6M7qpLa2a9vM5DSBj/UYKRx4JcjKUH8k9f043L15fNh/xZmSgYtyaRL4cSVs2l6q7xsCjcDVD3t3x5vru5SmEPCvMfgjGYMLVxJCi+SzZq2GQjCVvRMic9ZVG6YQphWte+B3k1TQD5PrVi6iG1zVs22EugzuqWAFwp4Qrk0h+jhUvWiXbIdku1UUC+v10wa9cIbzTBR72I+cF2nr6hXsl+E3s1mQF5PLWy6iHpzVs22YOgz6qWAFwp4Qr00h+jhUvWiXrIdku1UUC+vR08a9cIb9BDqVe2zoXeTHpDXU+ufLqLenFWzrSv6jHop4IUCnlAvzSF6uFS9qJdsh2Q7FdTL622URr3wxkaEelX7k+jd3Ajk9dQyqYuoN2fVbMuPPqNeCnihgCfUS3OIHi5VL+ol2yHZTgX18npCpVEvvCEUoV7Vvi56N4WCreuhVlMdRL05q2ZbpfQZ9VLACwU8oV6aQ/RwqZrX9ZLtcGyngnp5vbTSqBfeSItQr2o/HL2baYG8nlp0dRH15qyabTHTZ9RLAS8U8IR6aQ7Rw6XqRb1kOyTbSaPefx02Txy0G78EBbmXFc4EcqlBiciYuZ5/qKP+hjoqAXE5QHkI9rvTMRrkuNpsPkYqfXf/svzv/vB+FponGmUdYozZcbNMv+gn16LXn6M3Mj+5Op5Sl+ebp02iSEUUa2ZED3UOaV4bz7a7UjVDq4yMAuq715cgYDFGbVwWSFO1uf/uTzwahZwqx6Wekm3bTKK+uhhEv9dx051w09ea6XBOHmECddPWzyzys276GWqtrqLfavQW9X6rVLyjfmuQUUFFPOFxJWOU+sNSCcPk6OYW83QJb7VaRp2tN6mkR82GKRyy84OuxTEq7lHwNRkM1FJbK9tJFGE82wt8/zpy9miA9FVNy33kG61SPY19rr7SH/mcJj5XRxGQ134+XQSEt5+nImDB5tR+VmBUUBFQeFzJKKV2+VT0MDm6uUVAXcJbrepRZydyKgLS2QsUDtn5QdciGhUBKfiaDAY6YUQr20kUZPzAsz1298zsVU2LgOQbrVI9jX2uviIg+ZwmPldHEZB3Gk+6CAg/jYeKgAWbUzd+gVFBRUDhcSWjlE4PoqKHydHNLQLqEt5qVY86D2ahIqBiEVDHeKBwoCKghvffj8lIs+DTtghItpO0nUxBxvV9b3QdOXtwZPqqpkVA8o1WqZ7GPldfEZB8ThOfq6MIyDucMF0EhB9OSEXAgs3pcCKBUUFFQOFxJaOUDlOkoofJ0c0tAuoS3mpVjzrPqaMioGIRUMd4oHCgIqCG99+PyUiz4NO2CEi2k7SdREHGC7yJP7mOnD1HO31V0yIg+UarVE9jn6uvCEg+p4nP1VEE5J3VnC4Cws9qpiJgcQs4ndVYPSqsJ6DouLKb9ulsaSp6GBzd/J6AmoS3YhO0Go/tpSKgak9ADeOBwoGKgBrefz8mI82CT9siINlO0nYyBRnLnwXZTmy3gdNXNS0Ckm+0SvU09rkaewKSz+nhc3UUAV2BIuDl8GMqAiIUAenoaoFRQUVA4XElo1ToqG0qAna86GFudHOLgLqEt1rVQyg8qQjYThFQx3igcKAioIb334/JSLPg07YISLaTtJ1EQSbwfMcbXEdOF2TczFVNi4DkG61SPY19rr4iIPmcJj6HUQT8cf20eXv58Lx8Cu+weDTw+eW75HWFc4Eve6+p/Hcr+Q6i37y1s0eDn1PAPADX1qVlgErt0lIglXdpIbD1g5JiqOInV+FI851E6RcDBfFPXvWPy9WXz4f9Wwij7utdKEHx2Gg88oobyXUwvhrEPzl8db4vWSDVTNGvoUV45N76urdy7Q2xDAYdqqoKIhzAi0H0ywzg9LVmKHlDSauJZxamhHV4MpyUJMsTqsnJZZECsRREljKezwYL7mGVWBMHRApk6oDIAUweEDEgtiIviPhKV/gKRWY7kVkXBMgdC5w9MLrPzIUc3QBHJwbT+NnvunEYzU6815LFFA9e57EY+PHrxGIK2XARjOdjTqNdC20KgUgBnXQGkAM5+gwgBnaEu7QgYjFdYTEUme1EZn2FzMy5htkTL/vMYsjRDXB0YjH6HpjcUALT7MheLVlM8eRYHouBnx9LLKaQDef2YjHhdAq00aYQiBTIFAKRA5hCIGJALEZeELGYrrAYisx2IrMuEJA7mCl7ZFefWQw5ugGOTixG3xMfm2Ixep05qCWLKR59x2Mx8APwiMUUsuE0mMzmnJqOgzaFQKSADpwEyIGcQAkQA2Ix8oKIxXSFxVBkthOZdYGA3MkS2TNH+sxiyNENcHRiMfoeWdXUijK9Dk3SksUUz+7hsRj4CT7EYorrayeLgeews+EIbQqBSAEtSgbIgSxKBoiB7YuRFkQspisshiKzncisbV9MtjV2tml6n1kMOboBjk4sRt8zN5piMXqd+qAliykePsBjMfAjCIjFFDvOTeeDMScbumhTCEQKqNsfQA6k/R9ADOwYA2lBxGK6wmIoMtuJzLpAQK63Z7bra59ZDDm6AY5OLEbfpuFNJTC92lZrxWIqd/XDN/M7/SUt3NOKLrcPPOwo9fQElo0CywIekYYG16Sg5Ca5CTqXaaidbd5NMo6ReldN+LHDps9F1dn0xaDCQ2ZNRfVVYZ0zGXK0tmUrpl26olip45r004Q3iX5LskL6lQiArqPPoHMQ3e53t47gktKJN52BF3KK+3pVHGsGJ2jXNrkUbYBtqTfAJrZJbJPYZtdSksxiK/OaEBPfxDI+8U3jTIYer8Q4a1EtcU7inN0GGcQ59dO9Mues/mJTvV05cU7inMQ5u5aSJCZrA1tGE+fEMj5xTuNMhh6vxDlrUS1xTuKc3QYZxDn1070y56xsLm+pN5cnzkmckzhn11KSxGRtYINv4pxYxifOaZzJ0OOVOGctqiXOSZyz2yCDOKd+ulfmnJVHAVjqRwEQ5yTOSZyzaylJYrI2sB07cU4s4xPnNM5k6PFKnLMW1RLnJM7ZbZBBnFM/3StzzsqDGyz1gxuIcxLnJM7ZtZQks4nJvOb5xDmxjE+c0ziToccrcc5aVEuckzhnt0EGcU79dK/MOSuP2bDUj9kgzkmckzhn11KSDO0w76gD4pxoxifOaZzJsOOVOGctqiXOSZyz2yCDOKd+upfnnD9sjvxmtdGLCg1qR82QS5ZD5TqQJw6VbkGediXNmCnPoyoeCnYIVrWmOshiE+Mfgv3udIz87rjabD5Gz//u/mX53/3h/SyMwkjiOoRIs+NmmX7RT65Frz9Hb2R+cnU8pS7PN08bZTRbk32BOT3N9qrR4zBwpmOPdR8WTnLXLWrMOZCt9zpHO21uMYh+czD0fIfpa7WdNdfGjQIxR1Wj/DP2UO+STyAEF4TkOu0mD5VqtQsL7sphCYgYDkSELNxtKNJm7PQZjhiodzRI4tle4PvMqqZuoAT1VtVgCbeXchaWwBspEyzBhSW5ZoyZWLTgIV45LMESw2GJkIW7DUvajJ0+wxID9Y4GS/wgnO3Zm0qzV9uHJai3qgZLuO02s7AE3muTYAkuLMn168rEog0P8cphCZYYDkuELNxtWNJm7PQZlhiodzxY4vq+N2LO9bZusATzVtVgCbcjWxaWwNuxESzBhSW5li6ZWHTgIV45LMESw2GJkIW7DUvajJ0+wxID9Y73JU7gTfz88ubLPeoFS1BvVQ2WcJv2ZGEJvGMPwRLktSXZXf+ZWBzBQ7xyWIIlhsMSIQt3G5a0GTt9hiUG6h0Pllj+LMguzbjdo2awBPNW1WAJt69DFpbAmzoQLMGFJbmNoZlYdOEhXjkswRLDYYmQhbsNS9qMnT7DEgP1jgZLAs93vPzmlss96gVLUG8VBkvKl7rCV7i6jaKQFqaTPkCfyo186f3y+u4NzO3Bp53RVejs+NdFV1YSrce/FsfsNSU4Jd1nwXJQTG98i6U07sMGDVLRzrGaHv1r6u1sVY+/szWHZDlT9Z4G0Ipqr2V66qi7G96+qu19+JLVhK6pJ9XtSYUbYTv1sey2KkwmxmuBDEyoF4Kl3guBKFkHKJnAZmb5Wa+tHdIgsNPvXhEGUTNZ8xsPm+okZ5JxT/RMI3omYDtTNd84QZOeqjrq8oZTtPb7kmhO0upXENE0EE2r+MJMvTcM0bQO0DSB5g7yc19bHSNAoKffvXMMommy5jceOtVJ0yTjnmiaRjRNwHamar5xmiY9VXXU5Q2nae33adKcptWvIKJpIJpW3ivLUu+VRTStAzRNoNmN/NzXVgcdEOjpdy8xg2iarPmNh0510jTJuCeaphFNE7CdqZpvnKZJT1UddXnTaVrrfet0p2m1K4hoGoimlfcOtNR7BxJN6wBNE2j+JT/3tdVRDAR6+t1b0SCaJmt+46FTnTRNMu6JpmlE0wRsZ6rmG6dp0lNVR13ecJrWfh9PzWla/QoimgaiaeW9VC31XqpE0zpA0wSaIQIW/LfUYRG206PXvWYNommy5jceOtW6N00u7ommaUTTBGxnquab35smO1V11OVNp2mt9zXWnabVriCiaSCaVt5b2lLvLU00rQM0TaA5rPzc11bHWRDo6XfvbYNomqz5jYdOddI0ybgnmqYRTROwnamab5ymSU9VHXV5w2la+33eNadp9SuIaJowTfvXYfPE7fAYvajQ2HHcDCszieM4g+iXTdwuF89uPw/A5T5pGaAvqqSlQKrA0kJyOaxeMb/VK8ZQtiebZ1X6/gqTy8fzU4OOyhEYCs6KhpjeYs7ZQhhAUNjDJoPoV9DDxi0evIN4o0AoUNX0+QwJ1Js+EzYoBPx4PhssuC0ksdABRAoEH0DkABACRAwII8AFSaIEeUE9wQlqrSc7jBRgHkNYgells/E88MS9rE20gHqraniB2300ixfg3UcJLxR7mQXj+ZizSd5ihj2oYxpACqjbJ0AOpJkeQAwIL8AFSeIFeUE9wQtqPdA6jBdgHkN4gbM3aDaeucJe1iZeQL1VNbzAbYOXxQvwNniEFwphP7cXiwlnt6bNDHsIXoBIgeAFiBwAXoCIAeEFuCBJvCAvqC94QakZT4fxAsxjCC+wv+3yPG+2EPayNvEC6q2q4QVuP6YsXoD3YyK8UGzCF0xmcw5NcJhhD2r1B5ACalMLkAPpAgkQA8ILcEGSeEFeUE/wglpXiA7jBZjHEF5getk8mA85qyVZXtYmXkC9VTW8wG0MksUL8MYghBeKX0NOFgPPYYf9iBn2oPULACmg9QsAOZD1CwAxsPULYEGy6xekBfUFLyhtT+4wXoB5DOEF9qKAkTfy2d96sbys1fULmLeqhhe4O9SzeAG+Q53wQnG/23Q+GHPC3mWGPWhXHUAKaEc4QA5kwyVADAgvwAVJ4gV5QT3BC2r75DqMF2AeQ3iB7WXzxWzGnoRZXtYmXkC9VRheKF/nCF/eOGkGHlADmzoBTcVDQdBL5ZAQqFI5KACXVI4JAiGCo0oijmrn6wO8aHzbpVIKgM0Yvhv9Cj7jcCo9tVVgILwnFkRMFkZm6nGDnVZsqN6rx3gjsIDxVeP3F2vkrpBdCik9/hFN6ba0mTqzITuj3lJk4taCTICjgh2jbn2333AHSOeEtrtb6tvdid91gN85wWQ45+6zBTI8gUFB7Xmqh4X046keFdaAR3Rc2Y47VeP2hOu1sHW+DbbnBVbAXpBHfI/4HvE9bYxAfA/FLt7cH3FWFOnG+Npvq6HO+VRRCnhcsIPUr3XTmV/FF3rqjUuI+XWA+S0Go4HDiVALyvwEBgU1UqkeFtI3pXpUWJsU0XFlu6JUjdsT5tdCE5QWmF8w8T3fE35KYn4GgFtifl00AjE/HLtY3tybi6f1Fplf+w2S1JmfKkoBjwsvDdSuddOZX3kLKku9BRUxvw4wv+l8Ph9x9rLbUOYnMCioxUX1sJCOFtWjwhpYiI4r26+iaty+ML/m21m1wfxGIfdjVzhZT0nMzwRwS8yvg0Yg5odil6iHgMcudTHTeovMr/1Wd+rMTxWlgMeFLwKuXeumM7/yZoKWejNBYn4dYH6TgeukdptmItSBMj+BQSHMT2BYAPMTGBXE/ITHlWR+leP2hPm10JiwDeZn+UHArnCynpKYnwHglphfF41AzA+H+Y28wGcDe2Zab5H5td+0VJ35qaIU8LhgB6lf66Yzv/K2sJZ6W1hifh1gfs58tli47AgdQZmfwKCgfX7Vw0L2+VWPCtvnJzqu7D6/qnH7wvyabzHbzj4/N5gKPyUxPwPALTG/LhqBmB+KXbyZ7we2eFpvc59f6+2nEfb5KaIU8LjwfX61a9105lfe4NtSb/BNzK8DzC8YT12HE6EulPkJDApqOF49LKS/ePWosHbiouPKdg+vGrcnzK+FZuFtfOfnBw6nBM56SmJ+BoBbYn5dNAIxPxy7eP7UY5e6mGm9RebX/kEC6sxPFaWAx4U7SO1aN5X5le/vg2/rmzZD9IyiTVnjJu6dsy6IOokNDKJPYkNDKJTYyLAEJTO2bJISGbsndKqFwxE2OTC0KYdHotZSJCAmhrzlGBLzPNSIKBMMLHLwOx0B6JxaT9dXdSOa7sj1q2sRrfp+bS5a4keqYdUQ80bOf/r6QFV9BNd6OhZZEE1dVU3pLugyfeJRdSKtjrQhz2qdwaOMrRUsQvRwYEVP6LQeW/20HirxUYKgEl/HS3ytnImjC9I3P+ipyIcwpedOnsjGAJX59PV+U6a83jg/FfpMLfSh50B9vYBKfajGpmKfqdOPqhtpdpoZ+VbrbL575T5UH1cr+JUf0marH9JGBT9KEVTw63jBr5Wj0HTB++YHPRX8ECb13IFD2Riggp++3m/KlNcb56eCn6kFP/QcqK8XUMEP1dhU8DN1+lHuwKTXIZbkW62z+e4V/FB9XK3gV7F3V/1sTir4UYqggl/XC35tnICpC943P+ip4IcwqefOmcvGABX89PV+U6a83jg/FfxMLfih50B9vYAKfqjGpoKfqdOPct1Yr7OLybdaZ/PdK/ih+rhawa/8SGZb/UhmKvhRiqCCX8cLfq0cfKwL3jc/6Kngh9KlI3O8aDYGqOCnr/ebMuX1xvmp4GdqwQ89B+rrBVTwQzU2FfxMnX5U3UizI+vJt1pn890r+KH6uFrBbyRW8LucjE4FPyr4aZgiqODHeL3r593rgvfND3oq+GG0McueKp2NASr46ev9pkx5vXF+KviZWvBDz4H6egEV/FCNTQU/U6cf5R5+I2/ks+vGLO5ABb8O+lbXC36oPq5W8HPFCn4uFfyo4KdviqCCH+P1Bgt+gec7nG8wWMedU8FPr6Cngh/CpB6Mp67D5j8uFfw09n5TprzeOD8V/Ewt+KHnQH29gAp+qMamgp+p04+yG80XM85CURZ3oIJfB32r6wU/VB+XKfh5y8OXHzbHU6HKF71wF78CLOyNB80U9pLZWHkmb7A22MECzyD+yTnw+ZBp5UoOGuS6XixLiUPMnNg05BI0g2yFATolqupSwHh6QN0SvaevfXhePq1BECVDeesJA7YmcQ2qpTnm8uZIU09UHqiqYPODA2ANKDdUj45uqhJAhfqtSgjiTr5fH/KR9+W79UtoEwQnCI5xRi6BcALhXQThlmMHLnuRAcHwNmC47Y6CKXubFwHxFgKkAXv0B4o3pcxegHFcZSrA8eKR1QU4Dj6umuB4n+C48Al2BMcJjncRjruW5Vg2JwAIjrdwnoJju7YjbhCC48bboz9wvCll9gKO4ypTAY4XD5QswHHwYZIEx/sEx4XPlyE4TnC8i3Dc8d2hxW6yb7NyOsHxmg0ydqeWLXKMC8HxrtijP3C8KWX2Ao7jKlMBjhePeyrAcfBRTwTH+wTHhbu/ExwnON5FOG4H9nDE/sbTYeV0guM1G2QUONPxTNwgBMeNt0d/4HhTyuwFHMdVpgIcLx7GUIDj4IMYCI73CY4L92YlOE5wvItw3BqMJu6YEwAEx1tYOz6cONO5uEEIjhtvj/7A8aaU2Qs4jqtMBThebJVcgOPgNskEx/sEx4U7pxEcJzjeRTg+HTvjAS8ACI43D8d9210M2DUvpkEIjhtvj/7A8aaU2Qs4jqtMGTgex++nt3jgMAEU0Pjl9bvLG6BY/IJMWsDiOUCSpKw0ImkJhSt1f8/tlUyeKrVZsiy98watUFU5agUPWjLVg8csbYaK0vib0wxVaeyHgl90kqv5bvTLpAjpa+fGrcOpafRNJWbb7T6e9dGzXVj9eiEEjtluXrloAmCEyocPNdA7bTqV1jXrhIdm1Fsf4FKfDC4XM3ptuTEewLiMcxtMt20L9SdpKA0ieeIVm/hHcBp0gec/lBAoiZXH0a/gjQLqSrt1hCu4zi0I4IUkfa1BkgLh4na0zBMv9caWxMBMYGC5jpSZQRU4mMCwABYmMCrxMJ15mBdYAXt/KzExYmIdZWLWwlmM2Z06iIupcrGccnNTwuVyvWysAQMTH9PJGmiMbD5ZLHyRe22fk83G88DzhW+VWJk8Kys2NuWxMnh/U2JlJrAygUEhrEwWhaKNSqxMY1YWTHzP53XBLWZ2YmXEyjrAysZja2Gx18Aw+ycSK5NIrznl5gLrcrleVtaAgYmV6WQNNFbmj+aTOXujIWtCbJOVecFsPGMvwmbdKrEyeVZW7G/LY2XwNrfEypBZWa53VWYCcqCsLNefNjOozUq9aMMCWJnAqMTKdGZlo5CXsett2YaCnWNlArFLrKyjrGzkj0c2+3xAZhtNYmUS6TWn3NyUcLlcLytrwMDEynSyBhor81zfnot02G2flS08z5uJ32o5K1NnMMWWwDwGA+8MTAwGmcEI4Hd5BiMArSAMRhaxoY1KDEZnBmP5QcCuTWWXPHSOwcgyemIw3WEwzsKeu7yu6TiQqr8MJqfc3JRwuVwvg2nAwMRgdLIGGoNZLBYDj32+GWtCbJPBzIP50GMTQ9at0vdK8qys2Bmax8rgDaKJlSGzslzXt8wE5EJZWa6zc2bQESv1og0L2YNVPSqxMo1Zme8FbsCehLKtODvHygRil1hZR1mZNXZnY3ZFltmAlliZRHrNKTc3JVwu17wHq34DEyvTyRp4e7Bcz/PZm5JZE2Kre7BG3shnk13WrRIrk2dlxQbhPFYG7xNOrAyZlQlwEnlWJgAXIaxMFoWijUqsTGNWFviB47MnzOw3aJ1jZbJVCmJl3WFlc3fkDtjQi9mHmFiZRHrNKTc3JVwu18vKGjAwsTKdrIHGyoK558zZnTFYE2KbrCyYL2acc9JZt0qsTISVRScy8alY/CqUfl12dhP96vQBTY03/a514ik9vslqBb1NfXtms6cT5pbexaJmrJy7oQziKew6v96NInLmKr861jUhJCyIzCKKQDQGHao/Z9ssBtGvYKay8Q+2kVnAFP6I3qhddqNQTFDdwDh9liO8ezGBhH6AhOY70hJMIJhAMIFggrRFPdsLOB0BdAMK3twfBUPxW60TKpR01UxDBXhLTYIKvYAKLbRJJKhAUIGgAkEF+QNegxAssL+SYOWqNqFCYHlzby5+q3VChZJWb2moAO/zRlChH1Ch+d5dfYMKYcT4E4l9n7VDhdwNZaBCYWsyQQWCCrpABdf3vZFwrmoTKvizYOixGRjzVuuECiU9ldJQAd5QiaBCP6BC801y+gYVxv504Ug0uasdKuRuKAMVCn0YCSoQVNAEKniBN+HslGPlqlahwsgLOPspmLdaJ1QoafSRhgrwLh8EFXoBFVro3NA3qBBYY3vAPqWEuUK+dqiQu6HyTRwEFQgq6AIVrIisC+eqVtcqzHw/sMVvtU6oULL7PA0V4FvPCSr0Aiq0sJ24b1DBdibejF03ZbY4qR0q5G4oAxUKXXgIKhBU0AQqBJ7vcFqNsnJVq2sVPH/KaeDKvFV0qPCvw+aJDxHiV6HIwCZkUNaUpr7uKQ95Ud2EJCp7h0CApGxyE1+RGP8I3jVgE7rcFC8XKHo+cf0NOYQfNadO3qMmkGkuP/HU3ptCn0dFa/wwGUS/gv4H6KWABgYQbxQKBao3Q0bvUt8MSdiAsEGd2EBtu1B76GA+WSx8dpOazuKD+p9ZI4Rgu6NgKuKYXcAIDTwsGkqYjechGRf2wjZxAuqtKiKFkr2QaaQA3wtJSIGQQp1IQW23UHtIwR/NJ/Ox8H13AinU/8waIYWpY7s2GxYxt64ajRQaeFi8Y6OD2XjGXmDN8sI2kQLqrSoihZKtkGmkAN8KSUiBkEKdSEFts1B7SKH+Y+71Qwr1P7NGSGHsTi1b5GG7gBQaeFi841k9z5uJe2GbSAH1VhWRQslOyDRSgO+EJKRASKFWpKC0V6g9pFD/cdL6IYX6n1kjpDAKnOmYvRuF2ePCaKTQwMPiHRlY++noeh7krogUSjZCppECfCMkIQVCCnUiBbWtQu0hhfqPONUPKdT/zBohBXs4cabsr8WYm1GMRgoNPCzeOoXaT+zV83BhRaRQsg8yjRTg+yAJKRBSqBMpqO0Uag8p1H/snn5Iof5n1ggp+La7kOlwYTRSaOBhEQ+8rPsUST0PvLwhhcu/jv/4f1BLAwQUAAAACAB4JGtcYHmC0zk1AABzrwYAGgAAAHdvcmQvc3R5bGVzV2l0aEVmZmVjdHMueG1s7X1dl6NGsu37+RW16sVPnpYAIcnLfc4SAsZey+Pxmfb4Pqur1F2arpLqSiq37V9/QJ+AEsiPSMiE7X6YKUAZkLkzc8cOiPj+f/54eb77fbndrTbr998M/zb45m65ftg8rtaf33/z71/jbyff3O32i/Xj4nmzXr7/5s/l7pv/+e//+v7rd7v9n8/L3V3y+/Xuu6+vD+/vn/b71+/evds9PC1fFru/vawetpvd5tP+bw+bl3ebT59WD8t3Xzfbx3fOYDg4/L/X7eZhudslxuaL9e+L3f2puZcNX2svi4fz/3UGg0ny92p9aeP2jjavy3Vy8tNm+7LYJ39uPye/2H55e/02afN1sV99XD2v9n+mbfmXZn5/f/+2XX93auPby32kv/kuuYHvfn95Pl+8qbr2eKOn/zn/Ystzk8efhJuHt5flen+4vXfb5XNyw5v17mn1eu032daSk0/nRiofOPOwX1+Hntqgh9vF1+R/rg3y3P7j8Ucvz8c7r25xOOAYkbSJyy94biFv83wnWfB9leuabOd+Vuvbv283b6/X1lZqrf24/nJpK1kGRNo6jVH20XZqN/PhafGaTKCXh+9+/LzebBcfn5M7Snr8LkXk/X//191dsjw9bh7C5afF2/N+lx45HNv+sj0dOx46Hzz/dfw73qz3u7uv3y12D6vVr8n9Ja2/rBJDP8zWu9V9cma52O1nu9UiezI6HUvPP6UXMn/5sNtnDgerx9X9u5z13V/JVb8vnt/fO87Nqfmu9OTzYv35fHK5/vbfH7L3mTn0MTH5/n6x/fbD7NrC9+8y3XD6I9dRiYFXVt+9Fvpu97p4WB1uZPFpv0zWtmT4U6vPqxQ0ztg///Gvt3TMFm/7Tf4uXrN3kTeZHikM6uG598ki9uG4FyUXLD/9tHn4snz8sE9OvL8/WE8O/vvHX7arzTZZ3N/fT6engx+WL6sfVo+Py/X7++H5wvXT6nH5/56W63/vlo/X4/8bH+b/qcWHzdt6f3ygSwc97x6jPx6Wr+minFyyXqTD/HP6q+f0J7uMsUMbb6vrLR0PFEwfDv7/s93huaPKTD0tF+mufTestTYltOYwGxdvxyVqxyNqZ0TUjk/UzpionQlRO1PFdvabhyNSs224U56f3UCO72c3COP72Q2g+H52gx++n93Ahe9nN+jg+9kNGPh+djP29T97WBz+vvnhSAw1v672z8va9W1IsZye9pm7Xxbbxeft4vXpLuUFN6bqmvnw9nHPd9NDgpv+sN9uUvZbY8txCGxFL69Pi91qV2+NYjh+TVne3d+3q8dae6OS/a3Gwi/Pi4fl0+b5cbm9+3X5x16qkZ83dx+OHKh+wAl65afV56f9XcKHH3ks+iUDwWXkp9VuX2+h5KG4LHANrl8C3RoL/1g+rt5ezj3FwZF8l8KOU2/HU7GTDgrPw4yUjXA8ia9iJB18nicZKxvheJKJshG33ojcKhUutl/45uJYbrbPN8+b7ae3Z+5VZSw35y92+B5GbtpfjHCtLWO5OZ9bhO9mDw+JQ8oDZdXVWMCU6rIsYIpmfRYwSLNQCxgkWLEFrMkt3f9a/r7anQm3+LjvMry39hbdkg4RYjL/+7bZ15Nkh0K6+HG9X653yzs+ky4Fe83tpAKDT7ClClgj2FsFrBFssgLWFHdbfktE266AQYL9V8AawUYsYI1wR+bgfVQ7Mocpqh2ZwxTtjsxhkHZHbsaHErBG4EwJWCPcAjisEW4BzfhZAtaItoB6S8RbAIdBwi2AwxrhFsBhjXAL4PDKqbYADlNUWwCHKdotgMMg7RbAYZBwC+CwRrgFcFgj3AI4rBFuARzWCLcA/ZobvyXiLYDDIOEWwGGNcAvgsEa4BXjNbQEcpqi2AA5TtFsAh0HaLYDDIOEWwGGNcAvgsEa4BXBYI9wCOKwRbgEc1oi2gHpLxFsAh0HCLYDDGuEWwGGNcAsYNbcFcJii2gI4TNFuARwGabcADoOEWwCHNcItgMMa4RbAYY1wC+CwRrgFcFgj2gLqLRFvARwGCbcADmuEWwCHNcItwG9uC+AwRbUFcJii3QI4DNJuARwGCbcADmuEWwCHNcItgMMa4RbAYY1wC+CwRrQF1Fsi3gI4DBJuARzWCLcADmtyq0n6Dvbz8o77heUh5Vsm/K9Jk7wAfnzUfy0/LbfL9QPH6y0UVs/PKmCW4g30YLP5csf3SYBbghwxe6uPz6vN4aWoP28MjGvfYP/n/O6H5eWdysL3E4wbST94y37edjh2+u46uXz/52vS6mv2Na3H4zcLp3fLDxf++Hj5CO1ye+n93J2+FTydu9776S6uB7a7ZIqerh4M4rk/dePrDR6M1N/Z5V5OPTBk3831G7ar/Y+LZKz+uS694fXyj33pyefV+sv55Nn0/GmxzVxyHYjzhVO57jicznwRmfz1Zbl8/Tm5v3eFYz+t1std9uD1w8mPy0+bbdJ93uSAztN3lJc17nD15m2ffkT50+/Plzu53ELuI8rc163fl33buvhPxbet6cnSb1tzv7x+25oezn/bmo5j7o957vEf0v3g/CyuP4qnBwQf2jvsFe/vF4dN4no43RjTORnnjGQ+n50UTmQ+np1ke+vUQwpgdqrB7GgEsyME5vz6ZwDIT58Hc4J82CGQe/FkGIRlIC+BtF8OaZ8W0m41pF2NkHb7BGmnb5CmgadXDU9PIzw9IXheSWlnIOvaDdlV7g8z4DyqhvNII5xHfYezZz6cc7B0PDc+itMc7Hgc0wLVrwaqrxGoft+BOjIfqNxra6sgHleDeKwRxOO+g9jvEIi9QfqvCOJ90o1XCP+6SvNEBcQInlQjeKIRwZO+I3hsPoLVhYZB4URGaBjQQnlaDeWpRihP+w7liflQ1roYa0X9QwKuxUMyDhWBmVOOqcun9ocMU8z5UJKNqgq8Q3HwVj/RPk3BVPE0hxRN9bGmu8N11fNOduLtPz7noJv8/eM6nXlfT8G+45M8/rHIDXVy2Xz5/PyPRT6b5X7zWv3T48qy/LQ/XjYcTKou/LjZ7zcvHC1uD2/21DSZjlXxvk/HeOC5fnv5uNyeYpGlccNDbpaSsTwmbqEeRpmt5OfNOedW2a2ez/POF7UF/CYL6mG0TzlQvcsftzlQM+uwwOLy8LZLcHWIERdHMBfyZHbOD+eI611hNyzstsylqnJ7HXJvrTWda85uZHUIUxAzTj1mHHLMOD3GTPsRQUGEuPUIcckR4gIh1QhRdMuOr1MxB/V4SoM/dmi41hkbZt/zU9ugX4PHPNO7ULPD79Mc86d3yv5KvaS745aevpRzGM5jv/PO13d5eyx+4A54GcIJFOvUsXlbPJ94jfFuXA7Gw3GyPd50XPpETt3WeOm4vCJ+cpG3Fyze7JyXnzilC+bI0bZgXgFePrHoVsriPK2ZStaskx0FEXsdvuSNZiLmclbDanxuu35BpvOYEm+0UEli9cx47+vYlZmLzV3xaF8zYAF3OCqBp+OVwtPxtK1xOdhUgpZupWNMgxqYWrPYdQY/7OUt1Y6uGUaZcClkIeVf6W4h4HpkK9XqICemml/68cugsEHV8jKZvgo2j38eEtIzuyk9e8xXz99D2Tl0br0+GMLz2mW+L2ezYTgJ+XWyocN6j51mfco9Z3VP0i1Ql6Hj7tjyDlSBTskb6tcnFnlHnfWAHO+hNwOfixt1+n6iIaE13w91nU0PsBrhTD/CSl4Yvz60yCvjrCfkeC28rQWKQR2uu+mwXKIb6pPo8r1WNzT0eKyR6fjw2GC/lrOUcnKiREnowZplJu7x5bqnxfpzWsj18HcDTCXtlZKt5lREpOEucx0/ng64umzstNZlJWvnoctEls2mu2w4mLTWZ8Hb8/OyYnLenS4wq/dudY7kyI+X35cLHc10Z9XcPV5h3BSu6VGn5R6tmtqnHjVthtf0qNtaj/58eGWlokNPF1jVnaOWu7Nqyh+vaH7KO1PfnZYTnZoe9Vvu0aopf+rRxqe8Wo+OW+vRedL0av1WEgU5dOnlErO6tMp1ZNL1hnjTubuq5v35GuNmvlCnNqTOZju1aupfOtW0yS/UqQfK30Cv/mPxsN2Ui94v6ekSCeLyUx2CUU1f7hcfd7l1NDlw/nHagekzvm52ybY/zmxTlVcOh9lwc/Wl42zIuvJSxx14vJdOskNeeanrjXgfy0t4U35bufadSFQ3zSX2tl0dBbFDtO16JK8PXVwCfa/5VwhyeVQyQX24hDgAcZ1HknrcDeBNHgz2WnKs88fs8uMp/vWY+yWKQ8O1C5CjkmkqPxDc8eLB4T/2hzK6wH/tjfJRoMN8cVBr+r073ZzLUMLs6fN7uR75e7le9QKTOcv6EMSa1zI+5v4wIrOIIDpG9egYkaNj1A90NJrjQHDc/fpx98nH3e/HuBuT90IQE+N6TIzJMTEGJppLIyEIiEk9ICbkgJj0AxCGZWUQRMa0HhlTcmRM+4EMe5McsB3u+eKQ+ZqNlofTSSKnm/Gy76gGF3qydlxlVLEPvRlYrHIylFeRYflHxUPFj4qv3wLst5uyr/FP52RXCYYzn41SqIkoJR2v2BuXCgDM/ricJewRlS8lufQOxQXiVC+gQpg7VxTQJtBlb6FWp3Nb+vTU0/jpKTtlkDMpj/1M3UOFjkN2kuNfyquZ4ZLJDUjqsUrHgXKThBudFOudMUOT+7jseVm9kBarvNCtp8MWZPrJYHJ6u7KO6qnKBEW0V/fyTVkbwm1L5XtSq4F9rZtThezrVXR97tL1+S7ZbJ8T6l/eofPBaOCVdGj+k+q3wn5ICvCa3r4tZkTY3dq5KvlQVH4ur2ec0rpOFXlIMmWfCEfGbW9kyrpYNZXLP+fnelPMfswWpCrtSEYyL1F/vJ0smrfpLqfZfuV6N+mS8fDap+mRtGhdSZempw9F7cp7NJsmsarfRgJBavr8c4eG5NMpBpvt43JbeBfqkE6xxs0ZZNycfOKbI0E+JltUa4TX5app5pymUa2V1ToZ2uUPRO38ptLOKX1kYey+72V+zNupf6i3eyrHWfaeZ6bUsPoC4Av4dZoWgPzGxv2Cy/ngTQKezI5WtsAcHPF/bb4Gi/Xjh9Vfl84dFpeYw4WJ2doLdSxZk5IJxfHaD8cipNR6v2ZxDg2/bC+tfFptd/sERveZDshMksI0OYth+ezZfHOmMGuK86ZACW9J4bviNDs8Ww6cD4Xm9g83YNUK15utd716vjmvDdAFvJTeQGEnLb/kt5JLDsAqdu3x4C957J3QVgXA5wXwB/y1h7/DApg80D0BLMRQ37jRjwkDGP623O7vSVBcB7SWgHBcMp4uLPDhebnYFrl88uen1fNB4En/XZAdHw7m2Vl67Cghu3Fhx5XA22EQfths/8Ig6B8EFd/l29lJwa73Ye6Ol1aV4+6GMyOZr73z7gxnoFl6/xUIZMOlgUtDClltpFLsDiyjlXBrgMG2MQjXptesOnTDOIoKrLrI1eDcWDwMBO5NaX4ThntTkeakG+7N1HN91yt726O/7g3nWzDSuzDvWzZwb+DeUENWG7UUuwPLqCXcG2CwbQzCvek1r47ihFlfWVmWV+ePwr2xdBgI3JvSTIMM96Yi4WA33JuxP3XcOXs3cHvs3kyDIBhNy/pF3b3hbB/uDdwbcshqo5Zid2AZtYR7Awy2jUG4N/3m1X4UhSMmr3ZzR+HeWDoMBO6NJ+DeZDOPdtK9GcXedDxj7wbXoE7/3JvJwPdmTlm/qLs3nO3DvYF7Qw5ZbdRS7A4so5Zwb4DBtjEI96bXvDqMw0k0YfJqL3cU7o2lw0Dg3owE3JtsNtNOujfucOJNA/ZucHVQ++feeMFsPvfL+kXdveFsH+4N3BtyyGqjlmJ3YBm1hHsDDLaNQbg3/ebVTjSL85933HI1uDcWDwOBe+MLuDfZClGddG8i158PSqI3102if+5NPJ76XskuWSwiK7MLc7YP9wbuDTlktVFLsTuwjFrCvQEG28Yg3Jte8+o4jLywmLCryNXg3lg8DFLuzU+r3b7KpzmcV/djsmnWjEn4breXwZ+PuTyzvMG5nm+nNBJJ99U1upUe4sN/xVH+uHj48nm7eUu2nXs2h+DcgriX8wLasmkwlbfPnjsNj5u3j9fp7qutJXrXQd0roda1EG6GIW5GYynGgX1N2Cd2eAAIWwAh7XrxZKxOr6NMVw1fLHtZ48mkxSeeIamqZSceMmHDJ2vSJyvgLZ+9E15ZE16ZfJZoHTmgG84yTb3owjuz0jvDHDB0DrTtpQEYLQND1VurTMCd9dYosm+Xe2vzYOD72RQR8Nboc2OLTz9DMm/LTj0k9oa31qS3VsBbPhkpvLUmvDX5pNc6Ulo3nDSbetGFt2alt4Y5YOgcaNtbAzBaBoaqt1aZTzzrrVEkE4e3lr2s8VTf4tPPkETislMPecrhrTXprRXwls+tCm+tCW9NPoe3jgzdDecAp1504a1Z6a1hDhg6B9r21gCMloGh6q1VpkfPemsUudHhrWUvazxzufj0MyQvuuzUQ9p1eGtNemsFvOVTxcJba8Jbk09JriPheMMpzakXXXhrVnprmAOGzoG2vTUAo2VgqHprldnes94aRap3eGvZyxpPxC7xIrIZad5lpx6yyMNba/S7tTze8plv4a014a3JZ1jXkT+94Qzt1IsuvDUrvTXMAUPnQNveGoDRMjBUvbXK5PVZb40icz28texljeeVF59+hmStl516SIoPb61Jb62At3wiX3hrTXhr8gnjdaSDbzjhPPWiC2/NSm8Nc8DQOdC2twZgtAwMKW/t79vVY5WXdjiv7pxlE5PAOUM6/pbT8R8aL1Tn0NP8bxqah0tpnku5jTfr/S5te/ewWv2aDt77+5fFfzbbH2YJENLGlwldnO1Wi+zJ6HQsPf+UXsj85cNunzkcrB5XxSFp3GHqUn7oodkJolmLFUcpofaTVFuhNfRt4qLKBeatRpXFjulEqvHY8cjY+u1bQej1IRSP6R4g7oTCSPNB+u9iKVtALHvM2KKcwF27VKZBnmIBsB0AG8Am38KllXye6k7pdZTVnSDtZy9DdSdDqjuxvG9dBgSXDtSnyi18kPlt9/XtKTBSKvWbU2FEq2zYTgEcCP4tTmEUUMMMhvQP6R90wMa1pFMhAABDMzDEFNPQDeMoutjK163NHu1OMAAI1EJzGuUwVoC8zcAAQG4/yHWHCCpLimZDBBQlRREiyF6GkqKGlBRl+em6DAguHiiKmlv4ECKwXROwp6pdaYjAnLJ2WgXGdqouIkTQ4hRG1V7MYIQIECIAHbBxLelUiADA0AwMMfU0ikM3ZBd0yR/tTogACNRCcxrlMFaAvM0QAUBuP8h1hwgq69hnQwQUdewRIshehjr2htSxZ/npugwILh6cBhAiQIjADk3AnlLKpSECc2opaxUY2yn1jRBBi1OYL0RgzxTGDG5hBiNEYPEjgw7Yu5Z0KkQAYGgGhqB66kdROLrYyqqnbu5od0IEQKAWmtMoh7EC5G2GCABy+0GuO0Tg8YYIsvo9QgTGhAj4i73LzHCR1mXmt0j7ErNbpHmpEIG4AcHFg9MAQgQIEdihCXDPGN0rVu2aVRoiEDOhc9nSKjDyr20IEXRkCvOFCOyZwpjBLcxghAgsfmTQAXvXkk6FCAAMzcAQU0/DOJxEk4utrHrq5Y52J0QABGqhOY1yGCtA3maIACC3H+S6QwQj3hDBCCECE0MEXjCbz0tqVo8KfoJEKjGB1qUSiQm0L5NGTKB5uVoEwgZEs5TxGUCIACECOzQB7hmje8WqXbPKaxEImdC5bGkVGPnXNoQIOjKFOWsRWDOFMYNbmMEIEVj8yKAD9q4lnQoRABiagSGonjrRLM4nZL+ayh7tTogACNRCcxrlMFaAvNVaBAC59SDXHSLweUMEPkIEJoYI4vHU90rQ5Rf8BPEZLtK6zPwWaV9idos0LxUiEDcguHhwGkCIACECOzQB7hmje8WqXbNKQwRiJnQuW1oFRv61DSGCjkxhvhCBPVMYM7iFGYwQgcWPDDpg71rSqRABgKEZGGLqaRxGXji42Mqqp37uaHdCBECgFprTKIexAuRthggAcvtBTh0i+MfycfX28uFp8Zjc/JAdHzhec3e66O4igSsEB7KVDBAcoPl+YJD+K+Jqv/wjU379uJYFccFhkIgGyhuTCg3Km5OJE8pbk/v2QMoewgDmhQEqPO/DgcOAn8ERH/4rDvvHxcOXz9vNW8KH85bbe7NPcjo0vLQ0vrg0vbxIyoeFSwio8+DwX4E6H+9fmSMbFRIwVZTHjOz8jNQpyLciidMblVQtuZe5+SD9x1zmsseMFcFM2Cpa6UNCjaWdya3mxJ9e9uN05s+v/MGrN9KrHwezwbykcqUGv17JnMxWr2RQYqtXsifl3ctahH8P/74R/15+SjS+yLSwzDS/0BhD3rx4Mgyuj5ANkcHTb8bTx9zszdyEx9+6xx+6YRxFJQte9ih8fvN6EV5/2sOOmNef/UgPXr8xXv88HgfjkmJUTuUGJbXpK5mT2fKVDEps+Er2pLx+WYvw+uH1N+L1y0+JxheZFpaZ5hcaY+jbfDAaeGyv31FnavD6Obx+zM3ezE14/a17/VGceKxOyYKXPQqv37xehNef9rAr5vVnPXZ4/cZ4/YE7n09K6ku4lRuU1KavZE5my1cyKLHhK9mT8vplLcLrh9ffiNcvPyUaX2RaWGaaX2iMoW/TIAhGV+coS99cdaYGr5/D68fc7M3chNffvtfvR1E4Klnwskfh9ZvXi/D60x72xLz+rEsOr98Yr38aT2ZBiSztVW5QUpu+kjmZLV/JoMSGr2RPyuuXtQivH15/I16//JRofJFpYZlpfqExhr4VKhrnS2nD62/C68fc7M3chNffutcfxuEkmpQseNmj8PrN60V4/ccyQkJe/6XqELx+k7z+8WQ+CD32BjWq3KCkNn0lc1If9akYlPmkT8We3Hf9khbh9cPrb8Trl58SjS8yLSwzzS80xtC3QpHCfHVMeP1NeP2Ym72Zm/D62/f6ra95bcK2YX9RZYu9/pLSvWVeP0UBX3j92ctoCvhOg8G4ZIPyKzcoqU1fyZxUfQ4VgzLlOlTsyRUBlrQIrx9efyNev/yUaHyRaWGZaX6hMYa+FeoO5QtewetvwuvH3OzN3ITX37rXb38ZSyO2DevrJFro9fNl8aNI3pf14uHkCzn5w7INKU9XandSznbgQcKDJPUgufF7Qz5ZSygBwgsAqlmtUQOtEYjnIHz749Z8KYCUg7vlV5wjSEsXnBbcFRMXSdZIAVeNLn4dAFQdYno/zpLef6f7O5yk/yrW6+yZ1Atcpr9pTbYw/rnWy9TBoAEYaLRe4XP9tThWlUy0fZ4AFCigQE0dE6pw6VBWuIRclr0MchnkMshl/VzhxRhgj0oJQjCzF6YQzCCY2bn8dQBSnZBw9I40RDNzxCWIZvUAA5mGaAYUGCWacb5aRlkgFqJZ9jKIZhDNIJr1c4UXY4A9qsQJ0cxemEI0g2hm5/LXAUh1QsLRO9IQzcwRlyCa1QMMZBqiGVBglGjGV1/ZoayvDNEsexlEM4hmEM36ucKLMcAeFbKFaGYvTCGaQTSzc/nrAKQ6IeHoHWmIZuaISxDN6gEGMg3RDCgwSjTjK0/uUJYnh2iWvQyiGUQziGb9XOHFGGCP6kBDNLMXphDNIJrZufx1AFKdkHD0jjREM3PEJYhm9QADmYZoBhQYJZqNxESzS71eiGYQzSCaMaAJ0QwrvB5m26My6hDN7IUpRDOIZnYufx2AVCckHL0jDdHMHHEJolk9wECmIZoBBUaJZr6YaHYpdw3RDKIZRDMGNCGaYYXXpEaMp77H9iX8AswhmoniFKIZGUwhmkE0s3L56wCkOiHh6B1piGbmiEsQzeoBBjIN0QwoaFc0+2m1qymZmV5BUiYz+1paO+pYHrI58BeqWp/AnytrnQO9zVpbGdo5+oBjMim1Dl2uTJcrLt3beLPe79I5sXtYrX5Nu/T9/cviP5vtD7NkcUlvaZkw/9lutciejE7H0vNP6YXMXz7s9pnDwepx1YqfqA1mxBstQ2FScrGGsTcdh6xncBregxV72apRVBRf8ttDQ+45QEAMAkkfWiCrefqv4LcdHyl77NfVev/+3o3Nd0S1PZACn+UqBX/ktZR14EFwCxeaRnALdThPfXBTiFN6weJsHyQXJLcRoIHmKjIc7n62bCRBdQGERuhu6IZxFDEDXrYSXo2PpE55qwu55ikvRRVXUN7ChaZR3kIVrdyy4hBQXs72QXlBeRsBGiivItPh7mfLRhKUF0BohPJGccIQ2dnE8kftobwaH0md8laXYctTXooabKC8hQtNo7yFGhi5ZcUloLyc7YPygvI2AjRQXkWmw93Plo0kKC+A0Azl9aMoHDH5oWsr5dX3SOqUt7qISp7yUlRQAeUtXGga5S1ksM4tKx4B5eVsH5QXlLcRoIHyKjId7n62bCRBeQGEZl5siMNJVPz+8vxQdlJejY+kTnmrU6DnKS9F/nNQ3sKFplHeQv7J3LIyIqC8nO2D8oLyNgI0UF5FpsPdz5aNJCgvgNAM5XWiWZx/xfX6UJZSXn2PpE55qxOY5ikvRfZSUN7ChaZR3kL2qNyy4hNQXs72QXlBeRsBGiivItPh7mfLRhKUF0BohPLGYeSFxeQG54eyk/JqfCR5ysvx2RrF12q+YQy3PX4Adq2Q/SybatCazGo5Soy0bW25BLu/zt3vFF/M2f0137FONkjmVZJoOp4aPG8BampOYf154BmuSoNsUWTAxBBjbRrpdlL/GzLPa0eNHlb9GHOGR9nIkOtO74dpXjrkyNF/2+u25EQkUUgxCKXZ6SlVDu0TeSdx++IAklLLFHQY/syZDmXmTAgzEGZIsnaKsxxDcoLKkmqkHIVAw4nQUoFGLLmhFZSy6xKN2JBBpIFIowVY/Rh1s2UaldS0mOqlgw6hhvG6rDXZfDst1bQwDBBrOCDUkljD8/IMZc5niDUQa0jyTYtzHUOyWcuSayTLhljDidBSsUYsLa8VtLLrYo3YkEGsgVijBVj9GHWzxRqVpOqY6qWDDrGGkcHSmjz0nRZrWhgGiDUcEGpJrOGoVuBQViuAWAOxhqRSgjjXMaQOgyy5RpkHiDWcCC0Va8QSyltBK7su1ogNGcQaiDVagNWPUTdbrFEpB4KpXjroEGsYKoE1FVS6LdY0PwwQazgg1JJYw1Fnx6GsswOxBmINSY0fca5jSAUhWXKNAkUQazgRWirWiJVCsYJWdl2sERsyiDUQa7QAqx+jbrZYo1LIClO9dNAh1jC+v7Gm9lenxZoWhgFiDQeEWhJrOCrEOZQV4iDWQKwhqU4n8cm3GbXvZMk1SutBrOFEaHnOGqEiXlbQyq6LNWJDBrEGYo0WYPVj1M0Wa1RKMGKqlw46xBqGSmBN1cpuizXNDwPEGg4ItSTWcNQ2dShrm0KsgVhDUldVnOsYUrVVllyjKCzEGk6Eloo1YuUnraCVXRdrxIYMYg3EGi3A6seomy3WqBQPxlQvHXSINYxet6becqfFmhaGAWINB4QaFGv+vl09VleBSq8gKf40bl2b6Zyi4Q3Sf2xV53zwOHeDOAcwqWCOvDGpl1PkzcnEFuWtFVbyhuz91oS9Pqo9D/m1R3NdxcKqLq02fSz03HzHVpWUdAlZo7pEjSH57JLZeEV8/GaGrXGjkj4O9+SaDNJ/nJNr3J630P4DKbBArpqgRzZIWRMUtDB7GQktHAezwby0VBQ5MVQyJ0MNlQxKkEMle1L0kMCiIEGUtQiKqL+eE0hiBj9kJFFhjoEmGkkTZ+MgDvknmA1EUeMjqVPF6opkeapIUZEMVDF7GU0dr3gcjEtyHzrVi6BUXQwVc1KVvlQMypRqUbEnRRUJLApSRVmLoIr6q0mAKmbwQ0YVFeYYqKKRVDGMZ+OZzz3BbKCKGh9JnSpW10PJU0WKeiigitnLSKhi4M7nk5LMS271IihDFZXMyVBFJYMSVFHJnhRVJLAoSBVlLYIq6s9lDaqYwQ8ZVVSYY6CKRlLFeRiGszn3BLOBKmp8JHWqWJ2NPU8VKbKxgypmL6MpOBdPZkGJv+xVL4JSBVxUzEmVpFMxKFNTSMWeFFUksChIFWUtgirqz6QJqpjBDxlVVJhjoIpGUsUgDoYlH9SwJpgNVFHjI6lTxepcsHmqSJELFlQxexnNu4qT+SD02IvgqHoRlHpXUcWc1LuKKgZl3lVUsSf3rqK6RdF3FSUtgirqz+MFqpjBD927ivJzDFTRSKo4G4WjiP2GB2uC2UAVNT6SOlWszkSXp4oUmehAFbOX0eRvmwaDccki6FcvglL5UFTMSWV4UzEok6JHxZ4UVSSwKEgVZS2CKurPIgKqmMEPGVVUmGOgikZSxTiYz2ZsXsWaYDZQRY2PJE8VOT5nofiKZdI6M0SOYmM5LkcfnJoWJ7T8bcuwV/7WJagqf+NSvFS0eUESytU8GGcHcu1ILWpyjFHkBdHkH2dvDafq5EGNPTfYheKk21FbQG4WbiRRVsCZoothFtAaSNzcX6TwuIUXEBQ3xmuu/uIpgKd98MwP//FyAVcdS8h1Vg3LSgLuE+yflRRc0QABIBsfP0tSKivoMvyZ6RzKzHQQaiDUlCfUjSfDoDR7lKpUI9K6VHJlgfZlsikLNC+XPlnYgGi+ZD4DEG06kf3OSNkmjJ2Y/bEGhBsIN/o8Kgg3QkDrse8N4Qbgka8UHUSjkjfMbZVu7Mk/SirecJNxefmGn++rA7OFUeyNhMPzig1lxlhIOJBwyvOYDkYDr2RRcQqMQSLVrUDrUpltBdqXSWQr0Lxc3lphA6JpavkMQMLpRFZaEyWceBKFUcjdX5BwIOFcht5wxxwSDpACCafv4HHCIAz4+YAFEo49ecFJJRxuMi4v4fDzfQJtsflR7I2Ew5HJ3aHM5A4JBxJOedLIIAhGJSn03AJjkMgrKtC6VBpRgfZlsoYKNC+XJFTYgGhOUD4DkHA6kS3eSAlnFE9K3lpi9RckHEg4l6E33DGHhAOkQMLpOXjSLI8hO0TB5AMWSDj21OsglXC4ybi8hMPP9wm+62t+FHsj4XBUWHEoK6xAwoGEU+rjTwa+l0kFlVtUvAJjEJdwRFqXkXBE2peQcESal5JwxA0ISjicBiDhdKKKi5ESjhPFMTsYxOovSDiQcC5Db7hjDgkHSIGE03PwRKMwjtieMpMPWCDh2FNHi1TC4Sbj8hIOP99XB2YLo9gbCYej8plDWfkMEg4knPJkKcFsPvfZi8qowBgkcuEItC6VC0egfZlcOALNy+XCETYgmguHzwAknE5UVzNRwonC2I+n3P0FCQcSzmXoDXfMIeEAKZBweg6ecBZFscvPByyQcOypb0mbC4eXjMtLOPx8nyAXTvOj2BsJh6MiqUNZkRQSDiSc8jqZ46nvlSwqfoExSJRSFWhdqnKqQPsyhVIFmperiypsQLQMKp8BSDidqHpqooQTR7FXEqVk9RckHEg4l6E33DGHhAOkQMLpO3jCaBqyQxRMPmCBhGNP3WlSCYebjMtLOPx8nwCYzY9i5yUcjhw4FKlvppnT7Sg23dM58ug5zTwmfGS1DkELUnqHoA0ZzUPQhNxSK2VEdLHlNwL9oys1uFdlXHnFSaMFUaPd5SeZp00sabWLmuORmdG9rkl6FzpWWHUiWHAMszPWBKmtczOWEOdtT1k6K5ixhsxYCtHS1inbxHSqADrhwmCC8qV7X+kpSAVkVW3w6og2qxOhkiJsj/l/58kEPYAng/Qfp7NtoA4PVBuOar1mDKfZ2maXQoDh9I7okCPQcH5H9PqyIiIOiDgg4oCIQ7ciDqEbxiWp2BFzADtDzMFAalWo252fs4g6IOrQXZcKcxZxhxYmVH/iDvr3lp7CFJEHSzCK2AMohXYIz8ZBHPK73Yg+ANeIPpgxv9TjD45A/OFSxBnxB8QfEH+gNIL4gwHxhygO3ZD9IR2rqDziD+BniD+0TK7mg9HAY/vfDj+PqtKIMGcRf8CctWfOIv6A+IMNOEX8AfEH0zGK+AMohf7c2PFsPGOX72S53Yg/ANeIP5gxv9TjDzyJls7xB2RcQvwB8Yc7xB+6Gn/woyjMV104L9T50iGIP4CfIf5gBLmaBkEwYmeFJUgAi/gD4g+Ys3bNWcQfEH+wAaeIPyD+YDpGEX8ApdAfQgvDcMYuXMRyuxF/AK4RfzBjfqnHHzyB+EM2OID4A+IPiD8g/tCl+EMYh5NowlyoPcZCjfgD+BniDy2Tq8nA90qKf3n8PKpKI8KcRfwBc9aeOYv4A+IPNuAU8QfEH0zHKOIPoBTaIRzEwTAccLvdiD8A14g/mDG/1OMPI4H4wwjxB8QfEH9A/KGr8QcnmsX5lHjnhTr/VQTiD+BniD8YQa68YDafsz8uHfHzqCqNCHMW8QfMWXvmLOIPiD/YgFPEHxB/MB2jiD+AUuiv/zAKRxE7hMZyuxF/AK4RfzBjfqnHH3yB+IOP+APiD4g/IP7Q0fhDHEZeSaA4z/ARfwA/Q/zBCHIVj6e+x/a/fX4eVaURYc4i/oA5a8+cRfwB8QcbcIr4A+IPpmMU8QdQCv0QDuazkk94WG434g/ANeIPZswv0fhDuNh++Wm127ODDunZu8Np5TjDeJA53U6cIU+W5GhXjnQZFruAeJybZYPDf4VZtl/+sc8NpmaVuCWSzjpftZkMNe0mppL0emyouZEFwGggMoQjJgYcax2zijHPHvvwtHhc0pBalvBlyPSvHUVtaLMPCgEBFBjaUgtqDeEw9nJRoECCPgWngVUBw6hfsMAwahhGWb/49FLesMY/Pr+Qd3Ut4CjDUbbEUfbiyTBgV6yGqwxXGa6yLHCs3Ycdz4199nuXcJYZ49hpZ9n1R/GUnQQE7nLP3OU2sACHuUsDCZfZnoFUdJodTqfZgdMMp9k2p3k+GA08ttPsZIcTTjOcZjjNfdiJfcfxHLdkRYDT3C+neeq5vuvxgwFOc3ed5jawAKe5SwMJp9megVR0ml1Op/lSyx1OM5xmW5zmaRAEoylz3rnZ4YTTDKcZTnMfdmIv8ocOu8Kxy9qJ4TR32Gke+1PHnfODAU5zd53mNrAAp7lLAwmn2Z6BVHSaPU6nOevRwmmG02yF08xTUBdOM5xmOM192Ynd2B2O2O98eaydGE5zh53mUexNxzN+MMBp7q7T3AYW4DR3aSDhNNszkIpOc0mh8xunmaDIOZxmOM0Nf9PMUQUOTjOcZjjNfdmJncFo4o9LVgQ4zf1ymt3hxJsG/GCA09xdp7kNLMBp7tJAwmm2ZyAVneaS6pw3TjNBZU44zXCam3WaeUqXwGmG0wynuS878XTsjQdlKwKc5n45zZHrzwfsWAYTDHCau+s0t4EFOM1dGkg4zfYMpKjTfFgHP70dTCULKdtnPl90d75K3WPOpt82zmMu8OfTXnFTkMhUX5kFTvGS1IW0WadOKOTNYkzcfPNlrXN0MXPSU7dewQnVG6+spEdSjvV2uSI3clpjCqCCKsNa2P30H9Pvzh471gocTiHU0C9G9rCAwhQ8gqV8BtJINaJVsuWEZG14y6PFJ1lBtcptDDY3naqPLEt4KQ6tYUNnBt9X3+HPBxmjmTNpTO0dCrwxtB3AzdSNxZpCafza9uE/Tl7l+0QPJC57CHwomv7jfCAKpX69TCkv9/zlcnHyLjD/rXxt/FYURZHqymJFcYSywBhUksKFPVNJCvW+cq1T6CQi7UsoJSLNQyvpmVYSxk7MTicGtYRvVkMtgVpin1rizL35mJ3RF3qJaQ4s3w5ZGNLCPn8+3KJi0gbmoJnIQa6dT85aAIhu1SSYzOcRzzPZo5vMxkEcRtyPBOXEDOWkpLxcmXJCUWUOyknhwp4pJyKtyygnIu1LKCcizUM56ZdyEk+iMCqrZ3i7CUI5gXIC5UQMb2YqJ+OxM3fY7w0zayFBObmcNlU5KQxpYb05H25ROWkDc1BO5CDXyvbSBkB0KyfRKJgE7ARELIZlg3ISxrPxjP15KOuRoJyYoZyU1BgsU04oSg1COSlcaJxyUigzkCMNXm55l1FOCpX/cq27hdZllBOR9iWUE5HmoZz0TDkZxZOIHT7I18WBciKsnHAvSvZQWygnXVFORtF45Bbft64oiAXl5HLaVOWkMKSFff58uEXlpA3MQTmRg1w7BQdaAIhu5ST0IzfgqTxoj3IyD8Nwxv9IIsoJjUZQUlKxTCOgqKwIjaBwoXEagYgbLK4RiCgQMhqBSPsSGoFI89AIeqYROFEcs4Xy/LuU0AiENQLuRckeEgeNoCsagTd3A7+seK8mOg6N4PzQWjSCwpAW9vnz4RY1gjYwB41ADnKtbC9tAES3RjCfzwdhMZtHOcOyQSMI4mAYsqUc1iPh7Qoz3q4oqatZppxQlNeEclK40DjlpFBaI0ca/NzyLpXRI1/tMtf6qNC6VEYPgfZlMnoINA/lpF/KSRTGfsze1/O1oKCcCCsn3IuSPdQWyklXlBNn7M/G7AgZswgclJPLaVOVk8KQFvb58+E2M3q0gDkoJ3KQayejRwsA0Z7Rww/DiJ0zjcWwbFBOZqNwFLEFLtYjQTkxQzkpKa5appxQ1FiFclK40DjlREQcEFdORHQZGeVEpH0J5USkeSgn/VJO4ij2IjZXyb+JAuVEWDnhXpTsobZQTrqinAT+yB+wCT2zEiCUk8tpU5WTwpAW9vnz4RaVkzYwB+VEDnKtbC9tAES3chIHoRewc6GyGJYNykkczGcztnLCeiQoJ20pJz+tdvsaueRwibpEkk2cComEViKBy2pJqVNj2ESVuzp0DPFAppE7c9mbPTN913xunndZeIYc5b5Jold8AO2+ZulQ89aRtkEw4HEqeUUmUreC3qgkVYXPUflO+CD9x7mduFQlK3V+NX74j/eBXP4HUmGhnIUM00spqxiClhYuBC3tY1U5EFMQUxBTEFNdRkFMNRDT0A3jkpSRtlLTMIhG8ZD/kZolp3W1orLklKJQFMhp4UKQ0z4W7gE5BTkFOQU51WUU5FQDOY3ihJ6yXwFgbSg2kNPYCYMw4H+kZslpXTmOLDmlqMUBclq4EOS0j7URQE5FRjJZF6KJQM4oE8lp4Rly5PQmcxvIKcipklGQUx3k1I+icMS9odhATqNZPAzZAg7zkZolp3V54LPklCIJPMhp4UKQ0z4m5QY5FRnJcTSdewJFT0wkp4VnyJHTm9JDIKcgp0pGQU51hPXjcFKSSYe1oVhBTkdhXJJEgPlIzZLTulS7WXJKkWcX5LRwIchpH/OegpyKuRljdzBjjiTzy2cTyWnhGarzD4CcgpwqGQU51UFOnVRo5N5QbCCn4SyKYpf/kZolp3XZDLPklCKVIchp4UKQ0z6mlgM5FRlJ15uEM3Y8jZnQ2ERyWniGHDm9SSsOcgpyqmQU5FRH8skw8kpKnbE2FBvIafJI05KCdMxHaoCc/n27eqwhpYdL1LmoCy6av5CQi7Immu7czicMIu1y/byXzNHRADOWoTv8Hy8d/uN8bIpMiNQsUjyZn/VdaFrSXu6eKoxVWU+dKH9AwBYMyzVrcE/pTro6GaT/OGcJRX5S3URR2wOp0ETOpE7ppZRJncAbCxeCN/aFN0on0LCdOQaT+TxiZ9EGdzS4E61lj64/iqc8Mw38sZW+0s0gZ+MgDvmzL9nAITU+EgGLrMu+lGWRFNmXwCILF4JF9oVFSme6sJ1FRqNgEoy5Hxws0pBOtJZFTj3Xd9mMm5mtq88sso2+0s0iw3g2nrG/HmXNFRtYpMZHImCRdWmSsiySIk0SWGThQrDIvrBI6ZQUtrPI0I/cgP1iK+vBwSIN6URrWeTYnzouT1+BRbbSV7pZ5DwMwxn/XLGBRWp8JAIWWZfPKMsiKfIZgUUWLgSL7A2LlM0dYTuLnM/ng5JXv1kPDhZpSCdayyJHsTcds1MMMJOz9plFttFXullkEAfDks9nWHPFBhap8ZEIWGRd4qEsi6RIPAQWWbgQLLIvLFI6yYPtLDLww7AkmxzrwcEiDelEa1mkO5x4U/a7I8xcAH1mkW30lfb3IkfhKGKXeGDNFRtYpMZHImCRdRmCsiySIkMQWGThQrDIvrBI6WwMtrPIOAi9gP3qFevBwSIN6URrWWTk+nORdKd9ZpFt9JVuFhkH89mMTblYc8UGFqnxkTIs8vJ/k538/wBQSwMEFAAAAAgAeCRrXKM/Rl+/AwAA5wkAABEAAAB3b3JkL3NldHRpbmdzLnhtbLVW3XLaOBS+36dguOFmCbZxTOMp6SSw3k0mbDN1+gCyfQBt9DeSDKFP3yPbismWZpjt7BXy+c6/vnPEx08vnA12oA2VYj4KL4LRAEQpKyo289HXp2z8YTQwloiKMClgPjqAGX26/u3jPjVgLWqZAXoQJuXlfLi1VqWTiSm3wIm5kAoEgmupObH4qTcTTvRzrcal5IpYWlBG7WESBUEy7NzI+bDWIu1cjDkttTRybZ1JKtdrWkL34y30OXFbk6Usaw7CNhEnGhjmIIXZUmW8N/5fvSG49U527xWx48zr7cPgjHL3UlevFuek5wyUliUYgxfEmU+Qij5w/IOj19gXGLsrsXGF5mHQnPrMDTsnkRZ6oIUm+nCcBS/Tu42QmhQM5kPMZniNjPomJR/s0x1B5wUYm1E7nDgAi5Hr3BILCBsFjDl6DksGBJ3t040mHJnlJY1NBWtSM/tEitxK5d3OoqCFyy3RpLSgc0VK9LaQwmrJvF4l/5Z2gSzV2MTWwpAdPGrYUdg/0tLWGlpHDZXdqTaQ/fFADrK2R0jejgk6FoRjsW+ov5IVuAJqTc+/j6FPEtv2TiCJU61pBU+uybk9MMiwxpx+gxtR3dfGUvTYDMAvZPBeAiBc5M9Ii6eDggyI65n5n4I1F5YxqlZUa6nvRIWT+avBJsfXiyuyMv7wRUrrVYPgNp7Nph2xHNojwTROwuQkkgTJdHEKCS+DWXx7ComukunV8hQyjZLs6mQGNzfh8sNJm59nvbgNkiQ+hWSL5Gqadb3pOsJTt/setT85mg14a7EgvNCUDFZuO06cRqGfb6nweAG4L+AYyevCg+NxCxhOGMtwXD0QtPKKGrWEdXNmK6I3vd9OQ5+U4mq4f/VVIk9A/6llrVp0r4lq6eNVwjjuLKmwD5R7uamL3FsJ3HBHUC2qzzvd9Klvzz61SL9mDB9Iw91GF8T4a+6IB8TYG0PJfPgPGd8/dnRnOneshRVRqmV8sQnnQ0Y3Wxs6M4tfFb6rzUexiTosarCoxZoPUrpiUbs79LLIy470pl427WWxl8W97NLLLntZ4mWJk21x/DWu7GecQ3908rVkTO6h+qvHfxB1y9xN901tpV/J3QY27WbeEgXLdt8jH2Ur6B4AM9il8GKxzRU+JwOjaMXJC15qEM2c806bNXv7ja7DnLJ666Eilvj98Ma4mYl/5eLeoZIif/MDL/rn5aIti1GDi0zhS2Sl9tjvDRbGWHR5h6OHp0YexUESBUn4CrdB7jjZwFLRXnEaBN2A+r9o198BUEsDBBQAAAAIAHgka1zoWuVTAAEAALYBAAAUAAAAd29yZC93ZWJTZXR0aW5ncy54bWyN0MFqwzAMANB7vsLkklPjZIwxQpIyGB27lEG2D3AcJTG1LWO5zfr3M1k2GLv0JiHpIanefxrNLuBJoW2yMi8yBlbioOzUZB/vh91jxigIOwiNFprsCpTt26ReqgX6DkKIjcQiYqkysknnEFzFOckZjKAcHdhYHNEbEWLqJ26EP53dTqJxIqheaRWu/K4oHtKN8bcoOI5KwjPKswEb1nnuQUcRLc3K0Y+23KIt6AfnUQJRvMfob88IZX+Z8v4fZJT0SDiGPB6zbbRScbws1sjolBlZvU4Wveg1NGmE0jZhLH5QaI3L2/GFb/mARwyduMATdXENDQelIRZr/ufbbfIFUEsDBBQAAAAIAHgka1z7OaBzYwIAAPsKAAASAAAAd29yZC9mb250VGFibGUueG1s3ZbBbtowHMbvfYool5xKbJO1FBEqxoa0yw4bewATHLAW25HtQLnS+847bI8w7bBJu/RtkHrtK8wkAYIIGXRDSAMhOf/P+WL/9P0dWrd3LLImRCoquO/AGnAswgMxpHzkOx/6vcuGYymN+RBHghPfmRHl3LYvWtNmKLhWlrmdqyYLfHusddx0XRWMCcOqJmLCjRgKybA2l3LkMiw/JvFlIFiMNR3QiOqZiwC4snMbeYiLCEMakFciSBjhOr3flSQyjoKrMY3Vym16iNtUyGEsRUCUMltmUebHMOVrG+jtGDEaSKFEqGtmM/mKUitzOwTpiEW2xYLmmxEXEg8i4tvGyG5fWFbOzpo2OWam/n7GBiJKpVSMMReKQKNPcOTboORju+vZwRhLRfR6NipoIWY0mq0knGhREGOqg/FKm2BJl6ss6IqOjJqoAdiswc4q0LfhdgXtzKlvV4LUp7FdgYU56YNbbsamDFOfMqKst2RqvRMM8/28kPlegTp4ATzzQ2bkVfACp+D12uwIdXq9Da+uqVw3PLjD66aKV3oJM59jeXUxG5hFVnFa8sk4LXmh83ACqMjJW1a8deXAXGWcbp7F6enh29PDD+vx86fHL1//URc29tOSaXg3Khe6LxPSn8VkD8OQ3pFhdWPCDUDQANdljQn/BBA9tzG7OKImaVVB66WNiNLInSdosCxonW5J0A5oyL8K2mL+czH/tbi/X8y/nz5uTAyJ/M/yJhJJiazKGzB5O5DdafKWP7Ze4FRgcOTBlvM+llPHrLDibwUCL82x7+V9ic51/Je+Juunek2uRqp98RtQSwMEFAAAAAgAeCRrXJRBIrjGBgAAuyoAABUAAAB3b3JkL3RoZW1lL3RoZW1lMS54bWztWk1v2zYYvvdXELrk1PrbdYq6RezY7damDRK3Q4+0RFtsKFEg6SS+De1xwIBh3bDDCuy2w7CtQAvs0v2abh22DuhfGCnZiihRcubFTdolB8ci+Tx8v19S8NXrhx4B+4hxTP32WuVSeQ0g36YO9sfttXuD/sXWGuAC+g4k1EfttSnia9evXbgKrwgXeQhIuM+vwLblChFcKZW4LYchv0QD5Mu5EWUeFPKRjUsOgweS1iOlarncLHkQ+xbwoYfa1t3RCNsIDBSlde0CAHP+HpEfvuBqLBy1Cdu1w52TSCuaD1c4e5X5U/jMp7xLGNiHpG3J/R16MECHwgIEciEn2lY5/LNKMUdJI5EURCyiTND1wz+dLkEQSljV6dh4GPNV+vX1y5tpaaqaNAXwXq/X7VXSuyfh0LalRSv5FPV+q9JJSZACxTQFknTLjXLdSJOVppZPs97pdBrrJppahqaeT9MqN+sbVRNNPUPTKLBNZ6PbbZpoGhmaZj5N//J6s26kaSZoXIL9vXwSFbXpQNMgEjCi5GYxS0uytFLRr6PUSJx2cSKOqC8WZKIHH1LWl+u03QkU2AdiGqARtCWuCwkeMnwkQbgKwcSS1JzN8+eUWIDbDAeibX0cQFlijta+ffnj25fPwatHL149+uXV48evHv1cBL8J/XES/ub7L/5++in46/l3b558tQDIk8Dff/rst1+/XIAQScTrr5/98eLZ628+//OHJ0W4DQaHSdwAe4iDO+gA7FBPKl+0JRqyJaEDF+IkdMMfc+hDBS6C9YSrwe5MIYFFgA7SHXCfyWJbiLgxeagpteuyiUjHloa45XoaYotS0qGs2AC3lBhJ20388QK52CQJ2IFwv1CsbiqEepNA5hou3KTrIk2VbSKjCo6RjwRQc3QPoSL8A4w1/2xhm1FORwI8wKADcbEhB3gozOib2JOOnhbKLkNKs+jWfdChpHDDTbSvQ2S6QlK4CSKaF27AiYBesVbQI0nIbSjcQkV2p8zWHMeFDKYxIhT0HMR5Ifgum2oq3ZK1cUFkbZGpp0OYwHuFkNuQ0iRkk+51XegFxXph302CPuJ7MlMg2KaiWD6q57B6lo6F/uKIuo+RWLJC3cNj1xyMambCCnMVUb2GTMkIosR2qiFmepvqd9g/Vr/zZLtL22yV/U62kdffPv3AOt2GtGFhsqf720JAuqt1KXPwh9HUNuHE30Yygc972nlPO+9pZ6inLaxKq+9keteK7n/zu93Rdc9bdNsbYUJ2xZSg21xvgFyaxunL2aPRaDzkiy+igSu/atqUjFiJHDMYDgJGxSdYuLsuDKRMFSu1w5hrssSjIKBc3p8tfSpfqPS66P0UlpYOFzX090c6HxRb1InW1crmhaGi831T4paUvLkq1NTWJ6VG7fJpqVGJGE9Ij0rjmHrk+O1f6RGNpMJMnfrkmU+WSClNsxppJ7MSEuSoME0F+Tycz3KMV3KcHhG60EHHWZewfqV2tqOoMKmX0Pe0oq28KNrCgm+o3YrWNxZ04oODtrXeqDYsYMOgbY3kHUd+9QK5H1etEZKx37ZswdLRauwFx/eRbvt1c6KnA61sWpZr9pyuE9IGjItNyN2IOFyVti7xDaaqNurKJau1VWnVWtRalfdVi+jJEOFoNEK2MEZ5Yiq1dTRjKrt0IhDbdZ0DMCQTtgOldepROjqYywNZdf7AZIGpzzJVL/DmApZ+72+oc+FCSAIXzgpOK7/eRHTZjIjlT3vBoPLRcMpGq7Jd7R3aLqeynNvu9G03qx3IRzUnYwhbXk4YBKo4tC3KhEtluwtcbPeZvNOYVJRWALKYKQMAQv3wP0P7qcY5lyfiz2xL5FVM7OAxYFg2YeEyhLbFzN7/btdK1XigCAvYbJNMhczaQlkoMJhniPYRGahi3lRusoA7b07ZuqvhcwI2NazX1uG4/7+9Etbf5alQU6F+kofgetFVKnEQWz8tbU/izJ9QpHpMt1UbBUXuvx7mAyhcoD7keQozmyAro746rw/ojsw7EF9VgKwmF1uz0h4PDqWNWlmt1N5qi/fvImpQxuiis/mWIhFrOfffbKydhCIriLWGIdQM+X28SFNjpn4RXk69xMtINZD5ZZg6AQ0fSgk30QhOSOLnYjyQQ4mexINtVko8D6kz1UcIj3pZcoxnDmnE30EjgJ1DQyKkomH206ns5WTnSLLY0DFrbTnWGYfhQBkzV5djjll0meWpKmYO3yQvYCcGmSOOZCgkDB6dRWIvhrZfuU+XtNECn5ZX5tMlY/CEfCoOl/Bp7MXw/J/JXqXjoWCwO//hmSwJco84/a9d+AdQSwMEFAAAAAgAeCRrXJ6AOtenAAAABgEAABMAAABjdXN0b21YbWwvaXRlbTEueG1srYyxCsIwFAD3fkXJksmmOogU01IQJxGhCq5J+toGkrySpGL/3oi/4Hh3cMfmbU3+Ah80Ok63RUlzcAp77UZOH/fz5kDzEIXrhUEHnK4QaFNnR1l1uHgFIU8DFyrJyRTjXDEW1ARWhAJncKkN6K2ICf3IcBi0ghOqxYKLbFeWeya1NBpHL+ZpJb/Zf1YdGFAR+i6uBjhh7a0tnt0lha+4CptkcoTV2QdQSwMEFAAAAAgAeCRrXD7K5dW9AAAAJwEAAB4AAABjdXN0b21YbWwvX3JlbHMvaXRlbTEueG1sLnJlbHONz7FqwzAQBuC9TyG0aKplZyihWPYSAtlCcCGrkM+2iKUTuktI3r6iUwMZMt4d//dzbX8Pq7hBJo/RqKaqlYDocPRxNupn2H9ulSC2cbQrRjDqAaT67qM9wWq5ZGjxiURBIhm5MKdvrcktECxVmCCWy4Q5WC5jnnWy7mJn0Ju6/tL5vyG7J1McRiPzYWykGB4J3rFxmryDHbprgMgvKrS7EmM4h/WYsTSKweYZ2EjPEP5WTVVMqbtWP/3X/QJQSwMEFAAAAAgAeCRrXLW7TE3hAAAAYgEAABgAAABjdXN0b21YbWwvaXRlbVByb3BzMS54bWydkLFugzAURXe+wvLiyTGgBGgUiEgAKWvVSl0deIAlbCPbRI2q/ntNOjVjx3eudO7VOxw/5YRuYKzQKifRJiQIVKs7oYacvL81NCPIOq46PmkFObmDJcciOHR233HHrdMGLg4k8h7lmc3x6Ny8Z8y2I0huN3oG5cNeG8mdP83AdN+LFirdLhKUY3EYJqxdvEt+yAkj7xZeealy/FU3cZplUULrc9LQMtnu6EuYVjRt4l1Zn09RtS2/cREgtE767XyF3q7kia3exYj/DryK6yT0YPg83jF7NLKnygf485Yi+AFQSwMEFAAAAAgAeCRrXJDQh4lrAwAAiRUAABIAAAB3b3JkL251bWJlcmluZy54bWzNWN1u4jgYvd+nQJFGXLWJkzQENLSiQFZdjUYjtfMAJhiw6p/IMTDc7kvtY80rrJ0/qIozTBJ2y40Tf985/nxO/AX4/PCDkt4OiRRzNu6DW6ffQyzmS8zW4/73l+gm7PdSCdkSEs7QuH9Aaf/h/o/P+xHb0gUSKq+nKFg62ifx2NpImYxsO403iML0luJY8JSv5G3Mqc1XKxwje8/F0nYd4GRXieAxSlPFM4VsB1OroKP8MjYK4/LSdZxQ3WNWcbyviCeIqeCKCwqluhVrhRCv2+RGcSZQ4gUmWB40V1DR7MbWVrBRwXFT1aExI1XAaEdJmczrcvNCi6FEiEuKzCEzHm8pYjIrzxaIqII5Szc4OerWlE0FNyVJ7YZPNrtPgN/O9JmAezUcCS8pf5mDKMkrr2cEzgWOaIoKcUkJb9csKzl9+PbNpDkVd91O2z8F3yZHNtyO7Ym9VlyqE/wOV+HR6dbSdsU8b2CiDhCNR09rxgVcEFWRUrynn0jrXrUnuEilgLH8uqW9N3dPy7HlZCksxUsV20EytqLsM5hato7QLZH4C9oh8nJIUJmjFyYom87TJE1IGZx6wJlPfTePkJ0OYDWUi6kmKmSZDPIs1UIjWk0uUYwpJBXBC/pRxT6B22r+r7icJWgl8+nkm8gKUvssxjJHrWGp64QrxUHoODrfPmZipiXQREVY3W0gW+v+b3lBmZ7x29ny2Xii5y/FBiaxZ43FnvtOOHRc/0OL7fu1Yutw92K7JrHnjcWOHoEbDL1JR2Inz/JAqpW/4FSXrr5JeNf0wglrvdDh7r3wTF5Ejb3wQt8HwV1XXcbkhXtFLwZunRU62r0TvsGJEDR2AgzAZOpNWrSgxZYQJM8q/fPvf/7/DrQfiWKIOJOpVjWNsfoW8XygC04y6ERp+mYCM6mfsRVUihZkooVxdybj3ObtzJtPotl82o1x70/QYxY938068rVdN/sIvgYmX73mrXEG5lE06+hAmnw93xm78bVVZ/wIrg5MroaNXZ05k8B9zPvYFV94V3zfHX0656qOdv++C01GDBsb4Q4HAVBeXPd4XfF0tfLhPzpdLDOTnf5ueuNsua+woGNnYK4ZFtTAPDPsrgb27sf2EebXwO7MsEENLDDDvBrYwAxza2ChGQZqYEMzzDmF2Sf/od7/C1BLAwQUAAAACAB4JGtcosjWZ70FAACEIAAAFwAAAGRvY1Byb3BzL3RodW1ibmFpbC5qcGVn7VZrcBNVFD67ezcpbc0QKC0UB8K7MsCkLUIrAjZp2qaUNqQtr3GGSZNNE5omYXfTlk6dkfoA9Yc8fP+xFFR0nHFQ0YI6UkVARwcQCxQYxiJq8TU8FF8D8dzdpAlQhJFfzuzd2f2+nPPdc885e+duoseiX8PQ8hJ7CTAMA2V4QfS0vstuta5wOKtK7BU2dADot7nC4QBrAmgMyqKz1GJaumy5Sd8LLIyCNMiGNJdbChc5HBWAg2rhunHpCDAUD08f3P+vI80jSG4AJgV5yCO5G5G3APABd1iUAXRn0F7QLIeR6+9EniFigsjNlNervJjyOpUvVTQ1TitymovB7XN5kLchn1aXZK9P4moOysgoFYKC6HebaC8cYsjrDwhJ6d7EfYujMRCJrzcG73SpoXoBYg6t3SeWOWO8w+2yVSOfiHx/WLZQ+2TkP0UaaouQTwVgh3nFklpVz97b6qtZgjwTuccv22ti9tZgXWWVOpftbAgtcMY0+92SFXsG45Gf8gn2CjUfDjxCsY32C/kYX6QsFp8rl5qqbfE4rT5rpRqHE1e6yh3Is5GvE0POKjVnrlMIlDrV+NzesOyI5cD1BwOVFWpMYhAkpUbFLvtqytS5ZJaML1GdS5Z7/SX2mL4tHFD2IuZGtooRZ21Mc9Al2krVOOSCEKyNxeRHelzFtLczkM+DxYwLBAhBHT7dEITLYAInlIIFMQwierzghwBaBPQKaPEzd0AD2gbXORSNyhOKemV2P52NqwyuUVc4G9OESBYxk3y855AKMpcUkEIwkfnkPjKPFKO1kMwZmOtIWp+udXYgziqIYFSqWwyW9dmRnMR67eIKv/vAk+eumh26Lmchnk9yB0DCDsSV05Pr39f2/shEjB7Sdf/h9H1tUHWz/vJn+H6+B5+9/MmEgj/Bn8SrF4owt4CSUSPefiUPKSmD5Bq68ZbBhc8+1IWSdFet6A2uz054aCeEtZWXKqF9WsJqPmr+2dxj3mzeav7xmi4P2iVuE7eD+4Dbye3iPgcTt5vr5j7k9nJvcO8lvasb74+Bd6/UG6+WegbrtQABg8Uw2jDBUGwYa5hkqEjEM2QZcg1lhinoGT3w3pLXS67FD8vwGe/q4Gupulr0+qFZqUBSOhyE1dfs/9hsMobkEvs1u7aA7uW4QmfTFeuKwKSbqivU5erKKY/np5uCvkJ82q7ade4bVCAkqZLrnK7sOrpX6ewmxSeBIAstMj1oraHwatFf75NNeWbzbFMRfqoEkz3onjHN5AoETIpLMomCJIhNgmcG0O+gekRfdCrfNybzQMImLwSY+wueWQcTtuURgNclgKyZCVsOnokjXgTomuWOiE2xM59hvgCQvPl56q90C55Np6LRi3he6TcCXN4Qjf7dGY1e3oLxTwLsDkT7QLa1+L0ACxfSUx9SgDDZwNPZeM9jRg/wEiYHD3DKWYC1fiAxe2Vs7bLYbxXZDjauYJ7o4OKcVaTRE2Cl/x5ua9AgtxuDie4GYwqLKXKMEVgjwxmZ6B4Yi7nyqiD+YWVYjvA6fcqQ1DQU7BgKLMNxLOF4nmBpzAPoB2Lkh43LLdINX+TSj1+Vkbdmw+aUCZbt3SOch85NzK8T24ekZmaNHJU9afKUnLumzrx71uyCwnusxbaS0jJ7eXVN7eIl+HrdHsFb7/OvlORIU3PL6taHHn7k0bXrHnt846annn7m2eeef6Fzy9aXXn5l26uvvfnW2zveebdr566PPt7zyd59+z/97MvDX/UcOXqs93jf6W/OfPvd9/1nfzh/4eKvv136/Y8//6J1McANlD5oXdgEhiWEI3paF8M2U4GR8ONydcOKFuldq4aPz1uTkmHZsHl795AJ+c5zI+rEQ6mZE2f2TTpPS1Mqu7XC2v9TZQOFJeo6DukcbjgjZ4T5cOVKDnSwD6aCBhpooIEGGmiggQYaaKCBBhpooIEGGmiggQb/M4j2wj9QSwECFAMUAAAACAB4JGtcrVKlkZUBAADKBgAAEwAAAAAAAAAAAAAAgAEAAAAAW0NvbnRlbnRfVHlwZXNdLnhtbFBLAQIUAxQAAAAIAHgka1x5JktA+AAAAN4CAAALAAAAAAAAAAAAAACAAcYBAABfcmVscy8ucmVsc1BLAQIUAxQAAAAIAHgka1yIhgtTaQEAANECAAARAAAAAAAAAAAAAACAAecCAABkb2NQcm9wcy9jb3JlLnhtbFBLAQIUAxQAAAAIAHgka1z029sX6wEAAGwEAAAQAAAAAAAAAAAAAACAAX8EAABkb2NQcm9wcy9hcHAueG1sUEsBAhQDFAAAAAgAeCRrXA5skMb3AQAADAYAABEAAAAAAAAAAAAAAIABmAYAAHdvcmQvZG9jdW1lbnQueG1sUEsBAhQDFAAAAAgAeCRrXG6AGxIyAQAAywQAABwAAAAAAAAAAAAAAIABvggAAHdvcmQvX3JlbHMvZG9jdW1lbnQueG1sLnJlbHNQSwECFAMUAAAACAB4JGtcB9SvmXMvAAASVQUADwAAAAAAAAAAAAAAgAEqCgAAd29yZC9zdHlsZXMueG1sUEsBAhQDFAAAAAgAeCRrXGB5gtM5NQAAc68GABoAAAAAAAAAAAAAAIAByjkAAHdvcmQvc3R5bGVzV2l0aEVmZmVjdHMueG1sUEsBAhQDFAAAAAgAeCRrXKM/Rl+/AwAA5wkAABEAAAAAAAAAAAAAAIABO28AAHdvcmQvc2V0dGluZ3MueG1sUEsBAhQDFAAAAAgAeCRrXOha5VMAAQAAtgEAABQAAAAAAAAAAAAAAIABKXMAAHdvcmQvd2ViU2V0dGluZ3MueG1sUEsBAhQDFAAAAAgAeCRrXPs5oHNjAgAA+woAABIAAAAAAAAAAAAAAIABW3QAAHdvcmQvZm9udFRhYmxlLnhtbFBLAQIUAxQAAAAIAHgka1yUQSK4xgYAALsqAAAVAAAAAAAAAAAAAACAAe52AAB3b3JkL3RoZW1lL3RoZW1lMS54bWxQSwECFAMUAAAACAB4JGtcnoA616cAAAAGAQAAEwAAAAAAAAAAAAAAgAHnfQAAY3VzdG9tWG1sL2l0ZW0xLnhtbFBLAQIUAxQAAAAIAHgka1w+yuXVvQAAACcBAAAeAAAAAAAAAAAAAACAAb9+AABjdXN0b21YbWwvX3JlbHMvaXRlbTEueG1sLnJlbHNQSwECFAMUAAAACAB4JGtctbtMTeEAAABiAQAAGAAAAAAAAAAAAAAAgAG4fwAAY3VzdG9tWG1sL2l0ZW1Qcm9wczEueG1sUEsBAhQDFAAAAAgAeCRrXJDQh4lrAwAAiRUAABIAAAAAAAAAAAAAAIABz4AAAHdvcmQvbnVtYmVyaW5nLnhtbFBLAQIUAxQAAAAIAHgka1yiyNZnvQUAAIQgAAAXAAAAAAAAAAAAAACAAWqEAABkb2NQcm9wcy90aHVtYm5haWwuanBlZ1BLBQYAAAAAEQARAGEEAABcigAAAAA="
                    doc = Document(io.BytesIO(_b64.b64decode(_BLANK_B64)))
                    for p in list(doc.paragraphs):
                        p._element.getparent().remove(p._element)

                    section = doc.sections[0]
                    section.page_width  = Inches(8.5)
                    section.page_height = Inches(11)
                    section.left_margin = section.right_margin = Inches(1)
                    section.top_margin  = section.bottom_margin = Inches(1)

                    def _hd(doc, text, level=1):
                        p = doc.add_heading(text, level=level)
                        if p.runs:
                            p.runs[0].font.color.rgb = RGBColor(0x2E,0x5C,0x8A) if level==1 else RGBColor(0x1A,0x3F,0x6B)
                        return p

                    def _tbl(doc, headers, rows):
                        t = doc.add_table(rows=1+len(rows), cols=len(headers))
                        t.style = "Table Grid"
                        for i,h in enumerate(headers):
                            c = t.rows[0].cells[i]
                            c.text = str(h)
                            if c.paragraphs[0].runs:
                                run = c.paragraphs[0].runs[0]
                                run.bold = True; run.font.size = Pt(10)
                                run.font.color.rgb = RGBColor(255,255,255)
                            tcPr = c._tc.get_or_add_tcPr()
                            shd = OxmlElement("w:shd")
                            shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"2E5C8A")
                            tcPr.append(shd)
                        for ri,row in enumerate(rows):
                            for i,val in enumerate(row):
                                c = t.rows[ri+1].cells[i]
                                c.text = str(val) if val is not None else "—"
                                if c.paragraphs[0].runs:
                                    c.paragraphs[0].runs[0].font.size = Pt(9)
                                if ri%2==1:
                                    tcPr = c._tc.get_or_add_tcPr()
                                    shd = OxmlElement("w:shd")
                                    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"EEF3FA")
                                    tcPr.append(shd)

                    # 封面
                    doc.add_paragraph()
                    tp = doc.add_paragraph()
                    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    tr2 = tp.add_run(report_title)
                    tr2.bold=True; tr2.font.size=Pt(22); tr2.font.color.rgb=RGBColor(0x2E,0x5C,0x8A)
                    sp = doc.add_paragraph()
                    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    sp.add_run(f"Analyst: {analyst_name or chr(8212)}   \u00b7   Date: {today}").font.color.rgb=RGBColor(0x66,0x66,0x66)
                    doc.add_paragraph()

                    # 执行摘要
                    if exec_summary:
                        _hd(doc, "Executive Summary", 1)
                        for line in exec_summary.split("\n"):
                            if line.strip(): doc.add_paragraph(line.strip())
                        doc.add_paragraph()

                    # Section 1
                    if has_data:
                        _hd(doc, "1. Historical Financial Overview", 1)
                        df_r = st.session_state.fetched_df
                        cols = list(df_r.columns)
                        trows = [[str(round(v,2)) if isinstance(v,float) else str(v) for v in row] for _,row in df_r.iterrows()]
                        _tbl(doc, cols, trows)
                        doc.add_paragraph()

                    # Section 2
                    if has_forecast:
                        _hd(doc, "2. Income Statement Forecast (3-Year)", 1)
                        fp = st.session_state.forecast_params
                        doc.add_paragraph(f"GM {fp['gross_margin']}%  |  Opex {fp['expense_ratio']}%  |  Tax {fp['tax_rate']}%")
                        fd = st.session_state.forecast_df[["年份","营业收入","净利润","净利率%"]]
                        _tbl(doc, list(fd.columns), fd.values.tolist())
                        doc.add_paragraph()

                    # Section 3
                    if has_3stmt:
                        _hd(doc, "3. Three-Statement Model", 1)
                        for label, key in [("Income Statement","income_rows"),("Cash Flow","cashflow_rows"),("Balance Sheet","balance_rows")]:
                            rlist = st.session_state[key]
                            if rlist:
                                _hd(doc, label, 2)
                                hdrs = list(rlist[0].keys())
                                _tbl(doc, hdrs, [[r[k] for k in hdrs] for r in rlist])
                                doc.add_paragraph()

                    # Section 4
                    if has_dcf:
                        _hd(doc, "4. DCF Valuation", 1)
                        d=st.session_state.dcf_result; dp2=st.session_state.dcf_params
                        _tbl(doc, ["Metric","Value"], [
                            ["Base Case Price/Share", f"${d['price_per_share']}"],
                            ["Enterprise Value",      f"{d['ev']} bn"],
                            ["PV of FCFs",            f"{d['pv_fcf']} bn"],
                            ["PV Terminal Value",     f"{d['pv_terminal']} bn ({d['terminal_pct']}% of EV)"],
                            ["WACC",                  f"{dp2['wacc']}%"],
                            ["Terminal Growth",       f"{dp2['terminal_growth']}%"],
                            ["Valuation Range",       f"${dp2.get('price_low',0):.1f}—${dp2.get('price_high',0):.1f}"],
                        ])
                        doc.add_paragraph()

                    # Section 5
                    if has_pepb:
                        _hd(doc, "5. Relative Valuation (PE & PB)", 1)
                        pr=st.session_state.pepb_result
                        pe_bear=pr.get("pe_price_bear",pr.get("pe_bear",0))
                        pe_base=pr.get("pe_price_base",pr.get("pe_base",0))
                        pe_bull=pr.get("pe_price_bull",pr.get("pe_bull",0))
                        pb_bear=pr.get("pb_price_bear",pr.get("pb_bear",0))
                        pb_base=pr.get("pb_price_base",pr.get("pb_base",0))
                        pb_bull=pr.get("pb_price_bull",pr.get("pb_bull",0))
                        _tbl(doc, ["Method","Bear","Base","Bull"], [
                            ["PE", f"${pe_bear}", f"${pe_base}", f"${pe_bull}"],
                            ["PB", f"${pb_bear}", f"${pb_base}", f"${pb_bull}"],
                        ])
                        c_low=pr.get("consensus_low", round((pe_base+pb_base)/2*0.9,2))
                        c_high=pr.get("consensus_high",round((pe_base+pb_base)/2*1.1,2))
                        c_mid=pr.get("consensus_mid", round((pe_base+pb_base)/2,2))
                        cp=doc.add_paragraph(f"Consensus: ${c_low}—${c_high}  \u00b7  Avg: ${c_mid}")
                        if cp.runs: cp.runs[0].bold=True; cp.runs[0].font.color.rgb=RGBColor(0x2E,0x5C,0x8A)
                        doc.add_paragraph()

                    # 免责声明
                    disc=doc.add_paragraph()
                    dr=disc.add_run("Disclaimer: This report is for informational purposes only and does not constitute investment advice.")
                    dr.italic=True; dr.font.size=Pt(9); dr.font.color.rgb=RGBColor(0x88,0x88,0x88)

                    # 输出
                    buf=io.BytesIO()
                    doc.save(buf)
                    docx_bytes=buf.getvalue()
                    st.success("\u2705 Report generated!")
                    st.download_button(
                        label="\u2b07\ufe0f Download Report (.docx)",
                        data=docx_bytes,
                        file_name=f"{cname_r.replace(' ','_')}_Research_Report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )


                except Exception as e:
                    st.error(f"Error: {e}")

# ============================================================
# TAB 6: 地雷扫描
# ============================================================
with tab6:
    st.header("🔍 地雷扫描")
    st.write("基于已获取的财务数据，自动检测财务风险信号。请先在 **Data** 标签页加载数据。")

    if st.session_state.fetched_df is None:
        st.info("👈 请先在 Data 标签页输入股票代码获取数据。")
    else:
        df_scan  = st.session_state.fetched_df
        cname_sc = st.session_state.fetched_company or "该公司"
        is_us_sc = st.session_state.is_us_stock

        rev_col  = "营业收入(亿美元)" if is_us_sc else "营业收入(亿元)"
        ni_col   = "净利润(亿美元)"   if is_us_sc else "净利润(亿元)"
        nm_col   = "净利率%"
        dr_col   = "资产负债率%"
        roe_col  = "ROE%"

        st.subheader(f"{cname_sc} — 财务健康扫描")

        def col_vals(col):
            if col not in df_scan.columns:
                return []
            return [float(v) for v in df_scan[col].tolist()
                    if v is not None and str(v) not in ('nan','None')]

        def latest_val(col):
            vals = col_vals(col)
            return vals[-1] if vals else None

        def prev_val(col):
            vals = col_vals(col)
            return vals[-2] if len(vals) >= 2 else None

        # ── Piotroski F-Score ─────────────────────────────────
        st.markdown("---")
        st.markdown("### 📊 Piotroski F-Score（企业质量评分）")
        st.caption("满分9分。7-9分：优质企业 ✅ | 4-6分：普通 ⚠️ | 0-3分：高风险 🚨")

        f_scores = {}
        f_details = {}

        ni_now  = latest_val(ni_col);  ni_prev  = prev_val(ni_col)
        roe_now = latest_val(roe_col); dr_now   = latest_val(dr_col); dr_prev = prev_val(dr_col)
        nm_now  = latest_val(nm_col);  nm_prev  = prev_val(nm_col)
        rev_now = latest_val(rev_col); rev_prev = prev_val(rev_col)
        ni_vals  = col_vals(ni_col);   rev_vals = col_vals(rev_col); roe_vals = col_vals(roe_col)

        f_scores["净利润为正"]     = 1 if (ni_now  and ni_now  > 0) else 0
        f_details["净利润为正"]    = f"最新净利润：{ni_now}"
        f_scores["ROE为正"]        = 1 if (roe_now and roe_now > 0) else 0
        f_details["ROE为正"]       = f"最新ROE：{roe_now}%"
        f_scores["净利润同比增长"] = 1 if (ni_now and ni_prev and ni_now > ni_prev) else 0
        f_details["净利润同比增长"]= f"{ni_prev} → {ni_now}"
        f_scores["资产负债率下降"] = 1 if (dr_now and dr_prev and dr_now < dr_prev) else 0
        f_details["资产负债率下降"]= f"{dr_prev}% → {dr_now}%"
        f_scores["净利率改善"]     = 1 if (nm_now and nm_prev and nm_now > nm_prev) else 0
        f_details["净利率改善"]    = f"{nm_prev}% → {nm_now}%"
        f_scores["收入同比增长"]   = 1 if (rev_now and rev_prev and rev_now > rev_prev) else 0
        f_details["收入同比增长"]  = f"{rev_prev} → {rev_now}"

        if len(rev_vals) >= 3:
            f_scores["收入增速加快"] = 1 if (rev_vals[-1]-rev_vals[-2]) > (rev_vals[-2]-rev_vals[-3]) else 0
            f_details["收入增速加快"]= "近两年收入增量对比"
        else:
            f_scores["收入增速加快"] = 0; f_details["收入增速加快"] = "数据不足"

        if len(roe_vals) >= 2:
            f_scores["ROE改善"] = 1 if roe_vals[-1] > roe_vals[-2] else 0
            f_details["ROE改善"]= f"{roe_vals[-2]}% → {roe_vals[-1]}%"
        else:
            f_scores["ROE改善"] = 0; f_details["ROE改善"] = "数据不足"

        ni_pos = sum(1 for v in ni_vals if v > 0)
        f_scores["持续盈利（多年）"] = 1 if ni_pos >= 3 else 0
        f_details["持续盈利（多年）"]= f"近{len(ni_vals)}年中有{ni_pos}年净利润为正"

        total_f = sum(f_scores.values())
        f_color = "#2e7d32" if total_f >= 7 else ("#f57c00" if total_f >= 4 else "#c62828")
        f_label = "优质企业 ✅" if total_f >= 7 else ("普通 ⚠️" if total_f >= 4 else "高风险 🚨")

        st.markdown(f"""
        <div style="background:{f_color};color:white;padding:16px 24px;border-radius:12px;margin:8px 0;">
            <span style="font-size:2em;font-weight:bold;">{total_f} / 9</span>
            &nbsp;&nbsp;<span style="font-size:1.2em;">{f_label}</span>
        </div>""", unsafe_allow_html=True)

        with st.expander("查看9项指标明细"):
            groups = {
                "盈利能力": ["净利润为正","ROE为正","净利润同比增长","持续盈利（多年）"],
                "财务杠杆": ["资产负债率下降","净利率改善"],
                "运营效率": ["收入同比增长","收入增速加快","ROE改善"],
            }
            for group, keys in groups.items():
                st.markdown(f"**{group}**")
                for k in keys:
                    if k in f_scores:
                        icon = "✅" if f_scores[k] == 1 else "❌"
                        st.markdown(f"&nbsp;&nbsp;{icon} **{k}**：{f_details[k]}")

        # ── 地雷预警 ──────────────────────────────────────────
        st.markdown("---")
        st.markdown("### 🚨 地雷预警指标")

        warnings = []
        cautions = []
        nm_vals = col_vals(nm_col)
        dr_vals = col_vals(dr_col)

        if len(rev_vals) >= 3 and len(ni_vals) >= 3:
            if (rev_vals[-1]-rev_vals[-3]) > 0 and (ni_vals[-1]-ni_vals[-3]) < 0:
                warnings.append({"title":"💸 收入增长但利润下滑",
                    "detail":f"收入变化：+{round(rev_vals[-1]-rev_vals[-3],2)}，利润变化：{round(ni_vals[-1]-ni_vals[-3],2)}",
                    "explain":"可能存在成本失控、毛利率压缩或隐性损失"})

        if len(roe_vals) >= 3 and roe_vals[-1] < roe_vals[-2] < roe_vals[-3]:
            cautions.append({"title":"📉 ROE连续三年下滑",
                "detail":f"{roe_vals[-3]}% → {roe_vals[-2]}% → {roe_vals[-1]}%",
                "explain":"盈利能力持续恶化，需关注核心竞争力是否减弱"})

        if dr_now and dr_now > 70:
            if dr_prev and dr_now > dr_prev:
                warnings.append({"title":"⚠️ 高负债且持续上升",
                    "detail":f"资产负债率：{dr_prev}% → {dr_now}%（超70%警戒线）",
                    "explain":"偿债压力大，若遇行业下行风险较高"})
            else:
                cautions.append({"title":"⚠️ 资产负债率偏高",
                    "detail":f"当前：{dr_now}%（超70%警戒线）",
                    "explain":"需关注现金流是否能覆盖利息支出"})

        if len(nm_vals) >= 2:
            nm_chg = nm_vals[-1] - nm_vals[-2]
            if nm_chg < -5:
                warnings.append({"title":"📉 净利率单年骤降超5个百分点",
                    "detail":f"{nm_vals[-2]}% → {nm_vals[-1]}%",
                    "explain":"可能存在大额减值、竞争加剧或成本突增"})
            elif nm_chg < -3:
                cautions.append({"title":"📉 净利率明显下滑",
                    "detail":f"{nm_vals[-2]}% → {nm_vals[-1]}%",
                    "explain":"盈利质量有所下降，需了解原因"})

        if len(ni_vals) >= 2:
            if ni_vals[-1] < 0 and ni_vals[-2] < 0:
                warnings.append({"title":"🚨 连续两年亏损",
                    "detail":f"近两年净利润：{ni_vals[-2]}、{ni_vals[-1]}",
                    "explain":"持续亏损消耗净资产，警惕退市风险"})
            elif ni_vals[-1] < 0:
                cautions.append({"title":"⚠️ 最新年度亏损",
                    "detail":f"最新净利润：{ni_vals[-1]}",
                    "explain":"需判断是一次性还是趋势性问题"})

        if len(rev_vals) >= 3 and rev_vals[-1] < rev_vals[-2] < rev_vals[-3]:
            cautions.append({"title":"📉 营业收入连续三年下滑",
                "detail":f"{round(rev_vals[-3],2)} → {round(rev_vals[-2],2)} → {round(rev_vals[-1],2)}",
                "explain":"收入持续萎缩，关注行业景气度或市场份额变化"})

        if not warnings and not cautions:
            st.success("✅ 未发现明显财务风险信号，各项指标相对健康。")
        else:
            if warnings:
                st.markdown("#### 🔴 红色预警（需重点关注）")
                for w in warnings:
                    st.markdown(f"""<div style="background:#fff3f3;border-left:4px solid #c62828;
                    padding:12px 16px;border-radius:6px;margin:8px 0;">
                    <b>{w['title']}</b><br>
                    <span style="color:#555;font-size:0.9em;">{w['detail']}</span><br>
                    <span style="color:#888;font-size:0.85em;">💡 {w['explain']}</span>
                    </div>""", unsafe_allow_html=True)
            if cautions:
                st.markdown("#### 🟡 黄色注意（值得关注）")
                for c in cautions:
                    st.markdown(f"""<div style="background:#fffde7;border-left:4px solid #f9a825;
                    padding:12px 16px;border-radius:6px;margin:8px 0;">
                    <b>{c['title']}</b><br>
                    <span style="color:#555;font-size:0.9em;">{c['detail']}</span><br>
                    <span style="color:#888;font-size:0.85em;">💡 {c['explain']}</span>
                    </div>""", unsafe_allow_html=True)

        # ── AI 综合诊断 ───────────────────────────────────────
        st.markdown("---")
        st.markdown("### 🤖 AI 综合诊断")

        if st.button("🔍 生成AI诊断报告", key="ai_scan_btn"):
            with st.spinner("AI正在分析财务数据..."):
                try:
                    scan_summary = f"""公司：{cname_sc}
Piotroski F-Score：{total_f}/9（{f_label}）
红色预警（{len(warnings)}项）：{'; '.join([w['title']+': '+w['detail'] for w in warnings]) if warnings else '无'}
黄色注意（{len(cautions)}项）：{'; '.join([c['title']+': '+c['detail'] for c in cautions]) if cautions else '无'}
最新净利率：{nm_now}% | ROE：{roe_now}% | 资产负债率：{dr_now}%"""

                    ai_resp = client.messages.create(
                        model="claude-opus-4-6", max_tokens=500,
                        messages=[{"role":"user","content":f"""你是资深财务分析师。以下是财务健康扫描结果：

{scan_summary}

请用中文给出：
1. 综合评价（2-3句，直接说结论）
2. 最值得关注的1-2个风险点
3. 投资者应进一步核查的1-2个问题

专业但易懂，总字数200字以内。"""}])
                    st.session_state["ai_scan_result"] = ai_resp.content[0].text
                except Exception as e:
                    st.error(f"AI诊断失败：{e}")

        if st.session_state.get("ai_scan_result"):
            st.markdown(f"""<div style="background:#f0f4fa;border-left:4px solid #2E5C8A;
            padding:16px 20px;border-radius:8px;margin:8px 0;">
            {st.session_state['ai_scan_result'].replace(chr(10),'<br>')}
            </div>""", unsafe_allow_html=True)